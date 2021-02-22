"""Copyright 2020 Oakwood Technologies BVBA"""

from .utilities import activity, only_supported_for, interpret_path
import selenium.webdriver
import logging

"""
Decisions
Icon: la la-database
"""

@activity
def internal_decision(path, variable_names, variable_values, output_variable_name):
    """Internal decision engine

    Evaluate a decision table with pyDMNrules decision engine.

    :parameter path: File path to the decision table (.xlsx Excel worksheet)
    :type path: string

    :parameter variable_names: Names of the input variables of the decision (comma separated list)
    :type variable_names: list of strings

    :parameter variable_values: Values of the corresponding variables of the decision (comma separated list)
    :type variable_values: list of strings

    :parameter output_variable_name: Output variable name of the decision table
    :type output_variable_name: string

    :return: Decision result
    :rtype: any

    Keywords
        decision, decision engine, decision table

    Icon
        la la-th
    """
    import pyDMNrules

    dmnRules = pyDMNrules.DMN()

    if len(variable_names) != len(variable_values):
        raise Exception('Same number of input variable names and values required')

    status = dmnRules.load(path)
    if 'errors' in status:
        raise Exception('{} has errors: {}'.format(path, str(status['errors'])))

    data = {}
    for i, name in enumerate(variable_names):
        data[name] = variable_values[i]

    return dmnRules.decide(data)[1]['Result'][output_variable_name]


@activity
def camunda_decision_engine(camunda_engine_URL, decision_key, variable_names, variable_values, output_variable_name):
    """Camunda decision service

    Evaluate a decision table with camunda decision engine.

    :parameter camunda_engine_URL: URL to camunda engine (e.g. http://localhost:8080/engine-rest/)
    :type camunda_engine_URL: string

    :parameter decision_key: Key to identify decision "Decision_0tgwupa"
    :type decision_key: strings

    :parameter variable_names: Names of the input variables of the decision (comma separated list)
    :type variable_names: list of strings

    :parameter variable_values: Values of the corresponding variables of the decision (comma separated list)
    :type variable_values: list of strings

    :parameter output_variable_name: Output variable name of the decision table
    :type output_variable_name: string

    :return: Decision result
    :rtype: any

    Keywords
        decision, decision engine, decision table

    Icon
        la la-server
    """
    import requests

    if len(variable_names) != len(variable_values):
        raise Exception('Same number of input variable names and values required')

    variables = {}
    for i, name in enumerate(variable_names):
        variables[name] = { "value": variable_values[i] }

    task = {
        "variables" : variables
    }

    response = requests.post('{}decision-definition/key/{}/evaluate'.format(camunda_engine_URL, decision_key), json=task)
    
    return list(map(lambda x: x[output_variable_name]["value"], response.json()))


@activity
def human_decision(message, choices):
    """Human decision

    Allows a human to select options from a list

    :parameter message: Message the user gets prompted
    :type message: string

    :parameter choices: Options from the user can select
    :type choices: list of strings

    :return: Decision result
    :rtype: any

    Keywords
        decision, human decision, prompt, multiselect

    Icon
        la la-user
    """
    from easygui import multchoicebox

    title = "Human decision required"

    return multchoicebox(message, title, choices)


"""
Cryptography
Icon: las la-shield-alt
"""

@activity
def generate_random_key():
    """Random key

    Generate random Fernet key. Fernet guarantees that a message encrypted using it cannot be manipulated or read without the key. Fernet is an implementation of symmetric (also known as “secret key”) authenticated cryptography

    :return: Random key
    :rtype: bytes

        :Example:

    >>> # Generate a random key
    >>> generate_random_key()
    b'AYv6ZPVgnrUtHDbGZqAopRyAo9r0_UKrA2Rm3K_NjIo='

    Keywords
        random, key, fernet, hash, security, cryptography, password, secure

    Icon
        las la-key
    """
    import os
    from cryptography.fernet import Fernet

    key = Fernet.generate_key()

    return key


@activity
def encrypt_text_with_key(text, key):
    """Encrypt text 

    Encrypt text with (Fernet) key, 

    :parameter text: Text to be encrypted
    :type text: string
    :parameter key: Fernet Encryption key
    :type key: bytes

    :return: Encrypted text
    :rtype: bytes

        :Example:

    >>> # Generate a random key
    >>> key = generate_random_key()
    >>> # Encrypt text with this key
    >>> encrypt_text_with_key('Sample text', key)
    b'gAAAAABd8lpG8fNqcj5eXrPPHlx4KeCm-1TgX3jkyhStMfIlgGImIa-qaINZAj8XcxPcG8iu84iT56b_qAW9c5qpe7btUFhtxQ=='

    Keywords
        random, encryption, secure, security, hash, password, fernet, text

    Icon
        las la-lock
    """
    from cryptography.fernet import Fernet

    f = Fernet(key)

    return f.encrypt(text.encode("utf-8"))


@activity
def decrypt_text_with_key(encrypted_text, key):
    """Decrypt text

    Dexrypt bytes-like object to string with (Fernet) key

    :parameter encrypted_text: Text to be encrypted.
    :type encrypted: bytes
    :parameter key: Fernet Encryption key
    :type key: bytes

    :return: Decrypted text
    :rtype: string

        :Example:

    >>> # Generate a random key
    >>> key = generate_random_key()
    >>> # Encrypt text with generated key
    >>> encrypted_text = encrypt_text_with_key('Sample text', key)
    >>> # Decrypt text with same key
    >>> decrypt_text_with_key(encrypted_text, key)
    'Sample text'

    Keywords
        decrypt, random, unlock, un-lock hash, security, cryptography, password, secure, hash, text

    Icon
        las la-lock-open
    """
    from cryptography.fernet import Fernet

    f = Fernet(key)

    return f.decrypt(encrypted_text).decode("utf-8")


@activity
def encrypt_file_with_key(input_path, key, output_path=None):
    """Encrypt file 

    Encrypt file with (Fernet) key. Note that file will be unusable unless unlocked with the same key.

    :parameter inputh_path: Path to file to be encrypted
    :type input_path: input_file
    :parameter key: Fernet Encryption key
    :type key: bytes
    :parameter output_path: Output path, defaults to the same directory with "_encrypted" added to the name
    :type output_path: output_file, optional

    :return: Path to encrypted file
    :rtype: path

        :Example:

    >>> # Generate a random key
    >>> key = generate_random_key()
    >>> # Create a text file to illustrate file encryption
    >>> text_file_path = make_text_file()
    >>> # Encrypt the text file
    >>> encrypt_file_with_key(text_file_path, key=key)
    'C:\\Users\\<username>\\generated_text_file_encrypted.txt'

    Keywords
        encrypt, random, password, secure, secure file, lock

    Icon
        las la-lock
    """

    # Set path if not specified
    input_path = interpret_path(input_path, required=True)
    if not output_path:
        output_path = interpret_path(input_path, addition="_encrypted")
    else:
        output_path = interpret_path(output_path)

    from cryptography.fernet import Fernet

    with open(input_path, "rb") as f:
        data = f.read()

    fernet = Fernet(key)
    encrypted = fernet.encrypt(data)

    with open(output_path, "wb") as f:
        f.write(encrypted)

    return output_path


@activity
def decrypt_file_with_key(input_path, key, output_path=None):
    """Decrypt file

    Decrypts file with (Fernet) key

    :parameter input_file: Bytes-like file to be decrypted.
    :type input_file: input_file
    :parameter key: Path where key is stored.
    :type key: bytes
    :parameter output_file: Outputfile, make sure to give this the same extension as basefile before encryption. Default is the same directory with "_decrypted" added to the name 
    :type output_path: output_file, optional

    :return: Path to decrypted file

        :Example:

    >>> # Generate a random key
    >>> key = generate_random_key()
    >>> # Create a text file to encrypt file
    >>> text_file_path = make_text_file()
    >>> # Encrypt the text file
    >>> encrypted_text_file = encrypt_file_with_key(text_file_path, key=key)
    >>> # Decrypt the newly encrypted file
    >>> decrypt_file_with_key(encrypted_text_file, key=key)
    'C:\\Users\\<username>\\generated_text_file_encrypted_decrypted.txt'

    Keywords
        decrypt, random, password, secure, secure file, unlock

    Icon
        las la-lock-open
    """
    # Set path if not specified
    import os

    input_path = interpret_path(input_path, required=True)
    if not output_path:
        output_path = interpret_path(input_path, addition="_decrypted")
    else:
        output_path = interpret_path(output_path)

    from cryptography.fernet import Fernet

    with open(input_path, "rb") as f:
        data = f.read()

    fernet = Fernet(key)
    decrypted = fernet.decrypt(data)

    with open(output_path, "wb") as f:
        f.write(decrypted)

    return output_path


@activity
def generate_key_from_password(password, salt=None):
    """Key from password

    Generate key based on password and salt. If both password and salt are known the key can be regenerated.

    :parameter password: Passwords
    :type password: string
    :parameter salt: Salt to generate key in combination with password. Default value is the hostname. Take in to account that hostname is necessary to generate key, e.g. when files are encrypted with salt 'A' and password 'B', both elements are necessary to decrypt files.
    :type salt: string, optional

    :return: Bytes-like object

        :Example:

    >>> # Generate a key from password
    >>> key = generate_key_from_password(password='Sample password')
    b'7jGGF5w_xyI0CIZGCmLlnNyUvFpNvIUY08JCHopgAmm8='

    Keywords
        random, key, fernet, hash, security, cryptography, password, secure, salt

    Icon
        las la-lock
    """
    import base64
    from cryptography.hazmat.backends import default_backend
    from cryptography.hazmat.primitives import hashes
    from cryptography.hazmat.primitives.kdf.pbkdf2 import PBKDF2HMAC
    import socket

    # If no salt is set, use hostname as salt
    if not salt:
        salt = socket.gethostname().encode("utf-8")

    kdf = PBKDF2HMAC(
        algorithm=hashes.SHA256(),
        length=32,
        salt=salt,
        iterations=500000,
        backend=default_backend(),
    )
    key = base64.urlsafe_b64encode(kdf.derive(password.encode("utf-8")))

    return key


@activity
def generate_hash_from_file(input_path, method="md5", buffer_size=65536):
    """Hash from file

    Generate hash from file 

    Can be used to create unique identifier for file validation or comparison. Please note that MD5 and SHA1 are not cryptographically secure.

    :parameter input_path: File to hash
    :type input_path: input_file
    :parameter method: Method for hashing, choose between 'md5', 'sha256' and 'blake2b'. Note that different methods generate different hashes. Default method is 'md5'.
    :options method: ['md5', 'sha256', 'blake2b']
    :parameter buffer_size: Buffer size for reading file in chunks, default value is 64kb
    :type buffer_size: int, optional

    :return: Bytes-like object

        :Example:

    >>> # Generate a text file to illustrate hash
    >>> text_file_path = make_text_file()
    >>> # Get hash from text file
    >>> generate_hash_from_file(text_file_path)
    '1ba249ca5931f3c85fe44d354c2f274d'

    Keywords
        hash, mdf5, sha256, blake2b, identifier, unique, hashing, fingerprint, comparison

    Icon
        las la-fingerprint

    """
    import sys
    import hashlib

    # Arbitrary buffer size. 64kb for compatibility with most systems
    buffer_size = 65536

    if method == "md5":
        hash_list = hashlib.md5()  # nosec
    if method == "sha256":
        hash_list = hashlib.sha1()  # nosec
    if method == "blake2b":
        hash_list = hashlib.blake2b()

    with open(input_path, "rb") as f:
        while True:
            data = f.read(buffer_size)
            if data:
                hash_list.update(data)
            else:
                return hash_list.hexdigest()


@activity
def generate_hash_from_text(text, method="md5"):
    """Hash from text

    Generate hash from text. Keep in mind that MD5 is not cryptographically secure.

    :parameter text: Text to hash
    :type text: string
    :parameter method: Method for hashing, choose between 'md5', 'sha256' and 'blake2b'. Note that different methods generate different hashes. Default method is 'md5'.
    :options method: ['md5', 'sha256', 'blake2b']

        :Example:

    >>> # Generate a hast from text
    >>> generate_hash_from_text('Sample text')
    '1ba249ca5931f3c85fe44d354c2f274d'

    Keywords
        Hash, mdf5, sha256, blake2b, identifier, unique, hashing, fingerprint, text, comparison

    Icon
        las la-fingerprint
    """
    import sys
    import hashlib

    encoded_text = text.encode("utf-8")

    if method == "md5":
        return hashlib.md5(encoded_text).hexdigest()  # nosec
    if method == "sha256":
        return hashlib.sha256(encoded_text).hexdigest()
    if method == "blake2b":
        return hashlib.blake2b(encoded_text).hexdigest()


"""
Random
Icon: las la-dice-d6
"""


@activity
def generate_random_number(lower_limit=0, upper_limit=100, fractional=False):
    """Random number

    Random numbers can be integers (not a fractional number) or a float (fractional number).

    :parameter lower_limit: Lower limit for random number
    :type lower_limit: int, optional
    :parameter upper_limit: Upper limit for random number
    :type upper_limit: int, optional
    :parameter fractional: Setting this to True will generate fractional number. Default value is False and only generates whole numbers.
    :type fractional: bool, optional

    :return: Random integer or float

        :Example:

    >>> # Generate a random number
    >>> generate_random_number()
    7

    Keywords
        random number, random integer, dice, gamble, rng, random

    Icon
        las la-dice
    """
    import random

    if fractional:
        return random.uniform(lower_limit, upper_limit)
    else:
        return random.randrange(lower_limit, upper_limit, 1)


@activity
def generate_random_data(locale=None, type=None):
    """Random data

    Generates all kinds of random data. Specifying locale changes format for some options
    
    :parameter attribute: Choose a specific characteristic or attribute from fake person
    :options type: ['email', 'food dish', 'food drink', 'fruit', 'vegetable', 'company type', 'ean code', 'imei code', 'isbn code', 'issn code', 'ip v4', 'ip v6', 'mac address', 'filename', 'filename extension', 'chemical element', 'academic degree', 'occupation', 'word', 'phone number' , 'political view', 'university']
    :parameter locale: Add a locale to generates typical data for selected locale.
    :options locale: ['cs', 'da', 'de', 'de-at', 'de-ch', 'el', 'en', 'en-au', 'en-ca', 'en-gb', 'es', 'es-mx', 'et', 'fa', 'fi', 'fr', 'hu', 'is', 'it', 'ja', 'kk', 'ko', 'nl', 'nl-be', 'no', 'pl', 'pt', 'pt-br', 'ru', 'sk', 'sv', 'tr', 'uk', 'zh']

    -	cs	    -	Czech
    -	da	    -	Danish
    -	de	    -	German
    -	de-at	-	Austrian german
    -	de-ch	-	Swiss german
    -	el	    -	Greek
    -	en	    -	English
    -	en-au	-	Australian English
    -	en-ca	-	Canadian English
    -	en-gb	-	British English
    -	es	    -	Spanish
    -	es-mx	-	Mexican Spanish
    -	et	    -	Estonian
    -	fa	    -	Farsi
    -	fi	    -	Finnish
    -	fr	    -	French
    -	hu	    -	Hungarian
    -	is	    -	Icelandic
    -	it	    -	Italian
    -	ja	    -	Japanese
    -	kk	    -	Kazakh
    -	ko	    -	Korean
    -	nl	    -	Dutch
    -	nl-be	-	Belgium Dutch
    -	no	    -	Norwegian
    -	pl	    -	Polish
    -	pt	    -	Portuguese
    -	pt-br	-	Brazilian Portuguese
    -	ru	    -	Russian
    -	sk	    -	Slovak
    -	sv	    -	Swedish
    -	tr	    -	Turkish
    -	uk	    -	Ukrainian
    -	zh	    -	Chinese

    :return: Random data as string

        :Example:

    >>> # Generate random data
    >>> generate_random_data()
    'Banana'

    Keywords
        random, lorem ipsum, gsm, cell, cellphone, telephone, mobile, number, smartphone, text generater, filler, place holder, noise, random text, random txt, text generation, fake, code, email, generate, generator, generic

    Icon
        las la-digital-tachograph
    """
    from mimesis import Generic

    if locale:
        gen = Generic(locale)
    else:
        gen = Generic()

    if type == "email":
        return gen.person.email()
    elif type == "company type":
        return gen.business.company_type()
    elif type == "food dish":
        return gen.food.dish()
    elif type == "food drink":
        return gen.food.drink()
    elif type == "fruit":
        return gen.food.fruit()
    elif type == "vegetable":
        return gen.food.vegetable()
    elif type == "ean code":
        return gen.code.ean()
    elif type == "imei code":
        return gen.code.imei()
    elif type == "isbn code":
        return gen.code.isbn()
    elif type == "issn code":
        return gen.code.issn()
    elif type == "filename":
        return gen.file.file_name()
    elif type == "filename extension":
        return gen.file.extension()
    elif type == "ip v4":
        return gen.internet.ip_v4()
    elif type == "ip v6":
        return gen.internet.ip_v6()
    elif type == "mac address":
        return gen.internet.mac_address()
    elif type == "chemical element":
        return gen.science.chemical_element()
    elif type == "academic degree":
        return gen.person.academic_degree()
    elif type == "occupation":
        return gen.person.occupation()
    elif type == "word":
        return gen.text.word()
    elif type == "phone number":
        return gen.person.telephone()
    elif type == "political view":
        return gen.person.political_views()
    elif type == "university":
        return gen.person.university()
    else:
        return gen.text.word()


@activity
def generate_random_boolean():
    """Random boolean

    Generates a random boolean (True or False)

    :return: Boolean

        :Example:

    >>> # Generate a random boolean
    >>> generate_random_boolean()
    True

    Keywords
        random, dice, gamble, rng, coin, coinflip, heads, tails

    Icon
        las la-coins
    """
    import random

    return bool(random.getrandbits(1))


@activity
def generate_random_name(locale=None, name=None):
    """Random name

    Generates a random name. Adding a locale adds a more common name in the specified locale. Provides first name and last name.

    :parameter locale: Add a locale to generate popular name for selected locale.
    :parameter name: Choose to generate first, last or full name.
    :options name: ['full', 'first', 'last']
    :options locale: ['cs', 'da', 'de', 'de-at', 'de-ch', 'el', 'en', 'en-au', 'en-ca', 'en-gb', 'es', 'es-mx', 'et', 'fa', 'fi', 'fr', 'hu', 'is', 'it', 'ja', 'kk', 'ko', 'nl', 'nl-be', 'no', 'pl', 'pt', 'pt-br', 'ru', 'sk', 'sv', 'tr', 'uk', 'zh']

    -	cs	    -	Czech
    -	da	    -	Danish
    -	de	    -	German
    -	de-at	-	Austrian german
    -	de-ch	-	Swiss german
    -	el	    -	Greek
    -	en	    -	English
    -	en-au	-	Australian English
    -	en-ca	-	Canadian English
    -	en-gb	-	British English
    -	es	    -	Spanish
    -	es-mx	-	Mexican Spanish
    -	et	    -	Estonian
    -	fa	    -	Farsi
    -	fi	    -	Finnish
    -	fr	    -	French
    -	hu	    -	Hungarian
    -	is	    -	Icelandic
    -	it	    -	Italian
    -	ja	    -	Japanese
    -	kk	    -	Kazakh
    -	ko	    -	Korean
    -	nl	    -	Dutch
    -	nl-be	-	Belgium Dutch
    -	no	    -	Norwegian
    -	pl	    -	Polish
    -	pt	    -	Portuguese
    -	pt-br	-	Brazilian Portuguese
    -	ru	    -	Russian
    -	sk	    -	Slovak
    -	sv	    -	Swedish
    -	tr	    -	Turkish
    -	uk	    -	Ukrainian
    -	zh	    -	Chinese

    :return: Random name as string

        :Example:

    >>> # Generate a random name
    >>> generate_random_name()
    'Michelle Murphy'

    Keywords
        random, dummy name, name, name generater, fake person, fake, person, surname, lastname, fake name generator

    Icon
        las la-user-tag
    """
    from mimesis import Person

    if locale:
        person = Person(locale)
    else:
        person = Person()

    if name == "first":
        return person.first_name()
    elif name == "last":
        return person.last_name()
    else:
        return person.full_name()


@activity
def generate_random_words(locale=None, type=None):
    """Random words

    Generates a random sentence. Specifying locale changes language and content based on locale.

    :parameter type: Specify type of words to generate
    :options type: ['sentence', 'quote', 'answer', 'single word', 'color', 'swear word']
    :parameter locale: Add a locale to generate text for selected locale.
    :options locale: ['cs', 'da', 'de', 'de-at', 'de-ch', 'el', 'en', 'en-au', 'en-ca', 'en-gb', 'es', 'es-mx', 'et', 'fa', 'fi', 'fr', 'hu', 'is', 'it', 'ja', 'kk', 'ko', 'nl', 'nl-be', 'no', 'pl', 'pt', 'pt-br', 'ru', 'sk', 'sv', 'tr', 'uk', 'zh']

    -	cs	    -	Czech
    -	da	    -	Danish
    -	de	    -	German
    -	de-at	-	Austrian german
    -	de-ch	-	Swiss german
    -	el	    -	Greek
    -	en	    -	English
    -	en-au	-	Australian English
    -	en-ca	-	Canadian English
    -	en-gb	-	British English
    -	es	    -	Spanish
    -	es-mx	-	Mexican Spanish
    -	et	    -	Estonian
    -	fa	    -	Farsi
    -	fi	    -	Finnish
    -	fr	    -	French
    -	hu	    -	Hungarian
    -	is	    -	Icelandic
    -	it	    -	Italian
    -	ja	    -	Japanese
    -	kk	    -	Kazakh
    -	ko	    -	Korean
    -	nl	    -	Dutch
    -	nl-be	-	Belgium Dutch
    -	no	    -	Norwegian
    -	pl	    -	Polish
    -	pt	    -	Portuguese
    -	pt-br	-	Brazilian Portuguese
    -	ru	    -	Russian
    -	sk	    -	Slovak
    -	sv	    -	Swedish
    -	tr	    -	Turkish
    -	uk	    -	Ukrainian
    -	zh	    -	Chinese

    :return: Random words as string

        :Example:

    >>> # Generate a random sentence
    >>> generate_random_words()
    'The age of automation is going to be the age of do-it-yourself'

    Keywords
        random, sentence, lorem ipsum, text generater, filler, place holder, noise, random text, random txt, text generation, nlp

    Icon
        las la-comment
    """
    from mimesis import Generic

    if locale:
        gen = Generic(locale).text
    else:
        gen = Generic().text

    if type == "sentence":
        return gen.sentence()
    elif type == "quote":
        return gen.quote()
    elif type == "answer":
        return gen.answer()
    elif type == "single word":
        return gen.word()
    elif type == "color":
        return gen.color()
    elif type == "swear word":
        return gen.swear_word()
    else:
        return gen.sentence()


@activity
def generate_random_address(locale=None, format=None):
    """Random address

    Generates a random address. Specifying locale changes random locations and streetnames based on locale.

    :parameter format: Choose a specific part or format for the address
    :options format: ['address', 'street', 'street number', 'city', 'continent', 'country', 'country code', 'postal code']
    :parameter locale: Add a locale to generate typical address for selected locale.
    :options locale: ['cs', 'da', 'de', 'de-at', 'de-ch', 'el', 'en', 'en-au', 'en-ca', 'en-gb', 'es', 'es-mx', 'et', 'fa', 'fi', 'fr', 'hu', 'is', 'it', 'ja', 'kk', 'ko', 'nl', 'nl-be', 'no', 'pl', 'pt', 'pt-br', 'ru', 'sk', 'sv', 'tr', 'uk', 'zh']

    -	cs	    -	Czech
    -	da	    -	Danish
    -	de	    -	German
    -	de-at	-	Austrian german
    -	de-ch	-	Swiss german
    -	el	    -	Greek
    -	en	    -	English
    -	en-au	-	Australian English
    -	en-ca	-	Canadian English
    -	en-gb	-	British English
    -	es	    -	Spanish
    -	es-mx	-	Mexican Spanish
    -	et	    -	Estonian
    -	fa	    -	Farsi
    -	fi	    -	Finnish
    -	fr	    -	French
    -	hu	    -	Hungarian
    -	is	    -	Icelandic
    -	it	    -	Italian
    -	ja	    -	Japanese
    -	kk	    -	Kazakh
    -	ko	    -	Korean
    -	nl	    -	Dutch
    -	nl-be	-	Belgium Dutch
    -	no	    -	Norwegian
    -	pl	    -	Polish
    -	pt	    -	Portuguese
    -	pt-br	-	Brazilian Portuguese
    -	ru	    -	Russian
    -	sk	    -	Slovak
    -	sv	    -	Swedish
    -	tr	    -	Turkish
    -	uk	    -	Ukrainian
    -	zh	    -	Chinese

    :return: Name as string

        :Example:

    >>> # Generate a random address
    >>> generate_random_address()
    '123 Robot Avenue'

    Keywords
        random, address, data, street, city, postal, dummy name, name, name generater, fake person, fake, person, surname, lastname, fake name generator

    Icon
        las la-map-marker
    """
    from mimesis import Address

    if locale:
        ad = Address(locale)
    else:
        ad = Address()

    if format == "address":
        return ad.address()
    elif format == "street":
        return ad.street_name()
    elif format == "street number":
        return ad.street_number()
    elif format == "city":
        return ad.city()
    elif format == "continent":
        return ad.continent()
    elif format == "country":
        return ad.country()
    elif format == "country code":
        return ad.country_code()
    elif format == "postal code":
        return ad.postal_code()
    else:
        return ad.address()


@activity
def generate_random_beep(max_duration=2000, max_frequency=5000):
    """Random beep

    Generates a random beep, only works on Windows

    :parameter max_duration: Maximum random duration in miliseconds. Default value is 2 miliseconds
    :type max_duration: int, optional
    :parameter max_frequency: Maximum random frequency in Hz. Default value is 5000 Hz.
    :type max_frequency: int, optional

    :return: Sound

        :Example: 

    >>> # Generate a random beep
    >>> generate_random_beep()

    Keywords
        beep, sound, random, noise, alert, notification

    Icon
        las la-volume-up
    """
    import winsound
    import random

    frequency = random.randrange(5000)
    duration = random.randrange(2000)
    winsound.Beep(frequency, duration)


@activity
def generate_random_date(formatting="%m/%d/%Y %I:%M", days_in_past=1000):
    """Random date

    Generates a random date.

    -   %a	Abbreviated weekday name.	 
    -   %A	Full weekday name.	 
    -   %b	Abbreviated month name.	 
    -   %B	Full month name.	 
    -   %c	Predefined date and time representation.	 
    -   %d	Day of the month as a decimal number [01,31].	 
    -   %H	Hour (24-hour clock) as a decimal number [00,23].	 
    -   %I	Hour (12-hour clock) as a decimal number [01,12].	 
    -   %j	Day of the year as a decimal number [001,366].	 
    -   %m	Month as a decimal number [01,12].	 
    -   %M	Minute as a decimal number [00,59].	 
    -   %p	AM or PM.
    -   %S	Second as a decimal number [00,61].	
    -   %U	Week number of the year (Sunday as the first day of the week) as a decimal number [00,53]. All days in a new year preceding the first Sunday are considered to be in week 0.	
    -   %w	Weekday as a decimal number [0(Sunday),6].	 
    -   %W	Week number of the year (Monday as the first day of the week) as a decimal number [00,53]. All days in a new year preceding the first Monday are considered to be in week 0.	
    -   %x	Predefined date representation.	 
    -   %X	Predefined time representation.	 
    -   %y	Year without century as a decimal number [00,99].	 
    -   %Y	Year with century as a decimal number.
    -   %Z	Time zone name (no characters if no time zone exists).

    :parameter days_in_past: Days in the past for which oldest random date is generated, default is 1000 days
    :type days_in_past: int, optional
    :parameter formatting: Formatting of the dates, replace with 'None' to get raw datetime format. e.g. format='Current month is %B' generates 'Current month is Januari' and format='%m/%d/%Y %I:%M' generates format 01/01/1900 00:00. 
    :type formatting: string, optional

    :return: Random date as string

        :Example: 

    >>> # Generate a random date
    >>> generate_random_date()
    01/01/2020 13:37'

    Keywords
        random, date, datetime, random date, fake date , calendar

    Icon
        las la-calendar
    """

    import random
    import datetime

    latest = datetime.datetime.now()
    earliest = latest - datetime.timedelta(days=days_in_past)
    delta_seconds = (latest - earliest).total_seconds()

    random_date = earliest + datetime.timedelta(seconds=random.randrange(delta_seconds))

    if formatting:
        return random_date.strftime(formatting)
    else:
        return random_date


@activity
def generate_date_today(formatting="%m/%d/%Y"):
    """Today's date

    Generates today's date.

    -   %a	Abbreviated weekday name.	 
    -   %A	Full weekday name.	 
    -   %b	Abbreviated month name.	 
    -   %B	Full month name.	 
    -   %c	Predefined date and time representation.	 
    -   %d	Day of the month as a decimal number [01,31].	 
    -   %H	Hour (24-hour clock) as a decimal number [00,23].	 
    -   %I	Hour (12-hour clock) as a decimal number [01,12].	 
    -   %j	Day of the year as a decimal number [001,366].	 
    -   %m	Month as a decimal number [01,12].	 
    -   %M	Minute as a decimal number [00,59].	 
    -   %p	AM or PM.
    -   %S	Second as a decimal number [00,61].	
    -   %U	Week number of the year (Sunday as the first day of the week) as a decimal number [00,53]. All days in a new year preceding the first Sunday are considered to be in week 0.	
    -   %w	Weekday as a decimal number [0(Sunday),6].	 
    -   %W	Week number of the year (Monday as the first day of the week) as a decimal number [00,53]. All days in a new year preceding the first Monday are considered to be in week 0.	
    -   %x	Predefined date representation.	 
    -   %X	Predefined time representation.	 
    -   %y	Year without century as a decimal number [00,99].	 
    -   %Y	Year with century as a decimal number.
    -   %Z	Time zone name (no characters if no time zone exists).

    :parameter formatting: Formatting of the dates, replace with 'None' to get raw datetime format. e.g. format='Current month is %B' generates 'Current month is Januari' and format='%m/%d/%Y %I:%M' generates format 01/01/1900 00:00. 
    :type formatting: string, optional

    :return: Random date as string

        :Example: 

    >>> # Generate a random date
    >>> generate_date_today()
    '01/01/2022'

    Keywords
        random, date, today, now, today date, time, datetime, random date, fake date , calendar

    Icon
        las la-calendar
    """

    import datetime

    today = datetime.datetime.now()

    if formatting:
        return today.strftime(formatting)
    else:
        return today


@activity
def generate_unique_identifier():
    """Generate unique identifier

    Generates a random UUID4 (universally unique identifier). While the probability that a UUID will be duplicated is not zero, it is close enough to zero to be negligible.

    :return: Identifier as string

        :Example:

    >>> # Generate unique identifier
    >>> generate_unique_identifier()
    'd72fd7ea-d682-4f78-8ca1-0ed34142a992'

    Keywords
        unique, identifier, primary key, random

    Icon
        las la-random
    """
    from uuid import uuid4

    return str(uuid4())


"""
Output
Icon: lab la-wpforms
"""


@activity
def display_osd_message(message="Example message", seconds=5):
    """Display overlay message

    Display custom OSD (on-screen display) message. Can be used to display a message for a limited amount of time. Can be used for illustration, debugging or as OSD.

    :parameter message: Message to be displayed
    :type message: string, optional
    :parameter seconds: Duration in seconds for message to be displayed
    :type seconds: int, optional

        :Example:

    >>> # Display overlay message
    >>> display_osd_message()

    Keywords
        message box, osd, overlay, info warning, info, popup, window, feedback, screen, login, attended

    Icon
        las la-tv
    """
    only_supported_for("Windows")

    if "DISABLE_AUTOMAGICA_OSD" in globals():
        return

    from threading import Thread

    def load_osd():
        import tkinter
        import win32con
        import pywintypes
        import win32api

        screen_width = win32api.GetSystemMetrics(0)
        screen_height = win32api.GetSystemMetrics(1)

        root = tkinter.Tk()
        label = tkinter.Label(
            text=message,
            font=("Helvetica", "30"),
            fg="white",
            bg="black",
            borderwidth=10,
        )
        label.master.overrideredirect(True)
        label.config(anchor=tkinter.CENTER)
        label.master.geometry(
            "+{}+{}".format(int(screen_width / 2), int(screen_height - 250))
        )
        label.master.lift()
        label.master.wm_attributes("-topmost", True)
        label.master.wm_attributes("-disabled", True)
        label.master.wm_attributes("-transparentcolor", "black")

        hWindow = pywintypes.HANDLE(int(label.master.frame(), 16))

        exStyle = (
            win32con.WS_EX_COMPOSITED
            | win32con.WS_EX_LAYERED
            | win32con.WS_EX_NOACTIVATE
            | win32con.WS_EX_TOPMOST
            | win32con.WS_EX_TRANSPARENT
        )
        win32api.SetWindowLong(hWindow, win32con.GWL_EXSTYLE, exStyle)

        label.after(seconds * 1000, lambda: root.destroy())
        label.pack()
        label.mainloop()

    t = Thread(target=load_osd)

    try:
        t.start()
    except Exception:
        logging.exception()

    finally:
        try:
            t.kill()
        except Exception:
            logging.Exception()


@activity
def print_console(data="Example print"):
    """Print message in console

    Print message in console. Can be used to display data in the Automagica Flow console

    :parameter data: Data to be printed
    :type data: string, optional

        :Example:

    >>> # Print in console
    >>> print_console()

    Keywords
        print, box, osd, data, debugging info, popup, window, feedback, screen, login, attended

    Icon
        las la-tv
    """
    print(data)


"""
Browser
Icon: lab la-chrome
"""


class Chrome(selenium.webdriver.Chrome):
    @activity
    def __init__(
        self,
        load_images=True,
        headless=False,
        incognito=False,
        disable_extension=False,
        maximize_window=True,
        focus_window=True,
        auto_update_chromedriver=False,
    ):
        """Open Chrome Browser

        Open the Chrome Browser with the Selenium webdriver. Canb be used to automate manipulations in the browser.
        Different elements can be found as:

        -   Xpath: e.g. browser.find_element_by_xpath() or browser.xpath()
        One can easily find an xpath by right clicking an element -> inspect. Look for the element in the menu and right click -> copy -> xpath
        find_element_by_id
        -   Name: find_element_by_name
        -   Link text: find_element_by_link_text
        -   Partial link text: find_element_by_partial_link_text
        -   Tag name: find_element_by_tag_name
        -   Class name: find_element_by_class_name
        -   Css selector: find_element_by_css_selector

        Elements can be manipulated by:

        - Clicking: e.g. element.click()
        - Typing: e.g. element.send_keys()

        :parameter load_images: Do not load images (bool). This could speed up loading pages
        :type load_images: bool, optional
        :parameter headless: Run headless, this means running without a visible window (bool)
        :type headless: bool, optional
        :parameter incognito: Run in incognito mode
        :type incognito: bool, optional
        :parameter disable_extension: Disable extensions
        :type disable_extension: bool, optional
        :parameter auto_update_chromedriver: Automatically update Chromedriver
        :type auto_update_chromedriver: bool, optional

        return: wWbdriver: Selenium Webdriver

            :Example:

        >>> # Open the browser
        >>> browser = Chrome()
        >>> # Go to a website
        >>> browser.get('https://automagica.com')
        >>> # Close browser
        >>> browser.quit()

        Keywords
            chrome, browsing, browser, internet, surfing, web, webscraping, www, selenium, crawling, webtesting, mozilla, firefox, internet explorer

        Icon
            lab la-chrome

        """
        import platform
        import os

        def download_latest_driver(chromedriver_path):
            # Downloads latest Chrome driver on Windows
            import subprocess  # nosec
            from automagica.httpclient import http_client
            import os
            from io import BytesIO
            import zipfile
            import shutil

            try:
                driver_path = (
                    os.path.abspath(__file__).replace(
                        os.path.basename(os.path.realpath(__file__)), ""
                    )
                    + chromedriver_path
                )

                if os.path.exists(driver_path):

                    current_version = str(
                        subprocess.check_output(
                            ["cmd.exe", "/c", str(driver_path + " --v")]
                        )
                    )
                    latest_version = http_client.get(
                        "https://chromedriver.storage.googleapis.com/LATEST_RELEASE"
                    ).text

                    if latest_version in current_version:
                        return

                    request = http_client.get(
                        "https://chromedriver.storage.googleapis.com/"
                        + str(latest_version)
                        + "/chromedriver_win32.zip"
                    )

                    file = zipfile.ZipFile(BytesIO(request.content))
                    shutil.rmtree(os.path.dirname(driver_path))
                    os.makedirs(os.path.dirname(driver_path))
                    file.extractall(os.path.dirname(driver_path))
                    return

                else:

                    latest_version = http_client.get(
                        "https://chromedriver.storage.googleapis.com/LATEST_RELEASE"
                    ).text

                    request = http_client.get(
                        "https://chromedriver.storage.googleapis.com/"
                        + str(latest_version)
                        + "/chromedriver_win32.zip"
                    )

                    file = zipfile.ZipFile(BytesIO(request.content))
                    if not os.path.exists(driver_path):
                        os.makedirs(os.path.dirname(driver_path))
                    file.extractall(os.path.dirname(driver_path))

            except Exception:
                raise Exception

        # Check what OS we are on
        if platform.system() == "Linux":
            chromedriver_path = "bin/linux64/chromedriver"
        elif platform.system() == "Windows":
            chromedriver_path = "\\bin\\win32\\chromedriver.exe"
            if auto_update_chromedriver:
                download_latest_driver(chromedriver_path)
        else:
            chromedriver_path = "bin/mac64/chromedriver"

        chrome_options = selenium.webdriver.ChromeOptions()
        if incognito:
            chrome_options.add_argument("--incognito")
        if disable_extension:
            # To disable the error message popup: "Loading of unpacked extensions is disabled by the administrator"
            chrome_options.add_experimental_option("useAutomationExtension", False)
        if headless:
            chrome_options.add_argument("--headless")

        if not load_images:
            prefs = {"profile.managed_default_content_settings.images": 2}
            chrome_options.add_experimental_option("prefs", prefs)

        selenium.webdriver.Chrome.__init__(
            self,
            os.path.abspath(__file__).replace(
                os.path.basename(os.path.realpath(__file__)), ""
            )
            + chromedriver_path,
            chrome_options=chrome_options,
        )

        if maximize_window:
            self.maximize_window()

        if focus_window:
            self.switch_to.window(self.current_window_handle)

    @activity
    def save_all_images(self, output_path=None):
        """Save all images

        Save all images on current page in the Browser

        :parameter output_path: Path where images can be saved. Default value is home directory.
        :type output_path: output_dir, optional

        :return: List with paths to images

            :Example:

        >>> # Open the browser
        >>> browser = Chrome()
        >>> # Go to a website
        >>> browser.get('https://www.nytimes.com/')
        >>> # Save all images
        >>> browser.save_all_images()
        >>> browser.quit()
        ['C:\\Users\\<username>\\image1.png', 'C:\\Users\\<username>\\image2.jpg', 'C:\\Users\\<username>\\image4.gif']

        Keywords
            image scraping, chrome, internet, browsing, browser, surfing, web, webscraping, www, selenium, crawling, webtesting, mozilla, firefox, internet explorer

        Icon
            las la-images

        """
        from automagica.httpclient import http_client
        import os
        from urllib.parse import urlparse

        output_path = interpret_path(output_path)

        paths = []

        images = self.find_elements_by_tag_name("img")

        for image in images:
            url = image.get_attribute("src")
            a = urlparse(url)
            filename = os.path.basename(a.path)

            if filename:
                with open(os.path.join(output_path, filename), "wb") as f:
                    try:
                        r = http_client.get(url)
                        f.write(r.content)
                        paths.append(os.path.join(output_path, filename))
                    except Exception:
                        logging.exception()

        return paths

    @activity
    def browse_to(self, url):
        """Browse to URL

        Browse to URL.

        :parameter url: Url to browser to
        :type url: string

        :return: Webpage

            :Example:

        >>> # Open the browser
        >>> browser = Chrome()
        >>> # Go to a website
        >>> browser.browse_to('https://nytimes.com')

        Keywords
            chrome, element, browse to, browse, surf, surf to, go to, get, internet, browsing, browser, surfing, web, webscraping, www, selenium, crawling, webtesting, mozilla, firefox, internet explorer

        Icon
            lab la-chrome

        """
        return self.get(url)

    @activity
    def find_elements_by_text(self, text):
        """Find elements by text

        Find all elements by their text. Text does not need to match exactly, part of text is enough.

        :parameter text: Text to find elements by
        :type text: string

        :return: Elements that matched with text

            :Example:

        >>> # Open the browser
        >>> browser = Chrome()
        >>> # Go to a website
        >>> browser.get('https://nytimes.com')
        >>> # Find elements by text
        >>> browser.find_elements_by_text('world')
        [webelement1, webelement2 , .. ]

        Keywords
            element, element by text, chrome, internet, browsing, browser, surfing, web, webscraping, www, selenium, crawling, webtesting, mozilla, firefox, internet explorer

        Icon
            las la-align-center

        """
        return self.find_elements_by_xpath(
            "//*[contains(text(), '"
            + text.lower()
            + "')] | //*[@value='"
            + text.lower()
            + "']"
        )

    @activity
    def find_all_links(self, contains=""):
        """Find all links

        Find all links on a webpage in the browser

        :parameter contains: Criteria of substring that url must contain to be included
        :type contains: string, optional

        :return: Links

            :Example:

        >>> # Open the browser
        >>> browser = Chrome()
        >>> # Go to a website
        >>> browser.get('https://nytimes.com')
        >>> # Find elements by text
        >>> browser.find_all_links()
        [webelement1, webelement2 , .. ]

        Keywords
            random, element,link, links element by text, chrome, internet, browsing, browser, surfing, web, webscraping, www, selenium, crawling, webtesting, mozilla, firefox, internet explorer

        Icon
            las la-window-restore
        """
        links = []
        for element in self.find_elements_by_xpath("//a[@href]"):
            try:
                href_el = element.get_attribute("href")
                if contains:
                    if contains in element.get_attribute("href"):
                        links.append(element.get_attribute("href"))
                else:
                    links.append(href_el)
            except:
                pass

        if links:
            return links

    @activity
    def find_first_link(self, contains=None):
        """Find first link on a webpage

        Find first link on a webpage

        :parameter contains: Criteria of substring that url must contain to be included
        :type contains: string, optional

        :return: First link

            :Example:

        >>> # Open the browser
        >>> browser = Chrome()
        >>> # Go to a website
        >>> browser.get('https://nytimes.com')
        >>> # Find elements by text
        >>> browser.find_first_link()


        Keywords
            random, link, links, element, element by text, chrome, internet, browsing, browser, surfing, web, webscraping, www, selenium, crawling, webtesting, mozilla, firefox, internet explorer

        Icon
            las la-window-restore
        """
        for element in self.find_elements_by_xpath("//a[@href]"):
            try:
                href_el = element.get_attribute("href")
                if contains:
                    if contains in element.get_attribute("href"):
                        return element.get_attribute("href")
                else:
                    return element.get_attribute("href")
            except:
                pass

    @activity
    def get_text_on_webpage(self):
        """Get all text on webwpage

        Get all the raw body text from current webpage

        :return: Text

            :Example:

        >>> # Open the browser
        >>> browser = Chrome()
        >>> # Go to a website
        >>> browser.get('https://nytimes.com')
        >>> # Get text from page
        >>> browser.get_text_on_webpage()

        Keywords
            random, link, links, element, element by text, chrome, internet, browsing, browser, surfing, web, webscraping, www, selenium, crawling, webtesting, mozilla, firefox, internet explorer

        Icon
            las la-window-restore
        """

        return self.find_element_by_tag_name("body").text

    @activity
    def highlight(self, element):
        """Highlight element

        Highlight elements in yellow in the browser

        :parameter element: Element to highlight
        :type element: selenium.webdriver.remote.webelement.WebElement

            :Example:

        >>> # Open the browser
        >>> browser = Chrome()
        >>> # Go to a website
        >>> browser.get('https://wikipedia.org')
        >>> # Find first link on page
        >>> first_link = browser.find_elements_by_xpath("//a[@href]")[0]
        >>> # Highlight first link
        >>> browser.highlight(first_link)

        Keywords
            element, element by text, chrome, internet, browsing, browser, surfing, web, webscraping, www, selenium, crawling, webtesting, mozilla, firefox, internet explorer

        Icon
            las la-highlighter

        """
        driver = element._parent

        def apply_style(s):
            driver.execute_script(
                "arguments[0].setAttribute('style', arguments[1]);", element, s
            )

        apply_style("background: yellow; border: 2px solid red;")

    @activity
    def exit(self):
        """Exit the browser

        Quit the browser by exiting gracefully. One can also use the native 'quit' function

            :Example:

        >>> # Open the browser
        >>> browser = Chrome()
        >>> # Go to a website
        >>> browser.get('https://automagica.com')
        >>> # Close browser
        >>> browser.exit()


        Keywords
            quit, exit, close, element, element by text, chrome, internet, browsing, browser, surfing, web, webscraping, www, selenium, crawling, webtesting, mozilla, firefox, internet explorer

        Icon
            las la-window-close

        """
        self.quit()

    @activity
    def by_xpaths(self, element):
        """Find all XPaths

        Find all elements with specified xpath on a webpage in the the browser. Can also use native 'find_elements_by_xpath' 

        :parameter element: Xpath of element
        :type element: string, optional

        :return: Element by xpaths

            :Example:

        >>> # Open the browser
        >>> browser = Chrome()
        >>> # Go to a website
        >>> browser.get('https://wikipedia.org')
        >>> # Find elements by xpaths
        >>> browser.by_xpaths('//*[@id=\'js-link-box-en\']')
        [webelement1, webelement2 , .. ]

        Keywords
            random, element, xpath, xml, element by text, chrome, internet, browsing, browser, surfing, web, webscraping, www, selenium, crawling, webtesting, mozilla, firefox, internet explorer

        Icon
            las la-times

        """
        return self.find_elements_by_xpath(element)

    @activity
    def by_xpath(self, element):
        """Find XPath in browser

        Find all element with specified xpath on a webpage in the the browser. Can also use native 'find_elements_by_xpath' 

        :parameter element: Xpath of element
        :type element: string, optional

        :return: Element by xpath

            :Example:

        >>> # Open the browser
        >>> browser = Chrome()
        >>> # Go to a website
        >>> browser.get('https://wikipedia.org')
        >>> # Find element by xpath
        >>> element = browser.by_xpath('//*[@id=\'js-link-box-en\']')
        >>> # We can now use this element, for example to click on
        >>> element.click()

        Keywords
            random, xpath, element, xml element by text, chrome, internet, browsing, browser, surfing, web, webscraping, www, selenium, crawling, webtesting, mozilla, firefox, internet explorer

        Icon
            las la-times

        """
        return self.find_element_by_xpath(element)

    @activity
    def by_class(self, element):
        """Find class in browser

        Find element with specified class on a webpage in the the browser. Can also use native 'find_element_by_class_name'

        :parameter element: Class of element
        :type element: string, optional

        :return: Element by class

            :Example:

        >>> # Open the browser
        >>> browser = Chrome()
        >>> # Go to a website
        >>> browser.get('https://wikipedia.org')
        >>> # Find element by class
        >>> element = browser.by_class('search-input')
        >>> # We can now use this element, for example to click on
        >>> element.click()

        Keywords
            browser, class, classes, element, xml element by text, chrome, internet, browsing, browser, surfing, web, webscraping, www, selenium, crawling, webtesting, mozilla, firefox, internet explorer

        Icon
            las la-times

        """
        return self.find_element_by_class_name(element)

    @activity
    def by_classes(self, element):
        """Find class in browser

        Find all elements with specified class on a webpage in the the browser. Can also use native 'find_elements_by_class_name' function

        :parameter element: Class of element
        :type element: string, optional

        :return: Element by classes

            :Example:

        >>> # Open the browser
        >>> browser = Chrome()
        >>> # Go to a website
        >>> browser.get('https://wikipedia.org')
        >>> # Find elements by class
        >>> elements = browser.by_classes('search-input')

        Keywords
            browser, class, classes, element, xml element by text, chrome, internet, browsing, browser, surfing, web, webscraping, www, selenium, crawling, webtesting, mozilla, firefox, internet explorer

        Icon
            las la-times

        """
        return self.find_elements_by_class_name(element)

    @activity
    def by_class_and_by_text(self, element, text):
        """Find element in browser based on class and text

        Find all elements with specified class and text on a webpage in the the browser. 

        :parameter element: Class of element
        :type element: string, optional

        :return: Element by class and text

            :Example:

        >>> # Open the browser
        >>> browser = Chrome()
        >>> # Go to a website
        >>> browser.get('https://wikipedia.org')
        >>> # Find elements by class and text
        >>> element = browser.by_class_and_by_text('search-input', 'Search Wikipedia')
        >>> # We can now use this element, for example to click on
        >>> element.click()

        Keywords
            browser, class, text, name classes, element, xml element by text, chrome, internet, browsing, browser, surfing, web, webscraping, www, selenium, crawling, webtesting, mozilla, firefox, internet explorer

        Icon
            las la-times

        """
        for element in self.find_elements_by_class_name(element):
            if text in element.text:
                return element

    @activity
    def by_id(self, element):
        """Find id in browser

        Find element with specified id on a webpage in the the browser. Can also use native 'find_element_by_id' function

        :parameter element: Id of element
        :type element: string, optional

        :return: Element by id

            :Example:

        >>> # Open the browser
        >>> browser = Chrome()
        >>> # Go to a website
        >>> browser.get('https://wikipedia.org')
        >>> # Find element by class
        >>> elements = browser.by_id('search-input')
        >>> # We can now use this element, for example to click on
        >>> element.click()

        Keywords
            browser, class, classes, element, xml element by text, chrome, internet, browsing, browser, surfing, web, webscraping, www, selenium, crawling, webtesting, mozilla, firefox, internet explorer

        Icon
            las la-times

        """
        return self.find_element_by_id(element)

    @activity
    def switch_to_iframe(self, name="iframe"):
        """Switch to iframe in browser

        Switch to an iframe in the browser

        :parameter name: Name of the Iframe
        :type name: string, optional

            :Example:

        >>> # Open the browser
        >>> browser = Chrome()
        >>> # Go to a website
        >>> browser.get('https://www.w3schools.com/html/html_iframe.asp')
        >>> # Switch to iframe
        >>> browser.switch_to_iframe()

        Keywords
            browser, class, classes, element, xml element by text, chrome, internet, browsing, browser, surfing, web, webscraping, www, selenium, crawling, webtesting, mozilla, firefox, internet explorer

        Icon
            las la-times

        """

        return self.switch_to.frame(self.find_element_by_tag_name("iframe"))


"""
Credential Management
Icon: las la-key
"""


@activity
def set_credential(username=None, password=None, system="Automagica"):
    """Set credential

    Add a credential which stores credentials locally and securely. All parameters should be Unicode text. 

    :parameter username: Username for which credential will be added.
    :type username: string, optional
    :parameter password: Password to add
    :type password: string, optional
    :parameter system: Name of the system for which credentials are stored. Extra safety measure and method for keeping passwords for similar usernames on different applications a part. Highly recommended to change default value.
    :type system: string

    :return: Stores credentials locally

        :Example:

    >>> set_credential('SampleUsername', 'SamplePassword')

    Keywords
        credential, login, password, username, store, vault, secure, credentials, store, log in, encrypt

    Icon
        las la-key

    """
    import keyring

    keyring.set_password(system, username, password)


@activity
def delete_credential(username=None, password=None, system="Automagica"):
    """Delete credential

    Delete a locally stored credential. All parameters should be Unicode text. 

    :parameter username: Username for which credential (username + password) will be deleted.
    :type username: string
    :parameter password: Password to delete
    :type password: string, optional
    :parameter system: Name of the system for which password will be deleted. 

        :Example:

    >>> set_credential('SampleUsername', 'SamplePassword')
    >>> delete_credential('SampleUsername', 'SamplePassword')

    Keywords
        credential, delete, login, password, username, store, vault, secure, credentials, store, log in, encrypt

    Icon
        las la-key

    """
    import keyring

    keyring.delete_password(system, username)


@activity
def get_credential(username=None, system="Automagica"):
    """Get credential

    Get a locally stored redential. All parameters should be Unicode text. 

    :parameter username: Username to get password for.
    :type username: string
    :parameter system: Name of the system for which credentials are retreived.
    :type system: string, optional

    :return: Stored credential as string

        :Example:

    >>> set_credential('SampleUsername', 'SamplePassword')
    >>> get_credential('SampleUsername')
    'SamplePassword'

    Keywords
        credential, get, delete, login, password, username, store, vault, secure, credentials, store, log in, encrypt

    Icon
        las la-key

    """
    import keyring

    return keyring.get_password(system, username)


"""
FTP
Icon: las la-key
"""


class FTP:
    @activity
    def __init__(self, server, username, password):
        """Create FTP connection (insecure)

        Can be used to automate activites for FTP

        :parameter server: Name of the server
        :type server: string
        :parameter username: Username 
        :type username: string
        :parameter password: Password
        :type password: string

            :Example:

        >>> # This example uses the Rebex FPT test server.
        >>> # Take caution uploading and downloading from this server as it is public
        >>> ftp = FTP('test.rebex.net', 'demo', 'password')

        Keywords
            FTP, file transfer protocol, filezilla, winscp, server, remote, folder, folders

        Icon
            las la-folder-open

        """
        import ftplib  # nosec

        self.connection = ftplib.FTP(server)  # nosec
        self.connection.login(username, password)

    @activity
    def download_file(self, input_path, output_path=None):
        """Download file
        
        Downloads a file from FTP server. Connection needs to be established first.

        :parameter input_path: Path to the file on the FPT server to download
        :type input_path: input_file
        :parameter output_path: Destination path for downloaded files. Default is the same directory with "_downloaded" added to the name
        :type output_path: output_dir, optional

        :return: Path to output file as string 

            :Example:

        >>> # This example uses the Rebex FPT test server.
        >>> # Take caution uploading and downloading from this server as it is public
        >>> ftp = FTP('test.rebex.net', 'demo', 'password')
        >>> # Download Rebex public file 'readme.txt'
        >>> ftp.download_file('readme.txt')
        'C:\\Users\\<username>\\readme_downloaded.txt'

        Keywords
            FTP, file transfer protocol, download, filezilla, winscp, server, remote, folder, folders

        Icon
            las la-download

        """
        # Set path if not specified
        input_path = interpret_path(input_path)
        if not output_path:
            output_path = interpret_path(default_filename="downloaded_readme.txt")
        else:
            output_path = interpret_path(output_path)

        self.connection.retrbinary("RETR " + input_path, open(output_path, "wb").write)

        return output_path

    @activity
    def upload_file(self, input_path, output_path=None):
        """Upload file
        
        Upload file to FTP server

        :parameter input_path: Path file that will be uploaded
        :type input_path: input_file
        :parameter output_path: Destination path to upload. 
        :type output_path: output_dir, optional

        :return: Path to uploaded file as string

            :Example:

        >>> # This example uses the Rebex FPT test server.
        >>> # Take caution uploading and downloading from this server as it is public
        >>> ftp = FTP('test.rebex.net', 'demo', 'password')
        >>> # Create a .txt file for illustration
        >>> text_file = make_text_file()
        >>> # Upload file to FTP test server
        >>> # Not that this might result in a persmission error for public FPT's
        >>> ftp.upload_file(input_path = text_file)

        Keywords
            FTP, upload, fptfile transfer protocol, filezilla, winscp, server, remote, folder, folders

        Icon
            las la-upload
        """
        # Set to user home if no path specified

        input_path = interpret_path(input_path)

        if not output_path:
            output_path = "/"

        self.connection.retrbinary("RETR " + input_path, open(output_path, "wb").write)

    @activity
    def enumerate_files(self, path="/"):
        """List FTP files

        Generate a list of all the files in the FTP directory

        :parameter path: Path to list files from. Default is the main directory
        :type path: input_dir, optional

        :return: Prints list of all files and directories

            :Example:

        >>> # This example uses the Rebex FPT test server.
        >>> # Take caution uploading and downloading from this server as it is public
        >>> ftp = FTP('test.rebex.net', 'demo', 'password')
        >>> # Show all files in main directory
        >>> ftp.enumerate_files()
        10-27-15  03:46PM       <DIR>          pub
        04-08-14  03:09PM                  403 readme.txt
        '226 Transfer complete.'

        Keywords
            FTP, list, upload, fptfile transfer protocol, filezilla, winscp, server, remote, folder, folders

        Icon
            las la-list-ol
        """
        self.connection.cwd(path)
        lines = self.connection.retrlines("LIST")
        return lines

    @activity
    def directory_exists(self, path="/"):
        """Check FTP directory
        
        Check if FTP directory exists

        :parameter path: Path to check on existence. Default is main directory
        :type path: input_dir, optional

        :return: Boolean

            :Example:

        >>> # This example uses the Rebex FPT test server.
        >>> # Take caution uploading and downloading from this server as it is public
        >>> ftp = FTP('test.rebex.net', 'demo', 'password')
        >>> # Check if 'pub' folder exists in main directory
        >>> ftp.directory_exists('\\pub')
        True

        Keywords
            FTP, list, upload, fptfile transfer protocol, filezilla, winscp, server, remote, folder, folders

        Icon
            las la-list-ol

        """
        try:
            self.connection.cwd(path)
            return True
        except:
            return False

    @activity
    def create_directory(self, directory_name, path="/"):
        """Create FTP directory

        Create a FTP directory. Note that sufficient permissions are present

        :parameter directory_name: Name of the new directory, should be a string e.g. 'my_directory'
        :type directory_name: string
        :parameter path: Path to parent directory where to make new directory. Default is main directory
        :type path: output_dir, optional

        :return: Boolean if creation was succesful (True) or failed (False)
            :Example:

        >>> # This example uses the Rebex FPT test server.
        >>> # Trying to create a directory will most likely fail due to permission
        >>> ftp = FTP('test.rebex.net', 'demo', 'password')
        >>> # Create directory
        >>> ftp.create_directory('brand_new_directory')      
        False

        Keywords
            FTP, create, create folder, new, new folder, fptfile transfer protocol, filezilla, winscp, server, remote, folder, folders

        Icon
            las la-folder-plus

        """
        try:
            self.connection.cwd(path)
            try:
                self.connection.mkd(directory_name)
                return True
            except Exception as e:
                if not e.args[0].startswith("550"):  # Exists already
                    raise
        except:
            return False


"""
Keyboard
Icon: las la-keyboard
"""


def easy_key_translation(key):
    """Activity supporting key translations
    """

    if not key:
        return ""

    key_translation = {
        "backspace": "{BACKSPACE}",
        "break": "{BREAK}",
        "capslock": "{CAPSLOCK}",
        "del": "{DELETE}",
        "alt": "%",
        "ctrl": "^",
        "shift": "+",
        "downarrow": "{DOWN}",
        "end": "{END}",
        "enter": "{ENTER}",
        "escape": "{ESC}",
        "help": "{HELP}",
        "home": "{HOME}",
        "insert": "{INSERT}",
        "win": "^{Esc}",
        "left": "{LEFT}",
        "numlock": "{NUMLOCK}",
        "pagedown": "{PGDN}",
        "pageup": "{PGUP}",
        "printscreen": "{PRTSC}",
        "right": "{RIGHT}",
        "scrolllock": "{SCROLLLOCK}",
        "tab": "{TAB}",
        "uparrow": "{UP}",
        "f1": "{F1}",
        "f2": "{F2}",
        "f3": "{F3}",
        "f4": "{F4}",
        "f5": "{F5}",
        "f6": "{F6}",
        "f7": "{F7}",
        "f8": "{F8}",
        "f9": "{F9}",
        "f10": "{F10}",
        "f11": "{F11}",
        "f12": "{F12}",
        "f13": "{F13}",
        "f14": "{F14}",
        "f15": "{F15}",
        "f16": "{F16}",
        "+": "{+}",
        "^": "{^}",
        "%": "{%}",
        "~": "{~}",
        "(": "{(}",
        ")": "{)}",
        "[": "{[}",
        "]": "{]}",
        "{": "{{}",
        "}": "{}}",
    }

    if key_translation.get(key):
        return key_translation.get(key)

    return key


@activity
def press_key(key=None, delay=1, perform_n_times=1, delay_between=0.5):
    """Press key

    Press and release an entered key. Make sure your keyboard is on US layout (standard QWERTY). 
    If you are using this on Mac Os you might need to grant acces to your terminal application. (Security Preferences > Security & Privacy > Privacy > Accessibility)

    :parameter key: Key to press. This can also be a scan code (e.g: 33 for '!')
    :options key: [' ', '!', '"', '#', '$', '%', '&', "'", '(', ')', '*', '+', ',', '-', '.', '/', '0', '1', '2', '3', '4', '5', '6', '7', '8', '9', ':', ';', '<','=', '>', '?', '@', '[', '\', ']', '^', '_', '`', 'a', 'b', 'c', 'd', 'e','f', 'g', 'h', 'i', 'j', 'k', 'l', 'm', 'n', 'o', 'p', 'q', 'r', 's', 't', 'u', 'v', 'w', 'x', 'y', 'z', '{', '|', '}', '~', 'alt', 'backspace', 'end', 'ctrl', 'del', 'down', 'right', 'left', 'up', 'enter', 'escape', 'f1', 'f2', 'f3', 'f4', 'f5', 'f6', 'f7', 'f8',  'f9', 'f10', 'f11', 'f12', 'f13', 'f14', 'f15', 'f16', 'home', 'insert', 'pagedown', 'pageup', 'help', 'space', 'tab', 'shift', 'win']
    :parameter delay: Delay before key is pressed in seconds, default is 1 second
    :type delay: int, optional
    :parameter perform_n_times: How many times to perform the key press
    :type perform_n_times: int, optional
    :parameter delay_between: Delay between key presses
    :type delay_between: float, optional


    :return: Keypress

        :Example:

    >>> # Open notepad to illustrate typing
    >>> run('notepad.exe')
    >>> # Press some keys
    >>> press_key('a')
    >>> press_key('enter')
    >>> press_key('b')
    >>> press_key('enter')
    >>> press_key('c')

    Keywords
        keyboard, typing, type, key, keystroke, hotkey, press, press key

    Icon
        las la-keyboard

    """
    if delay:
        from time import sleep

        sleep(delay)

    for i in range(perform_n_times):

        import platform

        # Check if system is not running Windows
        if platform.system() == "Windows":
            import win32com.client

            shell = win32com.client.dynamic.Dispatch("WScript.Shell")
            shell.SendKeys(easy_key_translation(key), 0)

        else:
            from keyboard import send

            if key:
                send(key)

        if perform_n_times > 1:
            sleep(delay_between)


@activity
def press_key_combination(
    first_key, second_key, third_key=None, compatibility=False, delay=1
):
    """Press key combination

    Press a combination of two or three keys simultaneously. Make sure your keyboard is on US layout (standard QWERTY).

    :parameter first_key: First key to press
    :options first_key: [' ', '!', '"', '#', '$', '%', '&', "'", '(', ')', '*', '+', ',', '-', '.', '/', '0', '1', '2', '3', '4', '5', '6', '7', '8', '9', ':', ';', '<','=', '>', '?', '@', '[', '\', ']', '^', '_', '`', 'a', 'b', 'c', 'd', 'e','f', 'g', 'h', 'i', 'j', 'k', 'l', 'm', 'n', 'o', 'p', 'q', 'r', 's', 't', 'u', 'v', 'w', 'x', 'y', 'z', '{', '|', '}', '~', 'alt', 'backspace', 'end', 'ctrl', 'del', 'down', 'right', 'left', 'up', 'enter', 'escape', 'f1', 'f2', 'f3', 'f4', 'f5', 'f6', 'f7', 'f8',  'f9', 'f10', 'f11', 'f12', 'f13', 'f14', 'f15', 'f16', 'home', 'insert', 'pagedown', 'pageup', 'help', 'space', 'tab', 'shift', 'win']
    :parameter second_key: Second key to press
    :options second_key: [' ', '!', '"', '#', '$', '%', '&', "'", '(', ')', '*', '+', ',', '-', '.', '/', '0', '1', '2', '3', '4', '5', '6', '7', '8', '9', ':', ';', '<','=', '>', '?', '@', '[', '\', ']', '^', '_', '`', 'a', 'b', 'c', 'd', 'e','f', 'g', 'h', 'i', 'j', 'k', 'l', 'm', 'n', 'o', 'p', 'q', 'r', 's', 't', 'u', 'v', 'w', 'x', 'y', 'z', '{', '|', '}', '~', 'alt', 'backspace', 'end', 'ctrl', 'del', 'down', 'right', 'left', 'up', 'enter', 'escape', 'f1', 'f2', 'f3', 'f4', 'f5', 'f6', 'f7', 'f8',  'f9', 'f10', 'f11', 'f12', 'f13', 'f14', 'f15', 'f16', 'home', 'insert', 'pagedown', 'pageup', 'help', 'space', 'tab', 'shift', 'win']
    :parameter third_key: Third key to press, this is optional.
    :options third_key: [' ', '!', '"', '#', '$', '%', '&', "'", '(', ')', '*', '+', ',', '-', '.', '/', '0', '1', '2', '3', '4', '5', '6', '7', '8', '9', ':', ';', '<','=', '>', '?', '@', '[', '\', ']', '^', '_', '`', 'a', 'b', 'c', 'd', 'e','f', 'g', 'h', 'i', 'j', 'k', 'l', 'm', 'n', 'o', 'p', 'q', 'r', 's', 't', 'u', 'v', 'w', 'x', 'y', 'z', '{', '|', '}', '~', 'alt', 'backspace', 'end', 'ctrl', 'del', 'down', 'right', 'left', 'up', 'enter', 'escape', 'f1', 'f2', 'f3', 'f4', 'f5', 'f6', 'f7', 'f8',  'f9', 'f10', 'f11', 'f12', 'f13', 'f14', 'f15', 'f16', 'home', 'insert', 'pagedown', 'pageup', 'help', 'space', 'tab', 'shift', 'win']
    :parameter compatibility: Set parameter to true to not use win32com. This could help with compatibility on certain systems or when certain keypresses do not work correctly.
    :parameter key: Delay before keys are pressed in seconds, default is 1 second
    :type key: int, optional

    :return: Key combination

        :Example:

    >>> # Open notepad to illustrate typing
    >>> run('notepad.exe')
    >>> # Press 'ctrl + s' to prompt save window 
    >>> press_key_combination('ctrl', 's')

    Keywords
        keyboard, key combination, shortcut, typing, type, key, keystroke, hotkey, press, press key

    Icon
        las la-keyboard

    """

    if delay:
        from time import sleep

        sleep(delay)

    import platform

    # Check if system is not running Windows
    if first_key == "win" or second_key == "win" or third_key == "win":
        compatibility = True
    if platform.system() != "Windows" or compatibility:
        from keyboard import send

        if not third_key:
            send(first_key + "+" + second_key)
            return
        if third_key:
            send(first_key + "+" + second_key + "+" + third_key)
            return

    import win32com.client

    shell = win32com.client.dynamic.Dispatch("WScript.Shell")
    key_combination = (
        easy_key_translation(first_key)
        + easy_key_translation(second_key)
        + easy_key_translation(third_key)
    )
    shell.SendKeys(easy_key_translation(key_combination), 0)


@activity
def typing(text, automagica_id=None, clear=False, interval_seconds=0.01, delay=1):
    """Type text

    Simulate keystrokes. If an element ID is specified, text will be typed in a specific field or element based on the element ID (vision) by the recorder.

    Supported keys: 
        ' ', '!', '"', '#', '$', '%', '&', "'", '(', ,')', '*', '+', ',', '-', '.', '/', '0', '1', '2', '3', '4', '5', '6', '7', '8', '9', ':', ';', '<','=', '>', '?', '@', '[', '\', ']', '^', '_', '`', 'a', 'b', 'c', 'd', 'e','f', 'g', 'h', 'i', 'j', 'k', 'l', 'm', 'n', 'o', 'p', 'q', 'r', 's', 't', 'u', 'v', 'w', 'x', 'y', 'z', '{', '|', '}', '~', 'alt', 'backspace',  'ctrl', 'delete' 'downarrow', 'rightarrow', 'leftarrow', 'uparrow', 'enter', 'escape', 'f1', 'f2', f3', 'f4', 'f5', 'f6', 'f7', 'f8',  'f9', 'f10', 'f11', 'f12', 'f13', 'f14', 'f15', 'f16', 'home', 'insert', 'pagedown', 'pageup', 'help', 'printscreen', 'space', 'scrollock', 'tab', shift, 'win'

    :parameter text: Text in string format to type. Note that you can only press single character keys. Special keys can not be part of the text argument.
    :type text: string
    :parameter automagica_id: ID of the element. To define an element and attach an ID one can use the Automagica Wand. The recorder uses vision to detect an element and can be invoked with the recorder() function.
    :type automagica_id: automagica_id, optional
    :parameter clear: Attempts to clear the element before typing using hotkeys. Be cautious when using this method as a vision mismatch could result in deleting unwanted data. Default value is False
    :type clear: bool, optional
    :parameter interval_seconds: Time in seconds between two keystrokes. Defautl value is 0.01 seconds.
    :type interval_seconds: int, optional
    :parameter delay: Delay before beginning to type, default is 1 second
    :type delay: int, optional

    :return: Keystrokes

        :Example:

    >>> # Open notepad to illustrate typing
    >>> run('notepad.exe')
    >>> # Type a story
    >>> typing('Why was the robot mad? \n They kept pushing his buttons!')

    Keywords
        keyboard, keystrokes, key combination, shortcut, typing, type, key, keystroke, hotkey, press, press key, send keys, keystrokes

    Icon
        las la-keyboard

    """

    if delay:
        from time import sleep

        sleep(delay)

    if automagica_id:
        location = detect_vision(automagica_id)
        x, y = get_center_of_rectangle(location)

        from mouse import click, move

        move(x, y)
        click()

    if clear:
        press_key_combination("ctrl", "a")
        press_key("del")

    import platform

    # Set keyboard layout for Windows platform
    if platform.system() != "Windows":
        from keyboard import write

        return write(text, delay=interval_seconds)

    import win32com.client

    shell = win32com.client.dynamic.Dispatch("WScript.Shell")
    import time

    for character in text:
        shell.SendKeys(easy_key_translation(character), 0)
        time.sleep(interval_seconds)


"""
Mouse
Icon: las la-mouse-pointer
"""


@activity
def get_mouse_position(delay=None, to_clipboard=False):
    """Get mouse coordinates

    Get the x and y pixel coordinates of current mouse position.
    These coordinates represent the absolute pixel position of the mouse on the computer screen. The x-coördinate starts on the left side and increases going right. The y-coördinate increases going down.

    :parameter delay: Delay in seconds before capturing mouse position.
    :type delay: int, optional
    :parameter to_clipboard: Put the coordinates in the clipboard e.g. 'x=1, y=1'
    :type to_clipboard: bool, optional

    :return: Tuple with (x, y) coordinates

        :Example:

    >>> get_mouse_position()
    (314, 271)

    Keywords
        mouse, mouse automation, click, right click, mouse button, move mouse, position, pixel

    Icon
        las la-mouse
    """
    from mouse import get_position
    from time import sleep

    if delay:
        sleep(delay)

    coord = get_position()

    if to_clipboard:
        set_to_clipboard("x=" + str(coord[0]) + ", y=" + str(coord[1]))

    return coord[0], coord[1]


@activity
def display_mouse_position(duration=10):
    """Display mouse position

    Displays mouse position in an overlay. Refreshes every two seconds. Can be used to find mouse position of element on the screen. 
    These coordinates represent the absolute pixel position of the mouse on the computer screen. The x-coördinate starts on the left side and increases going right. The y-coördinate increases going down.

    :parameter duration: Duration to show overlay.
    :type duration: int, optional

    :return: Overlay with (x, y) coordinates

        :Example:

    >>> display_mouse_position()

    Keywords
        mouse, osd, overlay, show, display, mouse automation, click, right click, mouse button, move mouse, position, pixel

    Icon
        lars la-search-location
    """

    if duration < 1 or type(duration) != int:
        return

    from mouse import get_position
    from time import sleep

    duration_half = int(duration / 2)
    for i in range(0, duration_half, 2):
        coord = get_position()
        message = "x=" + str(coord[0]) + ", y=" + str(coord[1])
        display_osd_message(message, seconds=2)
        sleep(2)


@activity
def click(automagica_id, delay=1):
    """Mouse click

    Clicks on an element based on the element ID (vision)

    :parameter automagica_id: ID of the element. To define an element and attach an ID one can use the Automagica Wand. The recorder uses vision to detect an element and can be invoked with the recorder() function.
    :type automagica_id: automagica_id
    :parameter delay: Delay before clicking in seconds. 
    :type delay: int, optional

        :Example:

    >>> # Click on a vision element, use the recorder() function to define elements
    >>> # Run the Windows calculator and try to perform the activity
    >>> run('calc.exe')
    >>> # Use the element ID found by the recorder, in this case ID 'qf41'. You can also view this on automagica.id/qf41 
    >>> #  If you have a vastly different version or layout the element might not be found, use the recorder 
    >>> click('qf41')

    Keywords
        mouse, vision, mouse, osd, overlay, show, display, mouse automation, click, right click, mouse button, move mouse, position, pixel

    Icon
        las la-mouse-pointer
    """
    from mouse import click as click_
    from mouse import move as move_

    if delay:
        from time import sleep

        sleep(delay)  # Default delay

    location = detect_vision(automagica_id)
    x, y = get_center_of_rectangle(location)

    move_(x, y)
    click_()


@activity
def click_coordinates(x=None, y=None, delay=1):
    """Mouse click coordinates

    Clicks on an element based on pixel position determined by x and y coordinates. To find coordinates one could use display_mouse_position().

    :parameter x: X-coordinate
    :type x: int
    :parameter y: Y-coordinate
    :type y: int
    :parameter delay: Delay before clicking in seconds. 
    :type delay: int, optional

        :Example:

    >>> # Click on pixel position
    >>> click_coordinates(x=100, y=100)

    Keywords
        mouse, vision, mouse, osd, overlay, show, display, mouse automation, click, right click, mouse button, move mouse, position, pixel

    Icon
        las la-mouse-pointer
    """

    from mouse import click as click_
    from mouse import move as move_

    if delay:
        from time import sleep

        sleep(delay)  # Default delay

    if x and y:
        move_(x, y)
        click_()

    else:
        raise Exception("Could not click, did you enter a valid ID or coordinates")


@activity
def double_click_coordinates(x=None, y=None, delay=1):
    """Double mouse click coordinates

    Double clicks on a pixel position determined by x and y coordinates.

    :parameter x: X-coordinate
    :type x: int
    :parameter y: Y-coordinate
    :type y: int
    :parameter delay: Delay before cliking in seconds. 
    :type delay: int, optional

        :Example:

    >>> # Click on coordinates
    >>> double_click_coordinates(x=100, y=100)

    Keywords
        mouse, osd, overlay, double, double click, doubleclick show, display, mouse automation, click, right click, mouse button, move mouse, position, pixel

    Icon
        las la-mouse-pointer
    """
    from mouse import double_click as double_click_
    from mouse import move as move_

    if delay:
        from time import sleep

        sleep(delay)  # Default delay

    if x and y:
        move_(x, y)
        double_click_()

    else:
        raise Exception("Could not click, did you enter valid coordinates")


@activity
def double_click(automagica_id=None, delay=1):
    """Double mouse click

    Double clicks on an element based on the element ID (vision) 

    :parameter automagica_id: ID of the element. To define an element and attach an ID one can use the Automagica Wand. The recorder uses vision to detect an element and can be invoked with the recorder() function.
    :type automagica_id: automagica_id
    :parameter delay: Delay before clicking in seconds. 
    :type delay: int, optional

        :Example:

    >>> # Click on a vision element, use the recorder() function to define elements
    >>> # Run the Windows calculator and try to perform the activity
    >>> run('calc.exe')
    >>> # Use the element ID found by the recorder, in this case ID 'qf41'. You can also view this on automagica.id/qf41 
    >>> #  If you have a vastly different version or layout the element might not be found, use the recorder 
    >>> double_click('qf41')

    Keywords
        mouse, osd, overlay, double, double click, doubleclick show, display, mouse automation, click, right click, mouse button, move mouse, position, pixel

    Icon
        las la-mouse-pointer
    """
    from mouse import double_click as double_click_
    from mouse import move as move_

    if delay:
        from time import sleep

        sleep(delay)  # Default delay

    if automagica_id:
        location = detect_vision(automagica_id)
        x, y = get_center_of_rectangle(location)

        move_(x, y)
        double_click_()

    else:
        raise Exception("Could not click, did you enter a valid ID")


@activity
def right_click(automagica_id=None, delay=1):
    """Right click

    Right clicks on an element based on the element ID (vision)

    :parameter automagica_id: ID of the element. To define an element and attach an ID one can use the Automagica Wand. The recorder uses vision to detect an element and can be invoked with the recorder() function.
    :type automagica_id: automagica_id    
    :parameter delay: Delay before cliking in seconds. 
    :type delay: int, optional

        :Example:

    >>> # Click on a vision element, use the recorder() function to define elements
    >>> # Run the Windows calculator and try to perform the activity
    >>> run('calc.exe')
    >>> # Use the element ID found by the recorder, in this case ID 'qf41'. You can also view this on automagica.id/qf41 
    >>> #  If you have a vastly different version or layout the element might not be found, use the recorder 
    >>> right_click('qf41')

    Keywords
        mouse, osd, right click, right, rightclick, overlay, show, display, mouse automation, click, right click, mouse button, move mouse, position, pixel

    Icon
        las la-mouse-pointer
    """
    from mouse import right_click as right_click_
    from mouse import move as move_

    if delay:
        from time import sleep

        sleep(delay)  # Default delay

    if automagica_id:
        location = detect_vision(automagica_id)
        x, y = get_center_of_rectangle(location)

        move_(x, y)
        right_click_()

    else:
        raise Exception("Could not click, did you enter a valid ID or coordinates")


@activity
def right_click_coordinates(x=None, y=None, delay=1):
    """Right click coordinates

    Right clicks on an element based pixel position determined by x and y coordinates.

    :parameter x: X-coördinate
    :type x: int
    :parameter y: Y-coördinate
    :type y: int
    :parameter delay: Delay before clicking in seconds 
    :type delay: int, optional

        :Example:

    >>> # Right click on coordinates
    >>> right_click_coordinates(x=100, y=100)

    Keywords
        mouse, osd, right click, right, rightclick, overlay, show, display, mouse automation, click, right click, mouse button, move mouse, position, pixel

    Icon
        las la-mouse-pointer
    """
    from mouse import right_click as right_click_
    from mouse import move as move_

    if delay:
        from time import sleep

        sleep(delay)  # Default delay

    if x and y:
        move_(x, y)
        right_click_()

    else:
        raise Exception("Could not click, did you enter a valid ID or coordinates")


@activity
def move_mouse_to(automagica_id=None, delay=1):
    """Move mouse

    Moves te pointer to an element based on the element ID (vision)

    :parameter automagica_id: ID of the element. To define an element and attach an ID one can use the Automagica Wand. The recorder uses vision to detect an element and can be invoked with the recorder() function.
    :type automagica_id: automagica_id, optional
    :parameter delay: Delay before movement in seconds 
    :type delay: int, optional


    :return: Move mouse to (x, y) coordinates

        :Example:

    >>> # Use recorder to find an element ID
    >>> # Run the Windows calculator and try to perform the activity
    >>> run('calc.exe')
    >>> # Use the element ID found by the recorder, in this case ID 'qf41'. You can also view this on automagica.id/qf41 
    >>> #  If you have a vastly different version or layout the element might not be found, use the recorder 
    >>> move_mouse_to('qf41')

    Keywords
        mouse, osd, move mouse, right click, right, rightclick, overlay, show, display, mouse automation, click, right click, mouse button, move mouse, position, pixel

    Icon
        las la-arrows-alt
    """
    from mouse import move as move_

    if delay:
        from time import sleep

        sleep(delay)  # Default delay

    if automagica_id:
        location = detect_vision(automagica_id)
        x, y = get_center_of_rectangle(location)
        move_(x, y)

        return
    if x and y:
        move_(x, y)


@activity
def move_mouse_to_coordinates(x=None, y=None, delay=1):
    """Move mouse coordinates

    Moves te pointer to an element based on the pixel position determined by x and y coordinates

    :parameter x: X-coördinate
    :type x: int
    :parameter y: Y-coördinate
    :type y: int
    :parameter delay: Delay between movements in seconds, standard value is 1s. 
    :type delay: int, optional

    :return: Move mouse to (x, y) coordinates

        :Example:

    >>> # Move mouse to coordinates
    >>> move_mouse_to_coordinates(x=100, y=100)
 

    Keywords
        mouse, osd, move mouse, right click, right, rightclick, overlay, show, display, mouse automation, click, right click, mouse button, move mouse, position, pixel

    Icon
        las la-arrows-alt
    """
    from mouse import move as move_

    if delay:
        from time import sleep

        sleep(delay)  # Default delay

    if x and y:
        move_(x, y)


@activity
def move_mouse_relative(x=None, y=None):
    """Move mouse relative

    Moves the mouse an x- and y- distance relative to its current pixel position.

    :parameter x: X-coördinate
    :type x: int
    :parameter y: Y-coördinate
    :type y: int

    :return: Move mouse (x, y) coordinates

        :Example:

    >>> move_mouse_to_coordinates(x=100, y=100)
    >>> wait(1)
    >>> move_mouse_relative(x=10, y=10)

    Keywords
        mouse, osd, move mouse, right click, right, rightclick, overlay, show, display, mouse automation, click, right click, mouse button, move mouse, position, pixel

    Icon
        las la-arrows-alt
    """

    from mouse import move as move_

    move_(x, y, absolute=False)


@activity
def drag_mouse_to_coordinates(x=None, y=None, delay=1):
    """Drag mouse

    Drags mouse to an element based on pixel position determined by x and y coordinates

    :parameter x: X-coördinate
    :type x: int
    :parameter y: Y-coördinate
    :type y: int
    :parameter delay: Delay between movements in seconds, standard value is 1s. 
    :type delay: int, optional

    :return: Drag mouse 

        :Example:

    >>> # Use coordinates to move and drag mouse
    >>> move_mouse_to_coordinates(x=100, y=100)
    >>> drag_mouse_to_coordinates(x=1, y=1)

    Keywords
        mouse, osd, move mouse, right click, right, rightclick, overlay, show, display, mouse automation, click, right click, mouse button, move mouse, position, pixel

    Icon
        las la-arrows-alt
    """
    if delay:
        from time import sleep

        sleep(delay)  # Default delay

    if x and y:
        from mouse import drag

        drag(x, y, absolute=False)


@activity
def drag_mouse_to(automagica_id=None, delay=1):
    """Drag mouse

    Drags mouse to an element based on the element ID (vision) 

    :parameter automagica_id: ID of the element. To define an element and attach an ID one can use the Automagica Wand. The recorder uses vision to detect an element and can be invoked with the recorder() function.
    :type automagica_id: automagica_id
    :parameter delay: Delay before movement in seconds.
    :type delay: int, optional

    :return: Drag mouse 

        :Example:

    >>> # Use recorder to find an element ID
    >>> # Run the Windows calculator and try to perform the activity
    >>> run('calc.exe')
    >>> # Use the element ID found by the recorder, in this case ID 'qf41'. You can also view this on automagica.id/qf41 
    >>> #  If you have a vastly different version or layout the element might not be found, use the recorder 
    >>> drag_mouse_to('qf41')

    Keywords
        mouse, osd, move mouse, right click, right, rightclick, overlay, show, display, mouse automation, click, right click, mouse button, move mouse, position, pixel

    Icon
        las la-arrows-alt
    """
    if delay:
        from time import sleep

        sleep(delay)  # Default delay

    if automagica_id:
        location = detect_vision(automagica_id)
        x, y = get_center_of_rectangle(location)

        from mouse import drag

        drag(x, y, absolute=False)

        return


"""
Image
Icon: las la-image
"""


@activity
def random_screen_snippet(size=100, output_path=None):
    """Random screen snippet

    Take a random square snippet from the current screen. Mainly for testing and/or development purposes.

    :parameter size: Size (width and height) in pixels for square snippet. Default value is 100 pixels
    :type size: int, optional
    :parameter output_path: Path where snippet will be saved. Default value is home directory with name 'random_screensnippet.jpg'
    :type output_path: output_file
    :extension output_path: jpg

    :return: Path to snippet

        :Example:

    >>> random_screen_snippet()
    'C:\\Users\\<username>\\random_screensnippet.jpg'

    Keywords
        image, random, testing, screengrab, snippet

    Icon
        las la-crop-alt

    """
    only_supported_for("Windows", "Darwin")

    import PIL.ImageGrab

    img = PIL.ImageGrab.grab()

    width, height = img.size

    import random

    random_left = random.randrange(1, width, 1)
    random_top = random.randrange(1, height, 1)

    left, top, right, bottom = (
        random_left,
        random_top,
        random_left + size,
        random_top + size,
    )
    cropped = img.crop((left, top, right, bottom))

    if not output_path:
        output_path = interpret_path(
            output_path, default_filename="random_screensnippet.jpg"
        )
    else:
        output_path = interpret_path(output_path)

    cropped.save(output_path, "JPEG")

    return output_path


@activity
def take_screenshot(output_path=None):
    """Screenshot

    Take a screenshot of current screen.

    :parameter output_path: Path to save screenshot. Default value is home directory with name 'screenshot.jpg'.
    :type output_path: output_path
    :extesion output_path: jpg

    :return: Path to save screenshot

        :Example:

    >>> new_screenshot = take_screenshot()
    >>> open_file(new_screenshot)
    'C:\\Users\\<username>\\screenshot.jpg'

    Keywords
        image, screenshot, printscreen,

    Icon
        las la-expand

    """
    only_supported_for("Windows", "Darwin")

    import PIL.ImageGrab

    img = PIL.ImageGrab.grab()

    output_path = interpret_path(path=output_path, default_filename="screenshot.jpg")

    img.save(output_path, "JPEG")

    return output_path


"""
Folder Operations
Icon: las la-folder-open
"""


@activity
def get_files_in_folder(
    input_path=None, extension=None, show_full_path=True, scan_subfolders=False
):
    """List files in folder

    List all files in a folder (and subfolders)
    Checks all folders and subfolders for files. This could take some time for large repositories. 

    :parameter input_path: Path of the folder to retreive files from. Default folder is the home directory.
    :type input_path: input_dir
    :parameter extension: Optional filter on certain extensions, for example 'pptx', 'exe,' xlsx', 'txt', .. Default value is no filter.
    :type extension: string, optional
    :parameter show_full_path: Set this to True to show full path, False will only show file or dirname. Default is True
    :type show_full_path: bool, optional
    :scan_subfolders: Boolean to scan subfolders or not. Not that depending on the folder and hardware this activity could take some time if scan_subfolders is set to True
    :type scan_subfolders: bool, optional

    :return: List of files with their full path

        :Example:

    >>> # List all files in the homedirectory
    >>> get_files_in_folder()
    ['C:\\Users\\<username>\\file1.jpg', 'C:\\Users\\<username>\\file2.txt', ... ]

    Keywords
        folder, files, explorer, nautilus, folder, file, create folder, get files, list files, all files, overview, get files

    Icon 
        las la-search
    """
    import os

    path = interpret_path(input_path)

    if scan_subfolders:
        paths = []
        for dirpath, _, filenames in os.walk(path):
            for f in filenames:
                full_path = os.path.abspath(os.path.join(dirpath, f))
                if extension:
                    if not full_path.endswith(extension):
                        continue
                if show_full_path:
                    paths.append(full_path)
                else:
                    paths.append(f)
        return paths

    if not scan_subfolders:
        paths = []
        for item in os.listdir(path):
            full_path = os.path.abspath(os.path.join(path, item))
            if extension:
                if not full_path.endswith(extension):
                    continue
            if show_full_path:
                paths.append(full_path)
            else:
                paths.append(item)
        return paths


@activity
def create_folder(path=None):
    """Create folder

    Creates new folder at the given path.

    :parameter path: Full path of folder that will be created. If no path is specified a folder called 'new_folder' will be made in home directory. If this folder already exists 8 random characters will be added to the name.
    :type path: input_dir, optional

    :return: Path to new folder as string

        :Example:

    >>> # Create folder in the home directory
    >>> create_folder()
    'C:\\Users\\<username>\\new_folder'

    Keywords
        create folder, folder, new folder,  folders, make folder, new folder, folder manipulation, explorer, nautilus

    Icon
        las la-folder-plus
    """
    import os

    if not path:
        path = interpret_path(path, default_filename="new_folder", random_addition=True)

    else:
        path = interpret_path(path)

    if not os.path.exists(path):
        os.makedirs(path)
        return path


@activity
def rename_folder(input_path, output_name=None):
    """Rename folder

    Rename a folder

    :parameter input_path: Full path of folder that will be renamed
    :type input_path: input_dir, optional
    :parameter output_name: New name. By default folder will be renamed to original folder name with '_renamed' added to the folder name.
    :type output_name: string, optional

    :return: Path to renamed folder as a string.

        :Example:

    >>> # Make new folder in home directory for illustration
    >>> testfolder = create_folder()
    >>> # Rename the folder
    >>> rename_folder(testfolder, output_name='testfolder_brand_new_name')
    'C:\\Users\\<username>\\testfolder_brand_new_name'

    Keywords
        folder, rename, rename folder, organise folder, folders, folder manipulation, explorer, nautilus

    Icon
        las la-folder

    """
    import os

    input_path = interpret_path(input_path, required=True)
    if not output_name:
        output_path = interpret_path(input_path, addition="_renamed")
    else:
        output_path = interpret_path(input_path, replace_filename=output_name)

    os.rename(input_path, output_path)

    return output_path


@activity
def show_folder(input_path=None):
    """Open a folder

    Open a folder with the default explorer.

    :parameter input_path: Full path of folder that will be opened. Default value is the home directory
    :type input_path: input_dir, optional

    :return: Path to opend folder as a string

        :Example:

    >>> # Make new folder in home directory for illustration
    >>> testfolder = create_folder()
    >>> # Open folder
    >>> show_folder(testfolder)
    'C:\\Users\\<username>\\new_folder'


    Keywords
        folder, open, open folder, explorer, nautilus

    Icon
        las la-folder-open

    """

    path = interpret_path(input_path)

    import os

    if os.path.isdir(path):
        os.startfile(path)

    return path


@activity
def move_folder(input_path, output_path=None):
    """Move a folder

    Moves a folder from one place to another.
    
    :parameter input_path: Full path to the source location of the folder
    :type input_path: input_dir
    :parameter output_path: Full path to the destination location of the folder, defaults to input_path with '_moved' added
    :type output_path: output_dir, optional

     :return: Path to new location of folder as a string. None if folder could not be moved.

        :Example:

    >>> # Make new folder in home directory for illustration
    >>> # If no new_folder exists in home dir this will be called new_folder
    >>> testfolder = create_folder()
    >>> # Make a second new folder
    >>> # Since new_folder already exists this folder will get a random id added (in this case abc1)
    >>> testfolder_2 = create_folder()
    >>> # Move testfolder in testfolder_2
    >>> move_folder(testfolder, testfolder_2)
    'C:\\Users\\<username>\\new_folder_abc1\\new_folder'

    Keywords
        folder, move, move folder, explorer, nautilus, folder manipulation

    Icon
        las la-folder
    """
    import shutil

    input_path = interpret_path(input_path, required=True)
    if not output_path:
        output_path = interpret_path(input_path, addition="_moved")
    else:
        output_path = interpret_path(output_path)

    shutil.move(input_path, output_path)

    return output_path


@activity
def remove_folder(input_path, allow_root=False, delete_read_only=True):
    """Remove folder

    Remove a folder including all subfolders and files. For the function to work optimal, all files and subfolders in the main targetfolder should be closed.

    :parameter input_path: Full path to the folder that will be deleted
    :type input_path: input_dir
    :parameter allow_root: Allow paths with an arbitrary length of 10 characters or shorter to be deleted. Default value is False.
    :type allow_root: bool, optional
    :parameter delete_read_only: Option to delete read only
    :type delete_read_only: bool, optional

    :return: Path to deleted folder as a string

        :Example:

    >>> # Make new folder in home directory for illustration
    >>> testfolder = create_folder()
    >>> # Check if folder exists
    >>> print( folder_exists(testfolder) ) # Should print True
    >>> # Remove folder
    >>> remove_folder(testfolder)
    >>> # Check again if folder exists
    >>> folder_exists(testfolder)
    False

    Keywords
        folder, delete folder, delete, nautilus, folder manipulation, explorer, delete folder, remove, remove folder

    Icon
        las la-folder-minus

    """
    import os
    import shutil

    path = interpret_path(input_path, required=True)

    if len(path) > 10 or allow_root:
        if os.path.isdir(path):
            shutil.rmtree(path, ignore_errors=delete_read_only)
            return path


@activity
def empty_folder(input_path, allow_root=False):
    """Empty folder
    
    Remove all contents from a folder
    For the function to work optimal, all files and subfolders in the main targetfolder should be closed.

    :parameter input_path: Full path to the folder that will be emptied
    :type input_path: input_dir
    :parameter allow_root: Allow paths with an arbitrary length of 10 characters or shorter to be emptied. Default value is False.

        :Example:

    >>> # Make new folder in home directory for illustration
    >>> testfolder = create_folder()
    >>> # Make new text file in this folder
    >>> text_file_location = make_text_file(output_path = testfolder)
    >>> # Print all files in the testfolder
    >>> get_files_in_folder(testfolder) 
    >>> # Empty the folder
    >>> empty_folder(testfolder)
    >>> # Check what is in the folder
    >>> get_files_in_folder(testfolder)
    []

    Keywords
        folder, empty folder, delete, empty, clean, clean folder, nautilus, folder manipulation, explorer, delete folder, remove, remove folder

    Icon
        las la-folder-minus
    """
    import os

    path = interpret_path(input_path, required=True)
    if len(path) > 10 or allow_root:
        if os.path.isdir(path):
            for root, dirs, files in os.walk(path, topdown=False):
                for name in files:
                    os.remove(os.path.join(root, name))
                for name in dirs:
                    os.rmdir(os.path.join(root, name))


@activity
def folder_exists(path):
    """Checks if folder exists

    Check whether folder exists or not, regardless if folder is empty or not.

    :parameter input_path: Full path to folder
    :type input_path: input_dir

    :return: Boolean

        :Example:

    >>> # Make new folder in home directory for illustration
    >>> testfolder = create_folder()
    >>> # Check if folder exists
    >>> folder_exists(testfolder)
    True

    Keywords
        folder, folder exists, nautilus, explorer, folder manipulation, files

    Icon
        las la-folder
    """
    import pathlib

    return pathlib.Path(path).is_dir()


@activity
def copy_folder(input_path, output_path=None):
    """Copy a folder

    Copies a folder from one place to another.
    

    :parameter input_path: Full path to the source location of the folder
    :type input_path: input_dir
    :parameter output_path: Full path to the destination location of the folder. If no path is specified folder will get copied in the input directory with '_copied' added
    :type output_path: output_dir, optional

    :return: Path to new folder as string

        :Example:

    >>> # Make new folder in home directory for illustration
    >>> testfolder = create_folder()
    >>> # Copy this folder
    >>> # Since new_folder already exists in home dir this folder will get a random id added (in this case abc1)
    >>> copy_folder(testfolder)

    Keywords
        folder, move, move folder, explorer, nautilus, folder manipulation

    Icon
        lar la-folder
    """
    import shutil

    input_path = interpret_path(input_path, required=True)
    if not output_path:
        output_path = interpret_path(input_path, addition="_copied")
    else:
        output_path = interpret_path(output_path)

    shutil.copytree(input_path, output_path)

    return output_path


@activity
def zip_folder(input_path, output_path=None):
    """Zip

     Zip folder and its contents. Creates a .zip file. 

    :parameter input_path: Full path to the source location of the folder that will be zipped
    :type input_path: input_dir
    :parameter output_path: Full path to save the zipped folder. If no path is specified a folder with the original folder with '_zipped' added
    :type output_path: output_dir, optional

    :return: Path to zipped folder

        :Example:

    >>> # Make new folder in home directory for illustration
    >>> testfolder = create_folder()
    >>> # Zip this folder
    >>> zip_folder(testfolder)

    Keywords
        zip, zipping, winrar, rar, 7zip, compress, unzip

    Icon
        las la-archive
    """
    import zipfile
    import os
    import shutil

    path = interpret_path(input_path, required=True)

    if not output_path:
        output_path = interpret_path(path, addition="_zipped")

    if os.path.isdir(path):
        shutil.make_archive(output_path, "zip", path)

    return output_path


@activity
def unzip(input_path, output_path=None):
    """Unzip

    Unzips a file or folder from a .zip file.

    :parameter input_path: Full path to the source location of the file or folder that will be unzipped
    :type input_path: input_dir
    :parameter to_path: Full path to save unzipped contents. If no path is specified the unzipped contents will be stored in the same directory as the zipped file is located. 
    :type output_path: output_dir, optional

    :return: Path to unzipped folder

        :Example:

    >>> # Make new file in home directory for illustration
    >>> testfolder = create_folder()
    >>> # Add some files to this folder
    >>> make_text_file(output_path = testfolder)
    >>> # Zip this folder
    >>> zipped_folder = zip_folder(testfolder)
    >>> # Unzip this folder
    >>> unzip(zipped_folder)

    Keywords
        zip, zipping, winrar, rar, 7zip, compress, unzip

    Icon
        las la-archive
    """
    import zipfile
    import os
    import shutil

    path = interpret_path(input_path)

    if not output_path:
        output_path = interpret_path(path, addition="_unzipped")
    else:
        output_path = interpret_path(output_path)

    import zipfile

    with zipfile.ZipFile(path) as zip_ref:
        zip_ref.extractall(output_path)

    return output_path


@activity
def most_recent_file(input_path=None):
    """Return most recent file in directory

    Return most recent file in directory

    :parameter input_path: Path which will be scanned for most recent file, defaults to homedir
    :type input_path: input_dir, optional

    :return: Path to most recent file

        :Example:

    >>> # Find most recent file in homedir
    >>> most_recent_file()

    Keywords
        find file, file, recent, newest, latest, recent

    Icon
        las la-clock
    """

    import os

    path = interpret_path(input_path)

    files = os.listdir(path)
    paths = [os.path.join(path, basename) for basename in files]
    return max(paths, key=os.path.getatime)


"""
Delay
Icon: las la-hourglass
"""


@activity
def wait(seconds=1):
    """Wait

    Make the robot wait for a specified number of seconds. Note that this activity is blocking. This means that subsequent activities will not occur until the the specified waiting time has expired.

    :parameter seconds: Time in seconds to wait
    :type seconds: int, optional

        :Example:

    >>> print('Start the wait')
    >>> wait()
    >>> print('The wait is over')

    Keywords
        wait, sleep, time, timeout, time-out, hold, pause

    Icon
        las la-hourglass        
    """
    from time import sleep

    sleep(seconds)


@activity
def wait_folder_exists(input_path, timeout=60):
    """Wait for folder

    Waits until a folder exists.
    Not that this activity is blocking and will keep the system waiting.

    :parameter input_path: Full path to folder.
    :type input_path: input_dir, optional
    :parameter timeout: Maximum time in seconds to wait before continuing. Default value is 60 seconds.
    :type timeout: int, optional

        :Example:

    >>> # Create a random folder
    >>> testfolder = create_folder()
    >>> # Wait for the snippet to be visible
    >>> wait_folder_exists(testfolder)

    Keywords
        image matching, wait, pause, vision, template, template matching

    Icon
        las la-hourglass
    """
    from time import sleep
    import os

    path = interpret_path(input_path)

    while not os.path.exists(path):
        sleep(1)
    return

    for _ in range(timeout):
        if os.path.exists(path):
            break
            sleep(1)


"""
Word Application
Icon: las la-file-word
"""


class Word:
    @activity
    def __init__(
        self, file_path=None, visible=True,
    ):
        """Start Word Application

        For this activity to work, Microsoft Office Word needs to be installed on the system.

        :parameter file_path: Enter a path to open Word with an existing Word file. If no path is specified a document will be initialized, this is the default value.
        :type file_path: input_file, optional
        :extension file_path: docx
        :parameter visible: Show Word in the foreground if True or hide if False, defaults to True.
        :type visible: bool, optional

            :Example:

        >>> word = Word()

        Keywords
            word, editor, text, text edit, office, document, microsoft word, doc, docx

        Icon
            lar la-file-word

        """
        only_supported_for("Windows")

        if file_path:
            self.file_path = interpret_path(file_path)
        else:
            self.file_path = file_path

        self.app = self._launch()
        # self.app.Visible = visible TODO: gives error for some PP versions

    def _launch(self):
        """Utility function to create the Word application scope object

        Utility function to create the Word application scope object

        :return: Application object (win32com.client)
        """
        try:
            import win32com.client

            app = win32com.client.dynamic.Dispatch("Word.Application")
            # app = win32com.client.dynamic.Dispatch("Word.Application")

        except:
            raise Exception(
                "Could not launch Word, do you have Microsoft Office installed on Windows?"
            )

        if self.file_path:
            app.Documents.Open(self.file_path)
        else:
            app.Documents.Add()

        return app

    @activity
    def save(self):
        """Save

        Save active Word document

            :Example:

        >>> # Start Word
        >>> word = Word()
        >>> word.append_text('This is sample text')
        >>> word.save_as('automagica_document.docx')

        Keywords
            word, save, document, doc, docx

        Icon
            lar la-file-word
        """
        self.app.ActiveDocument.Save()

        return self.file_path

    @activity
    def save_as(self, output_path):
        """Save As

        Save active Word document to a specific location

        :parameter output_path: Enter a path to open Word with an existing Word file.
        :type output_path: output_file
        :extension output_file: docx

            :Example:

        >>> # Start Word
        >>> word = Word()
        >>> word.append_text('This is sample text')
        >>> word.save_as('document.odt')

        Keywords
            word, save as, document, doc, docx

        Icon
            lar la-file-word
        """
        file_path = interpret_path(file_path)
        self.app.ActiveDocument.SaveAs(file_path)

        return file_path

    @activity
    def append_text(self, text):
        """Append text

        Append text at end of Word document.

        :parameter text: Text to append to document
        :type text: string, optional

            :Example:

        >>> # Start Word
        >>> word = Word()
        >>> word.append_text('This is sample text')

        Keywords
            word, editor, text, text edit, office, document, microsoft word, doc, docx

        Icon
            lar la-file-word
        """
        import win32com.client

        wc = win32com.client.constants
        self.app.Selection.EndKey(Unit=wc.wdStory)
        self.app.Selection.TypeText(text)

    @activity
    def replace_text(self, placeholder_text, replacement_text):
        """Replace text

        Can be used for example to replace arbitrary placeholder value. For example when 
        using template document, using 'XXXX' as a placeholder. Take note that all strings are case sensitive.

        :parameter placeholder_text: Placeholder text value in the document, this will be replaced, e.g. 'Company Name'
        :type placeholder_text: string
        :parameter replacement_text: Text to replace the placeholder values with. It is recommended to make this unique to avoid wrongful replacement, e.g. 'XXXX_placeholder_XXX'
        :type replacement_text: string

            :Example:

        >>> # Start Word
        >>> word = Word()
        >>> word.append_text('This is sample text')
        >>> word.replace_text('sample', 'real')

        Keywords
            word, replace, text, template, doc, docx

        Icon
            lar la-file-word
        """

        self.app.Selection.GoTo(0)
        self.app.Selection.Find.Text = placeholder_text
        self.app.Selection.Find.Replacement.Text = replacement_text
        self.app.Selection.Find.Execute(Replace=2, Forward=True)

    @activity
    def read_all_text(self, return_as_list=False):
        """Read all text

        Read all the text from a document

        :parameter return_as_list: Set this paramater to True to return text as a list of strings. Default value is False.
        :type return_as_list: bool, optional

        :return: Text from the document

            :Example:

        >>> # Start Word
        >>> word = Word()
        >>> word.append_text('This is sample text')
        >>> word.replace_text('sample', 'real')
        >>> word.read_all_text()
        'This is real text'

        Keywords
            word, extract, text, document, doc, docx

        Icon
            lar la-file-word
        """

        if return_as_list:
            return self.app.ActiveDocument.Content.Text.split("\r")

        return self.app.ActiveDocument.Content.Text.replace("\r", "\n")

    @activity
    def export_to_pdf(self, output_path=None):
        """Export to PDF

        Export the document to PDF

        :parameter output_path: Output path where PDF file will be exported to. Default path is home directory with filename 'pdf_export.pdf'.
        :type output_path: output_file, optional
        :extension output_path: pdf

            :Example:

        >>> # Start Word
        >>> word = Word()
        >>> word.append_text('This is sample text')
        >>> word.replace_text('sample', 'real')
        >>> word.export_to_pdf('output.pdf')

        Keywords
            word, pdf, document, export, save as, doc, docx

        Icon
            lar la-file-pdf

        """
        if not output_path:
            file_path = interpret_path(output_path, addition="pdf_export.pdf")
        else:
            file_path = interpret_path(output_path)

        self.app.ActiveDocument.ExportAsFixedFormat(
            OutputFileName=file_path,
            ExportFormat=17,
            OpenAfterExport=False,
            OptimizeFor=0,
            CreateBookmarks=1,
            DocStructureTags=True,
        )

    @activity
    def export_to_html(self, output_path=None):
        """Export to HTML

        Export to HTML

        :parameter file_path: Output path where HTML file will be exported to. Default path is home directory with filename 'html_export.html'.
        :type output_path: output_file, optional
        :extension output_path: html

            :Example:

        >>> # Start Word
        >>> word = Word()
        >>> word.append_text('This is sample text')
        >>> word.replace_text('sample', 'real')
        >>> word.export_to_html('output.html')

        Keywords
            word, html, document, export, save as, doc, docx

        Icon
            las la-html5

        """
        if not output_path:
            file_path = interpret_path(output_path, addition="html_export.html")
        else:
            file_path = interpret_path(output_path)

        import win32com.client

        wc = win32com.client.constants
        self.app.ActiveDocument.WebOptions.RelyOnCSS = 1
        self.app.ActiveDocument.WebOptions.OptimizeForBrowser = 1
        self.app.ActiveDocument.WebOptions.BrowserLevel = 0
        self.app.ActiveDocument.WebOptions.OrganizeInFolder = 0
        self.app.ActiveDocument.WebOptions.UseLongFileNames = 1
        self.app.ActiveDocument.WebOptions.RelyOnVML = 0
        self.app.ActiveDocument.WebOptions.AllowPNG = 1
        self.app.ActiveDocument.SaveAs(FileName=file_path, FileFormat=wc.wdFormatHTML)

    @activity
    def set_footers(self, text):
        """Set footers
        
        Set the footers of the document

        :parameter text: Text to put in the footer
        :type text: string

            :Example:

        >>> # Start Word
        >>> word = Word()
        >>> word.set_footers('This is a footer!')


        Keywords
            word, footer, footers, doc, docx

        Icon
            las la-heading
        """
        for section in self.app.ActiveDocument.Sections:
            for footer in section.Footers:
                footer.Range.Text = text

    @activity
    def set_headers(self, text):
        """Set headers

        Set the headers of the document

        :parameter text: Text to put in the header
        :type text: string

            :Example:

        >>> # Start Word
        >>> word = Word()
        >>> word.set_headers('This is a header!')

        Keywords
            word, header, headers, doc, docx

        Icon
            las la-subscript
        """
        for section in self.app.ActiveDocument.Sections:
            for footer in section.Headers:
                footer.Range.Text = text

    @activity
    def quit(self):
        """Quit Word

        This closes Word, make sure to use 'save' or 'save_as' if you would like to save before quitting.
        
            :Example:
            
        >>> # Open Word  
        >>> word = Word()
        >>> # Quit Word
        >>> word.quit()
        
        Keywords
            word, wordfile, doc quit, close, doc, docx

        Icon
             la-file-word
        """
        self.app.Application.Quit(0)


"""
Word File
Icon: las la-file-word
"""


class WordFile:
    @activity
    def __init__(self, file_path=None):
        """Read and Write Word files

        These activities can read, write and edit Word (docx) files without the need of having Word installed. 
        Note that, in contrary to working with the :func: 'Word' activities, a file get saved directly after manipulation.

        :parameter file_path: Enter a path to open Word with an existing Word file. If no path is specified a 'document.docx' will be initialized in the home directory, this is the default value. If a document with the same name already exists the file will be overwritten.
        :type file_path: input_file, optional
        :extension file_path: docx

            :Example:

        >>> wordfile = WordFile()
        >>> wordfile.append_text('Some sample text')
        >>> wordfile.read_all_text()
        'Some sample text'

        Keywords
            word, read, text, file

        Icon
            las la-file-word

        """
        self.file_path = file_path

        self.app = self._launch()

    def _launch(self):

        if not self.file_path:
            self.file_path = interpret_path(
                self.file_path, default_filename="document.docx"
            )

        else:
            self.file_path = interpret_path(file_path)

        from docx import Document

        document = Document()
        document.save(self.file_path)

    @activity
    def read_all_text(self, return_as_list=False):
        """Read all text

        Read all the text from the document

        :parameter return_as_list: Set this paramater to True to return text as a list of strings. Default value is False.
        :type return_as_list: bool, optional

        :return: Text of the document

            :Example:

        >>> wordfile = WordFile()
        >>> wordfile.append_text('Some sample text')
        >>> wordfile.read_all_text()
        'Some sample text'

        Keywords
            word, read, text, file

        Icon
            las la-file-word
        """

        from docx import Document

        document = Document(self.file_path)
        text = []
        for paragraph in document.paragraphs:
            text.append(paragraph.text)

        if return_as_list:
            return text
        return "\r".join(map(str, text))

    @activity
    def append_text(self, text, auto_save=True):
        """Append text

        Append text at the end of the document

        :parameter text: Text to append
        :type text: streing
        :parameter auto_save: Save document after performing activity. Default value is True
        :type auto_save: bool, optional

            :Example:

        >>> wordfile = WordFile()
        >>> wordfile.append_text('Some sample text')

        Keywords
            word, append text, add text

        Icon
            las la-file-word
        """
        from docx import Document

        document = Document(self.file_path)
        document.add_paragraph(text)

        if auto_save:
            document.save(self.file_path)

    @activity
    def save(self):
        """Save

        Save document

            :Example:

        >>> wordfile = WordFile()
        >>> wordfile.append_text('Some sample text')
        >>> wordfile.save()

        Keywords
            word, save, store

        Icon
            las la-file-word
        """
        from docx import Document

        document = Document(self.file_path)
        document.save(self.file_path)

    @activity
    def save_as(self, output_path):
        """Save as

        Save file on specified path

        :param output_path: Path to save Wordfile to
        :type output_path: output_file
        :extension output_path: docx

            :Example:
            
        >>> wordfile = WordFile()
        >>> wordfile.append_text('Some sample text')
        >>> wordfile.save_as('document.docx')

        Keywords
            word, save as, store

        Icon
            las la-file-word
        """
        from docx import Document

        document = Document(self.file_path)
        path = interpret_path(output_path)
        self.file_path = path
        document.save(path)

    @activity
    def set_headers(self, text, auto_save=True):
        """Set headers

        Set headers of Word document

        :parameter text: Text to put in the header
        :type text: string
        :parameter auto_save: Save document after performing activity. Default value is True
        :type auto_save: bool, optional

            :Example:
            
        >>> wordfile = WordFile()
        >>> wordfile.append_text('Some sample text')
        >>> wordfile.set_headers('This is a header')

        Keywords
            word, header text

        Icon
            las la-file-word
        """
        from docx import Document

        document = Document(self.file_path)
        document.add_heading(text)

        if auto_save:
            document.save(self.file_path)

    @activity
    def replace_text(self, placeholder_text, replacement_text, auto_save=True):
        """Replace all

        Replaces all occurences of a placeholder text in the document with a replacement text.

        Can be used for example to replace arbitrary placeholder value. 
        For example when using template slidedeck, using 'XXXX' as a placeholder.
        Take note that all strings are case sensitive.

        :parameter placeholder_text: Placeholder text value (string) in the document, this will be replaced, e.g. 'Company Name'
        :type placeholder_text: string
        :parameter replacement_text: Text (string) to replace the placeholder values with. It is recommended to make this unique to avoid wrongful replacement, e.g. 'XXXX_placeholder_XXX'
        :type replacement_text: string
        :parameter auto_save: Save document after performing activity. Default value is True
        :type auto_save: bool, optional

            :Example:
            
        >>> wordfile = WordFile()
        >>> wordfile.append_text('Some sample text')
        >>> wordfile.replace_text('sample', 'real')

        Keywords
            word, replace text, template

        Icon
            las la-file-word
        """
        from docx import Document

        document = Document(self.file_path)
        for paragraph in document.paragraphs:
            paragraph.text = paragraph.text.replace(placeholder_text, replacement_text)

        if auto_save:
            document.save(self.file_path)


"""
Outlook Application
Icon: las la-envelope
"""


class Outlook:
    @activity
    def __init__(self, account_name=None):
        """Start Outlook Application

        For this activity to work, Outlook needs to be installed on the system.

        :parameter account_name: Name of the account
        :type account_name: string, optional

            :Example:
            
        >>> outlook = Outlook()

        Keywords
            outlook, send e-mail, send mail

        Icon
            las la-mail-bulk
        
        """
        only_supported_for("Windows")

        self.app = self._launch()
        self.account_name = account_name

    def _launch(self):
        """Utility function to create the Outlook application scope object

        Utility function to create the Outlook application scope object        

        :return: Application object (win32com.client)
        """
        try:
            import win32com.client

            app = win32com.client.dynamic.Dispatch("outlook.application")

        except:
            raise Exception(
                "Could not launch Outlook, do you have Microsoft Office installed on Windows?"
            )

        return app

    @activity
    def send_mail(
        self, to_address, subject="", body="", html_body=None, attachment_paths=None,
    ):
        """Send e-mail

        Send an e-mail using Outlook

        :parameter to_address: The e-mail address the e-mail should be sent to
        :type to_addres: string, optional
        :parameter subject: The subject of the e-mail
        :type subject: string, optional
        :parameter body: The text body contents of the e-mail
        :type body: text, optional
        :parameter html_body: The HTML body contents of the e-mail (optional)
        :type html_body: text, optional
        :parameter attachment_paths: List of file paths to attachments
        :type attachment_paths: string, optional

            :Example:
            
        >>> outlook = Outlook()
        >>> outlook.send_mail('test@test.com', subject='Hello world', body='Hi there')

        Keywords
            outlook, send e-mail, send mail

        Icon
            las la-mail-bulk
        """
        # mapi = self.app.GetNamespace("MAPI")

        # Create a new e-mail
        mail = self.app.CreateItem(0)

        mail.To = to_address
        mail.Subject = subject
        mail.Body = body

        if html_body:
            mail.HTMLBody = html_body

        # Add attachments
        if attachment_paths:
            for attachment_path in attachment_paths:
                mail.Attachments.Add(attachment_path)

        # Send the e-mail
        mail.Send()

    @activity
    def get_folders(self, limit=999):
        """Retrieve folders

        Retrieve list of folders from Outlook

        :parameter limit: Maximum number of folders to retrieve
        :type limit: int, optional

            :Example:
            
        >>> outlook = Outlook()
        >>> outlook.get_folders()
        ['Inbox', 'Sent', ...]

        Keywords
            outlook, get folders, list folders

        Icon
            las la-mail-bulk
        """

        folders = []

        if self.account_name:
            found_folders = (
                self.app.GetNamespace("MAPI").Folders.Item(self.account_name).Folders
            )
        else:
            found_folders = self.app.GetNamespace("MAPI").Folders.Item(1).Folders
        for folder in found_folders:
            name = folder.Name
            folders.append(name)

        return folders

    @activity
    def get_mails(self, folder_name="Inbox", fields=None):
        """Retrieve e-mails

        Retrieve list of messages from Outlook

        :parameter folder_name: Name of the Outlook folder, can be found using `get_folders`.
        :type folder_name: string, optional
        :parameter limit: Number of messages to retrieve
        :type limit: int, optional
        :parameter fields: Fields (properties) of e-mail messages to give, requires tupl Stadard is 'Subject', 'Body', 'SentOn' and 'SenderEmailAddress'.
        :type fields: tuple, optional

        :return: List of dictionaries containing the e-mail messages with from, to, subject, body and html.

            :Example:
            
        >>> outlook = Outlook()
        >>> outlook.get_mails()

        Keywords
            outlook, retrieve e-mail, receive e-mails, process e-mails, get mails

        Icon
            las la-mail-bulk
        """

        if not fields:
            fields = ("Subject", "Body", "SenderEmailAddress")

        messages = []

        if self.account_name:
            found_folders = (
                self.app.GetNamespace("MAPI").Folders.Item(self.account_name).Folders
            )
        else:
            found_folders = self.app.GetNamespace("MAPI").Folders.Item(1).Folders
        for folder in found_folders:
            name = folder.Name
            if name == folder_name:
                break
        else:
            raise Exception(
                "Could not find the folder with name '{}'.".format(folder_name)
            )

        # Loop over the items in the folder
        for item in folder.Items:
            message = {}

            for key in fields:
                try:
                    message[key] = getattr(item, key)
                except AttributeError:
                    pass

            messages.append(message)

        return messages

    @activity
    def delete_mails(
        self,
        folder_name="Inbox",
        limit=0,
        subject_contains="",
        body_contains="",
        sender_contains="",
    ):
        """Delete e-mails

        Deletes e-mail messages in a certain folder. Can be specified by searching on subject, body or sender e-mail.

        :parameter folder_name: Name of the Outlook folder, can be found using `get_folders`
        :type folder_name: string, optional
        :parameter limit: Maximum number of e-mails to delete in one go
        :type limit: int
        :parameter subject_contains: Only delete e-mail if subject contains this
        :type subject_contains: string, optional
        :parameter body_contains: Only delete e-mail if body contains this
        :type body_contains: string, optional
        :parameter sender_contains: Only delete e-mail if sender contains this
        :type sender_contains: string, optional

            :Example:

        >>> outlook = Outlook()
        >>> outlook.delete_mails(subject_contains='hello')

        Keywords
            outlook, remove e-mails, delete mail, remove mail

        Icon
            las la-mail-bulk

        """
        # Find the appropriate folder
        if self.account_name:
            found_folders = (
                self.app.GetNamespace("MAPI").Folders.Item(self.account_name).Folders
            )
        else:
            found_folders = self.app.GetNamespace("MAPI").Folders.Item(1).Folders
        for folder in found_folders:
            name = folder.Name
            if name == folder_name:
                break
        else:
            raise Exception(
                "Could not find the folder with name '{}'.".format(folder_name)
            )

        # Loop over the items in the folder
        for i, item in enumerate(folder.Items):

            if limit:
                if i > limit:
                    break

            if subject_contains in item.Subject:
                if body_contains in item.Body:
                    if sender_contains in item.SenderEmailAddress:
                        item.Delete()

    @activity
    def move_mails(
        self,
        source_folder_name="Inbox",
        target_folder_name="Archive",
        limit=0,
        subject_contains="",
        body_contains="",
        sender_contains="",
    ):
        """Move e-mails

        Move e-mail messages in a certain folder. Can be specified by searching on subject, body or sender e-mail.

        :parameter source_folder_name: Name of the Outlook source folder from where e-mails will be moved, can be found using `get_folders`
        :type source_folder_name: string, optional
        :parameter target_folder_name: Name of the Outlook destination folder to where e-mails will be moved, can be found using `get_folders`
        :type target_folder_name: string, optional
        :parameter limit: Maximum number of e-mails to move in one go
        :type limit: int
        :parameter subject_contains: Only move e-mail if subject contains this
        :type subject_contains: string, optional
        :parameter body_contains: Only move e-mail if body contains this
        :type body_contains: string, optional
        :parameter sender_contains: Only move e-mail if sender contains this
        :type sender_contains: string, optional

            :Example:

        >>> outlook = Outlook()
        >>> outlook.move_mails(subject_contains='move me')

        Keywords
            outlook, move e-mail, move e-mail to folder

        Icon
            las la-mail-bulk
        """
        # Find the appropriate source folder
        if self.account_name:
            found_folders = (
                self.app.GetNamespace("MAPI").Folders.Item(self.account_name).Folders
            )
        else:
            found_folders = self.app.GetNamespace("MAPI").Folders.Item(1).Folders
        for source_folder in found_folders:
            name = source_folder.Name
            if name == source_folder_name:
                break
        else:
            raise Exception(
                "Could not find the folder with name '{}'.".format(source_folder_name)
            )

        # Find the appropriate target folder
        if self.account_name:
            found_folders = (
                self.app.GetNamespace("MAPI").Folders.Item(self.account_name).Folders
            )
        else:
            found_folders = self.app.GetNamespace("MAPI").Folders.Item(1).Folders
        for target_folder in found_folders:
            name = target_folder.Name
            if name == target_folder_name:
                break
        else:
            raise Exception(
                "Could not find the folder with name '{}'.".format(target_folder_name)
            )

        # Loop over the items in the folder
        for i, item in enumerate(source_folder.Items):

            if limit:
                if i > limit:
                    break

            if subject_contains in item.Subject:
                if body_contains in item.Body:
                    if sender_contains in item.SenderEmailAddress:
                        item.Move(target_folder)

    @activity
    def save_attachments(self, folder_name="Inbox", output_path=None):
        """Save attachments

        Save all attachments from certain folder

        :parameter folder_name: Name of the Outlook folder, can be found using `get_folders`.
        :type folder_name: string, optional
        :parameter output_path: Path where attachments will be saved. Default is the home directory.
        :type output_path: output_dir, optional

        :return: List of paths to saved attachments.

            :Example:
            
        >>> outlook = Outlook()
        >>> outlook.save_attachments()
        ['Attachment.pdf', 'Signature_image.jpeg']

        Keywords
            outlook, save attachments, download attachments, extract attachments

        Icon
            las la-mail-bulk
        """
        import os

        paths = []

        # Set to user home if no path specified
        output_path = interpret_path(output_path)

        # Find the appropriate folder
        if self.account_name:
            found_folders = (
                self.app.GetNamespace("MAPI").Folders.Item(self.account_name).Folders
            )
        else:
            found_folders = self.app.GetNamespace("MAPI").Folders.Item(1).Folders
        for folder in found_folders:
            name = folder.Name
            if name == folder_name:
                break
        else:
            raise Exception(
                "Could not find the folder with name '{}'.".format(folder_name)
            )
        # Loop over the items in the folder
        for item in folder.Items:
            for attachment in item.Attachments:
                path = os.path.join(output_path, attachment.FileName)
                attachment.SaveAsFile(path)
                paths.append(path)

        return paths

    @activity
    def get_contacts(self, fields=None):
        """Retrieve contacts

        Retrieve all contacts 

        :parameter fields: Fields can be specified as a tuple with their exact names. Standard value is None returning "LastName", "FirstName" and "Email1Address".
        :type fields: tuple, optional

        :return: List of dictionaries containing the contact details.

            :Example:

        >>> outlook = Outlook()
        >>> outlook.get_contacts()
        [
            {
                'LastName': 'Doe',
                'FirstName' : 'John',
                'Email1Address': 'john@test.com'
            }
        ]

        Keywords
            outlook, get contacts, download contacts, rolodex

        Icon
            las la-mail-bulk
        """
        import win32com.client

        if not fields:
            fields = ("LastName", "FirstName", "Email1Address")

        contacts = []

        mapi = self.app.GetNamespace("MAPI")

        data = mapi.GetDefaultFolder(win32com.client.constants.olFolderContacts)

        for item in data.Items:
            if item.Class == win32com.client.constants.olContact:
                contact = {}
                for key in item._prop_map_get_:
                    if key in fields:
                        if isinstance(getattr(item, key), (int, str)):
                            contact[key] = getattr(item, key)
                contacts.append(contact)

        return contacts

    @activity
    def add_contact(self, email, first_name="", last_name=""):
        """Add a contact

        Add a contact to Outlook contacts

        :parameter email: The e-mail address for the contact
        :type email: string, optional
        :parameter first_name: First name for the contact (optional)
        :type first_name: string, optional
        :parameter last_name: Last name for the contact (optional)
        :type last_name: string, optional

            :Example:

        >>> outlook = Outlook()
        >>> outlook.add_contact('sales@automagica.com')

        Keywords
            outlook, create contact, add contact

        Icon
            las la-mail-bulk
        """

        # Create a new contact
        contact = self.app.CreateItem(2)

        contact.Email1Address = email

        if first_name:
            contact.FirstName = first_name

        if last_name:
            contact.LastName = last_name

        contact.Save()

    @activity
    def quit(self):
        """Quit

        Close the Outlook application

            :Example:
            
        >>> outlook = Outlook()
        >>> outlook.quit()

        Keywords
            outlook, close, quit

        Icon
            las la-mail-bulk
        """
        self.app.Application.Quit()


"""
Excel Application
Icon: las la-file-excel
"""


class Excel:
    @activity
    def __init__(self, file_path=None, visible=True):
        """Start Excel Application

        For this activity to work, Microsoft Office Excel needs to be installed on the system.

        :parameter file_path: Enter a path to open Excel with an existing Excel file. If no path is specified a workbook will be initialized, this is the default value.
        :type file_path: input_file, optional
        :extension file_path: xlsx
        :parameter visible: Show Excel in the foreground if True or hide if False, defaults to True.
        :type visible: bool, optional
    
            :Example:

        >>> # Open Excel
        >>> excel = Excel()

        Keywords
            excel, add worksheet, add tab

        Icon
            las la-file-excel

        """
        only_supported_for("Windows")

        if not file_path:
            self.file_path = file_path
        else:
            self.file_path = interpret_path(file_path)

        self.app = self._launch()
        self.app.Visible = visible

    def _launch(self):
        """Utility function to create the Excel application scope object

        Utility function to create the Excel application scope object

        :return: Application object (win32com.client)
        
        """
        try:
            import win32com.client

            app = win32com.client.dynamic.Dispatch("Excel.Application")

        except:
            raise Exception(
                "Could not launch Excel, do you have Microsoft Office installed on Windows?"
            )

        if self.file_path:
            app.Workbooks.Open(self.file_path)
        else:
            app.Workbooks.Add()

        self.workbook = app.ActiveWorkbook

        return app

    @activity
    def add_worksheet(self, name=None):
        """Add worksheet

        Adds a worksheet to the current workbook

        :parmeter name: Give the sheet a name (optional)
        :type name: string, optional

            :Example:

        >>> # Open Excel
        >>> excel = Excel()
        >>> # Add a worksheet
        >>> excel.add_worksheet('My Example Worksheet')

        Keywords
            excel, add worksheet, add tab, insert worksheet, new worksheet

        Icon
            las la-file-excel
        
        """
        worksheet = self.workbook.Worksheets.Add()
        if name:
            worksheet.Name = name

    @activity
    def activate_worksheet(self, name):
        """Activate worksheet

        Activate a worksheet in the current Excel document by name

        :parameter name: Name of the worksheet to activate
        :type name: string, optional

            :Example:

        >>> # Open Excel   
        >>> excel = Excel()
        >>> # Add the first worksheet
        >>> excel.add_worksheet('My Example Worksheet')
        >>> # Add another worksheet
        >>> excel.add_worksheet('Another Worksheet')
        >>> # Activate the first worksheet
        >>> excel.activate_worksheet('My Example Worksheet)
        

        Keywords
            excel, activate worksheet, set worksheet, select worksheet, select tab, activate tab

        Icon
            las la-file-excel
        
        """
        for worksheet in self.workbook.Worksheets:
            if worksheet.Name == name:
                worksheet.Activate()

    @activity
    def save(self):
        """Save
        
        Save the current workbook. Defaults to homedir

            :Example:

        >>> # Open Excel
        >>> excel = Excel()
        >>> # Add the first worksheet
        >>> excel.add_worksheet('My Example Worksheet')
        >>> # Save the workbook
        >>> excel.save()

        Keywords
            excel, save, store

        Icon
            las la-file-excel
        """
        self.workbook.Save()

        return self.file_path

    @activity
    def save_as(self, output_path):
        """Save as

        Save the current workbook to a specific path

        :parameter output_path: Path where workbook will be saved.
        :type output_path: output_file
        :extension output_path: xlsx

            :Example:

        >>> # Open Excel
        >>> excel = Excel()
        >>> # Add the first worksheet
        >>> excel.add_worksheet('My Example Worksheet')
        >>> # Save the workbook to the current working directory
        >>> excel.save_as('output.xlsx')

        Keywords
            excel, save as, export

        Icon
            las la-file-excel
        
        """
        file_path = interpret_path(output_path)
        self.app.DisplayAlerts = False
        self.workbook.SaveAs(file_path)
        self.app.DisplayAlerts = True

        return file_path

    @activity
    def write_cell(self, column, row, value):
        """Write cell

        Write to a specific cell in the currently active workbook and active worksheet

        :parameter column: Column number (integer) to write
        :type column: int
        :parameter row: Row number (integer) to write
        :type row: int
        :parameter value: Value to write to specific cell
        :type value: string

            :Example:

        >>> # Open Excel
        >>> excel = Excel()
        >>> # Add the first worksheet
        >>> excel.add_worksheet('My Example Worksheet')
        >>> # Insert a text into the first cell
        >>> excel.write_cell(1,1, 'Hello World!')

        Keywords
            excel, cell, insert cell, insert data

        Icon
            las la-file-excel
        """
        self.workbook.ActiveSheet.Cells(row, column).Value = value

    @activity
    def read_cell(self, column, row):
        """Read cell

        Read a cell from the currently active workbook and active worksheet

        :parameter column: Column number (integer) to read
        :type column: int
        :parameter row: Row number (integer) to read
        :type row: int

        :return: Cell value

            :Example:
            
        >>> # Open Excel
        >>> excel = Excel()
        >>> # Add the first worksheet
        >>> excel.add_worksheet('My Example Worksheet')
        >>> # Insert a text into the first cell
        >>> excel.write_cell(1,1, 'Hello World!')
        >>> excel.read_cell(1,1)
        'Hello World!'

        Keywords
            excel, cell, read cell, read data

        Icon
            las la-file-excel
        """
        return self.workbook.ActiveSheet.Cells(row, column).Value

    @activity
    def write_range(self, range_, value):
        """Write range

        Write to a specific range in the currently active worksheet in the active workbook

        :parameter range_: Range to write to, e.g. "A1:D10"
        :type range_: string
        :parameter value: Value to write to range
        :type value: string
        
            :Example:

        >>> # Open Excel
        >>> excel = Excel()
        >>> # Add the first worksheet
        >>> excel.add_worksheet('My Example Worksheet')
        >>> # Insert a text in every cell in this range
        >>> excel.write_range('A1:D5', 'Hello World!')

        Keywords
            excel, cell, write range, read data

        Icon
            las la-file-excel
        """
        self.workbook.ActiveSheet.Range(range_).Value = value

    @activity
    def read_range(self, range_):
        """Read range
        
        Read a range of cells from the currently active worksheet in the active workbook

        :parameter range_: Range to read from, e.g. "A1:D10"
        :type range_: string

        :return value: Values in param range
        
            :Example:

        >>> # Open Excel
        >>> excel = Excel()
        >>> # Add the first worksheet
        >>> excel.add_worksheet('My Example Worksheet')
        >>> # Insert a text in every cell in this range
        >>> excel.write_range('A1:D5', 'Hello World!')
        >>> # Read the same range
        >>> excel.read_range('A1:D5')
        [['Hello World', 'Hello World', 'Hello World', 'Hello World'], ...]

        Keywords
            excel, cell, read range, read data

        Icon
            las la-file-excel
        """
        return self.workbook.ActiveSheet.Range(range_).Value

    @activity
    def run_macro(self, name):
        """Run macro

        Run a macro by name from the currently active workbook

        :parameter name: Name of the macro to run. 
        :type name: string

            :Example:

        >>> excel = Excel('excel_with_macro.xlsx')
        >>> # Run the macro
        >>> excel.run_macro('Macro1')

        Keywords
            excel, run macro, run vba

        Icon
            las la-file-excel
        """
        return self.app.Run(name)

    @activity
    def get_worksheet_names(self):
        """Get worksheet names

        Get names of all the worksheets in the currently active workbook

        :return: List of worksheet names

            :Example:

        >>> # Open Excel
        >>> excel = Excel()
        >>> # Add a worksheet
        >>> excel.add_worksheet('My Example Worksheet')
        >>> # Get all worksheet names
        >>> excel.get_worksheet_names()
        ['Sheet1', 'My Example Worksheet']

        Keywords
            excel, worksheet names, tab names

        Icon
            las la-file-excel
        """
        names = []

        for worksheet in self.workbook.Worksheets:
            names.append(worksheet.Name)

        return names

    @activity
    def get_table(self, name):
        """Get table

        Get table data from the currently active worksheet by name of the table

        :parameter name: Table name
        :type name: string

        :return: List of dictionaries for each row with as key the column name

            :Example:
            
        >>> # Open Excel
        >>> excel = Excel()
        >>> # Create a table (Table1)
        >>> data = [
            {
                'Column A': 'Data Row 1 for A',
                'Column B': 'Data Row 1 for B',
                'Column C': 'Data Row 1 for C',
            },
            {
                'Column A': 'Data Row 2 for A',
                'Column B': 'Data Row 2 for B',
                'Column C': 'Data Row 2 for C',
            }]
        >>> excel.insert_data_as_table(data)
        >>> # Get the table
        >>> excel.get_table('Table1')
        [['Column A', 'Column B', 'Column C'], ['Row 1 A Data', 'Row 1 B Data', 'Row 1 C Data'], ...]

        Keywords
            excel, worksheet names, tab names

        Icon
            las la-file-excel
        """
        data = []

        for worksheet in self.workbook.Worksheets:
            for list_object in worksheet.ListObjects:
                if list_object.Name == name:
                    for row in list_object.DataBodyRange.Value:
                        data_row = {}
                        for i, column in enumerate(list_object.HeaderRowRange.Value[0]):
                            data_row[column] = row[i]
                        data.append(data_row)

        return data

    @activity
    def activate_range(self, range_):
        """Activate range

        Activate a particular range in the currently active workbook

        :parameter range_: Range to activate, e.g. "A1:D10"

            :Example:

        >>> # Open Excel
        >>> excel = Excel()
        >>> # Activate a cell range
        >>> excel.activate_range('A1:D5')

        Keywords
            excel, activate range, make selection, select cells, select range

        Icon
            las la-file-excel
        """
        self.workbook.ActiveSheet.Range(range_).Select()

    @activity
    def activate_first_empty_cell_down(self):
        """Activate first empty cell down

        Activates the first empty cell going down

            :Example:

        >>> # Open Excel
        >>> excel = Excel()
        >>> # Write some cells
        >>> excel.write_cell(1, 1, 'Filled')
        >>> excel.write_cell(1, 2, 'Filled')
        >>> excel.write_cell(1, 3, 'Filled')
        >>> # Activate the first empty cell going down, in this case cell A4 or (1,4)
        >>> excel.activate_first_empty_cell_down()

        Keywords
            excel, first empty cell, down

        Icon
            las la-file-excel
        """
        column = self.app.ActiveCell.Column
        row = self.app.ActiveCell.Row
        for cell in self.workbook.ActiveSheet.Columns(column).Cells:
            if not cell.Value and cell.Row > row:
                cell.Select()
                break

    @activity
    def activate_first_empty_cell_right(self):
        """Activate first empty cell right

        Activates the first empty cell going right

            :Example:

        >>> # Open Excel
        >>> excel = Excel()
        >>> # Write some cells
        >>> excel.write_cell(1, 1, 'Filled')
        >>> excel.write_cell(1, 2, 'Filled')
        >>> excel.write_cell(1, 3, 'Filled')
        >>> # Activate the first empty cell going right, in this case cell B1 or (2,1)
        >>> excel.activate_first_empty_cell_right()

        Keywords
            excel, first empty cell, right

        Icon
            las la-file-excel
        """
        column = self.app.ActiveCell.Column
        row = self.app.ActiveCell.Row
        for cell in self.workbook.ActiveSheet.Rows(row).Cells:
            if not cell.Value and cell.Column > column:
                cell.Select()
                break

    @activity
    def activate_first_empty_cell_left(self):
        """Activate first empty cell left

        Activates the first empty cell going left

            :Example:

        >>> # Open Excel
        >>> excel = Excel()
        >>> excel.write_cell(1, 1, 'Filled')
        >>> excel.write_cell(1, 2, 'Filled')
        >>> excel.write_cell(1, 3, 'Filled')
        >>> excel.activate_first_empty_cell_left()

        Keywords
            excel, first empty cell, left

        Icon
            las la-file-excel
        """
        column = self.app.ActiveCell.Column
        row = self.app.ActiveCell.Row

        for i in range(column):
            if column - i > 0:
                cell = self.workbook.ActiveSheet.Cells(row, column - i)
                if not cell.Value:
                    cell.Select()
                    break

    @activity
    def activate_first_empty_cell_up(self):
        """Activate first empty cell up

        Activates the first empty cell going up

            :Example:

        >>> # Open Excel
        >>> excel = Excel()
        >>> # Write some cells
        >>> excel.write_cell(1, 1, 'Filled')
        >>> excel.write_cell(1, 2, 'Filled')
        >>> excel.write_cell(1, 3, 'Filled')
        >>> # Activate first empty cell
        >>> excel.activate_first_empty_cell_up()

        Keywords
            excel, first empty cell, up

        Icon
            las la-file-excel
        """
        column = self.app.ActiveCell.Column
        row = self.app.ActiveCell.Row

        for i in range(row):
            if row - i > 0:
                cell = self.workbook.ActiveSheet.Cells(row - i, column)
                if not cell.Value:
                    cell.Select()
                    break

    @activity
    def write_cell_formula(self, column, row, formula):
        """Write cell formula
        
        Write a formula to a particular cell

        :parameter column: Column number to write formula
        :type column: int
        :parameter row: Row number to write formula
        :type row: int
        :parameter value: Formula to write to specific cell e.g. "=10*RAND()"
        :type value: string

            :Example:

        >>> # Open Excel
        >>> excel = Excel()
        >>> # Write a formula to the first cell
        >>> excel.write_cell_formula(1, 1, '=1+1)

        Keywords
            excel, insert formula, insert calculation, insert calculated cell

        Icon
            las la-file-excel
        """
        self.workbook.ActiveSheet.Cells(row, column).Formula = formula

    @activity
    def read_cell_formula(self, column, row, formula):
        """Read cell formula

        Read the formula from a particular cell

        :parameter column: Column number to read formula
        :type column: int
        :parameter row: Row number to read formula
        :type row: int

        :return: Cell value

            :Example:

        >>> # Open Excel
        >>> excel = Excel()
        >>> # Write a formula to the first cell
        >>> excel.write_cell_formula(1, 1, '=1+1)
        >>> # Read the cell
        >>> excel.read_cell_formula(1, 1)
        '=1+1'
        
        Keywords
            excel, read formula, read calculation

        Icon
            las la-file-excel
        """
        return self.workbook.ActiveSheet.Cells(row, column).Formula

    @activity
    def insert_empty_row(self, row):
        """Insert empty row

        Inserts an empty row to the currently active worksheet

        :parameter row: Row number where to insert empty row e.g 1
        :type row: int
            
            :Example:

        >>> # Open Excel              
        >>> excel = Excel()
        >>> excel.write_cell(1, 1, 'Filled')
        >>> excel.write_cell(1, 2, 'Filled')
        >>> excel.write_cell(1, 3, 'Filled')
        >>> excel.insert_empty_row(2)
        
        Keywords
            excel, insert row, add row, empty row

        Icon
            las la-file-excel
        """
        row_range = "A" + str(row)
        self.workbook.ActiveSheet.Range(row_range).EntireRow.Insert()

    @activity
    def insert_empty_column(self, column):
        """Insert empty column

        Inserts an empty column in the currently active worksheet. Existing columns will shift to the right.

        :parameter column: Column letter where to insert empty column e.g. 'A'
        :type column: string
            
            :Example:

        >>> # Open Excel
        >>> excel = Excel()
        >>> excel.write_cell(1, 1, 'Filled')
        >>> excel.write_cell(2, 2, 'Filled')
        >>> excel.write_cell(3, 3, 'Filled')
        >>> excel.insert_empty_column('A')
        
        Keywords
            excel, insert column, add column

        Icon
            las la-file-excel
        """
        column_range = str(column) + "1"
        self.workbook.ActiveSheet.Range(column_range).EntireColumn.Insert()

    @activity
    def delete_row(self, row):
        """Delete row in Excel

        Deletes a row from the currently active worksheet. Existing data will shift up.

        :parameter row: Row number (integer) where to delete row e.g 1
        :type row: int

            :Example:

        >>> # Open Excel              
        >>> excel = Excel()
        >>> excel.write_cell(1, 1, 'Filled')
        >>> excel.write_cell(2, 2, 'Filled')
        >>> excel.write_cell(3, 3, 'Filled')
        >>> excel.delete_row(2)
        
        Keywords
            excel, delete row, remove row

        Icon
            las la-file-excel
        """
        row_range = "A" + str(row)
        self.workbook.ActiveSheet.Range(row_range).EntireRow.Delete()

    @activity
    def delete_column(self, column):
        """Delete column

        Delete a column from the currently active worksheet. Existing columns will shift to the left.

        :parameter column: Column letter (string) where to delete  column e.g. 'A'
        :type column: string

            :Example:

        >>> # Open Excel              
        >>> excel = Excel()
        >>> excel.write_cell(1, 1, 'Filled')
        >>> excel.write_cell(2, 2, 'Filled')
        >>> excel.write_cell(3, 3, 'Filled')
        >>> excel.delete_column('A')
        
        Keywords
            excel, delete column, remove column

        Icon
            las la-file-excel
        """
        column_range = str(column) + "1"
        self.workbook.ActiveSheet.Range(column_range).EntireColumn.Delete()

    @activity
    def export_to_pdf(self, output_path=None):
        """Export to PDF

        Export to PDF

        :parameter path: Output path where PDF file will be exported to. Default path is home directory with filename 'pdf_export.pdf'.
        :type output_path: output_file, optional
        :extension output_path: pdf
        
            :Example:

        >>> # Open Excel              
        >>> excel = Excel()
        >>> excel.write_cell(1, 1, 'Filled')
        >>> excel.write_cell(2, 2, 'Filled')
        >>> excel.write_cell(3, 3, 'Filled')
        >>> excel.export_to_pdf('output.pdf')
        
        Keywords
            excel, save as pdf, export to pdf, export as pdf

        Icon
            las la-file-excel
        """
        if not file_path:
            file_path = interpret_path(output_path, addition="pdf_export.pdf")
        else:
            file_path = interpret_path(output_path)

        self.workbook.ActiveSheet.ExportAsFixedFormat(0, output_path, 0, True, True)

    @activity
    def insert_data_as_table(self, data, range_="A1", table_style="TableStyleMedium2"):
        """Insert data as table
        
        Insert list of dictionaries as a table in Excel

        :parameter data: List of dictionaries to write as table
        :type data: string
        :parameter range_: Range or startingpoint for table e.g. 'A1'
        :type range_: string, optional
        
            :Example:
            
        >>> excel = Excel()
        >>> data = [
            {
                'Column A': 'Data Row 1 for A',
                'Column B': 'Data Row 1 for B',
                'Column C': 'Data Row 1 for C',
            },
            {
                'Column A': 'Data Row 2 for A',
                'Column B': 'Data Row 2 for B',
                'Column C': 'Data Row 2 for C',
            }
        >>> excel.insert_data_as_table(data)
        
        Keywords
            excel, insert data, insert table, create table

        Icon
            las la-file-excel
        """
        row = self.workbook.ActiveSheet.Range(range_).Row
        column = self.workbook.ActiveSheet.Range(range_).Column

        column_names = list(data[0].keys())
        data_values = [[d[key] for key in data[0].keys()] for d in data]

        values = [column_names] + data_values
        for i in range(len(values)):
            for j in range(len(values[0])):
                self.workbook.ActiveSheet.Cells(row + i, column + j).Value = values[i][
                    j
                ]

        start_cell = self.workbook.ActiveSheet.Cells(row, column)
        end_cell = self.workbook.ActiveSheet.Cells(row + i, column + j)
        self.workbook.ActiveSheet.Range(start_cell, end_cell).Select()
        self.app.ActiveSheet.ListObjects.Add().TableStyle = table_style

    @activity
    def read_worksheet(self, name=None, headers=False):
        """Read worksheet

        Read data from a worksheet as a list of lists

        :parameter name: Optional name of worksheet to read. If no name is specified will take active sheet
        :type name: string, optional
        :parameter headers: Boolean to treat first row as headers. Default value is False
        :type headers: bool, optional

        :return: List of dictionaries with sheet data
        
            :Example:

        >>> # Open excel    
        >>> excel = Excel()
        >>> Write some cells
        >>> excel.write_cell(1, 1, 'A')
        >>> excel.write_cell(1, 2, 'B')
        >>> excel.write_cell(1, 3, 'C')
        >>> excel.read_worksheet()
        [['A'],['B'],['C']]
        
        Keywords
            excel, read worksheet, export data, read data

        Icon
            las la-file-excel
        """
        if name:
            self.activate_worksheet(name)

        data = self.workbook.ActiveSheet.UsedRange.Value

        if isinstance(data, str):
            return data

        # Remove empty columns and rows
        data = [list(x) for x in data if any(x)]
        transposed = list(map(list, zip(*data)))
        transposed = [row for row in transposed if any(row)]
        data = list(map(list, zip(*transposed)))

        if headers:
            header_row = data[0]
            data = data[1:]
            data = [
                {column: row[i] for i, column in enumerate(header_row)} for row in data
            ]

        return data

    @activity
    def quit(self):
        """Quit Excel

        This closes Excel, make sure to use 'save' or 'save_as' if you would like to save before quitting.
        
            :Example:
            
        >>> # Open Excel  
        >>> excel = Excel()
        >>> # Quit Excel
        >>> excel.quit()
        
        Keywords
            excel, exit, quit, close

        Icon
            las la-file-excel
        """
        self.app.Application.Quit()


"""
Excel File
Icon: las la-file-excel
"""


class ExcelFile:
    @activity
    def __init__(self, file_path=None):
        """Read and Write xlsx files. 

        This activity can read, write and edit Excel (xlsx) files without the need of having Excel installed. 
        Note that, in contrary to working with the :func: 'Excel' activities, a file get saved directly after manipulation.

        :parameter file_path: Enter a path to open Excel with an existing Excel file. If no path is specified a 'workbook.xlsx' will be initialized in the home directory, this is the default value. If a workbook with the same name already exists the file will be overwritten.
        :type file_path: input_file, optional
        :extension file_path: xlsx
        
            :Example:

        >>> # Open a new Excel file
        >>> excel_file = ExcelFile()
        
        Keywords
            excel, open, start, xlsx

        Icon
            las la-file-excel    

        """
        import openpyxl
        import os

        if file_path:
            file_path = interpret_path(file_path)

        self.file_path = file_path
        self.sheet_name = None

        if self.file_path:
            self.book = openpyxl.load_workbook(self.file_path)

        else:
            path = os.path.join(os.path.expanduser("~"), "workbookske.xlsx")
            self.book = openpyxl.Workbook()
            self.book.create_sheet("My Example Worksheet")
            self.book.save(path)
            self.file_path = path

    @activity
    def to_dataframe(self):
        """Export file to dataframe 

        Export to pandas dataframe

            :Example:

        >>> # Open a new Excel file
        >>> excel_file = ExcelFile()
        >>> # Convert to Dataframe
        >>> df = excel_file.to_dataframe()
        
        Keywords
            excel, open, start, xlsx, dataframe,

        Icon
            las la-file-excel    

        """

        import pandas as pd

        return pd.read_excel(self.file_path)

    @activity
    def activate_worksheet(self, name):
        """Activate worksheet

        Activate a worksheet. By default the first worksheet is activated.

        :parameter name: Name of the worksheet to activate.        
        :type name: string, optional
        
            :Example:

        >>> # Open a new Excel file
        >>> excel_file = ExcelFile()
        >>> # Add some worksheets
        >>> excel_file.add_worksheet('My Example Worksheet')
        >>> excel_file.add_worksheet('Another Worksheet')
        >>> # Activate a worksheet
        >>> excel_file.activate_worksheet('My Example Worksheet')
        
        Keywords
            excel, activate tab, activate worksheet

        Icon
            las la-file-excel
        """

        self.sheet_name = name

    @activity
    def save_as(self, output_path):
        """Save as

        Save file as

        :parameter file_path: Path where workbook will be saved
        :type output_path: output_file
        :extension output_path: xlsx
        
            :Example:

        >>> # Open a new Excel file
        >>> excel_file = ExcelFile()
        >>> # Ad a worksheet
        >>> excel_file.add_worksheet('My Example Worksheet')
        >>> # Save the Excel file
        >>> excel_file.save_as('output.xlsx')
        
        Keywords
            excel, save as, export, save

        Icon
            las la-file-excel
        """
        file_path = interpret_path(output_path)
        self.book.save(file_path)

    @activity
    def save(self):
        """Save as

        Save file
        
            :Example:

        >>> # Open a new Excel file
        >>> excel_file = ExcelFile()
        >>> # Ad a worksheet
        >>> excel_file.add_worksheet('My Example Worksheet')
        >>> # Save the Excel file
        >>> excel_file.save()
        
        Keywords
            excel, save as, export, save

        Icon
            las la-file-excel
        """
        self.book.save(self.file_path)

    @activity
    def write_cell(self, column, row, value, auto_save=True):
        """Write cell

        Write a cell based on column and row

        :parameter column: Column number (integer) to write
        :type column: int
        :parameter row: Row number (integer) to write
        :type row: int
        :parameter value: Value to write to specific cell
        :type value: string
        :parameter auto_save: Save document after performing activity. Default value is True
        :type auto_save: bool, optional
        
            :Example:

        >>> # Open a new Excel file
        >>> excel_file = ExcelFile()
        >>> # Add a worksheet
        >>> excel_file.add_worksheet('My Example Worksheet')
        >>> excel_file.write_cell(1, 1, 'Filled!')
        
        Keywords
            excel, write cell, insert data

        Icon
            las la-file-excel
        """
        if self.sheet_name:
            sheet = self.book[self.sheet_name]
        else:
            sheet = self.book.active

        sheet.cell(row=row, column=column).value = value

        if auto_save:
            self.book.save(self.file_path)

    @activity
    def read_cell(self, column, row):
        """Read cell

        Read a cell based on column and row

        :parameter column: Column number (integer) to read
        :type colunm: int
        :parameter row: Row number (integer) to read
        :type row: int

        :return: Cell value

           :Example:

        >>> # Open a new Excel file
        >>> excel_file = ExcelFile()
        >>> # Add a worksheet
        >>> excel_file.add_worksheet('My Example Worksheet')
        >>> # Write the first cell
        >>> excel_file.write_cell(1, 1, 'Filled!')
        >>> # Read the first cell
        >>> excel_file.read_cell(1, 1)
        'Filled!'
        
        Keywords
            excel, read cell, read

        Icon
            las la-file-excel

        """
        if self.sheet_name:
            sheet = self.book[self.sheet_name]
        else:
            sheet = self.book.active

        return sheet.cell(row=row, column=column).value

    @activity
    def add_worksheet(self, name, auto_save=True):
        """Add worksheet

        Add a worksheet

        :parameter name: Name of the worksheet to add
        :type name: string
        :parameter auto_save: Save document after performing activity. Default value is True
        :type auto_save: bool, optional

           :Example:

        >>> # Open a new Excel file
        >>> excel_file = ExcelFile()
        >>> # Add a worksheet
        >>> excel_file.add_worksheet('My Example Worksheet')
        >>> # List all the worksheets
        >>> excel.get_worksheet_names()

        
        Keywords
            excel, add worksheet, worksheet

        Icon
            las la-file-excel
        """

        self.book.create_sheet(name)
        if auto_save:
            self.book.save(self.file_path)

    @activity
    def get_worksheet_names(self):
        """Get worksheet names

        Get worksheet names

        :return: List of worksheet names

           :Example:
            
        >>> # Open a new Excel file
        >>> excel_file = ExcelFile()
        >>> # Add some worksheets
        >>> excel_file.add_worksheet('My Example Worksheet')
        >>> excel_file.add_worksheet('Another Worksheet')
        >>> # Get the worksheet names
        >>> excel_file.get_worksheet_names()
        ['My Example Worksheet', 'Another Worksheet']
        
        Keywords
            excel, worksheet names, worksheet,

        Icon
            las la-file-excel

        """

        return self.book.sheetnames


"""
PowerPoint Application
Icon: las la-file-powerpoint
"""


class PowerPoint:
    @activity
    def __init__(self, file_path=None, visible=True, add_slide=True):
        """Start PowerPoint Application

        For this activity to work, PowerPoint needs to be installed on the system.

        :parameter file_path: Enter a path to open an existing PowerPoint presentation. If no path is specified a new presentation will be initialized, this is the default value.
        :type file_path: input_file, optional
        :extension file_path: pptx
        :parameter visible: Show PowerPoint in the foreground if True or hide if False, defaults to True.
        :type visible: bool, optional
        :parameter add_slide: Add an initial empty slide when creating new PowerPointfile, this prevents errors since most manipulations require a non-empty presentation. Default value is True
        :type add_slide: bool, optional
        
            :Example:
            
        >>> # Start PowerPoint
        >>> powerpoint = PowerPoint()

        Keywords
            powerpoint, ppt

        Icon
            las la-file-powerpoint
        
        """
        only_supported_for("Windows")

        if not file_path:
            self.file_path = file_path
        else:
            self.file_path = interpret_path(file_path)

        self.app = self._launch(self.file_path)
        self.app.Visible = visible

    def _launch(self, path):
        """Utility function to create the Excel application scope object

        Utility function to create the Excel application scope object

        :return: Application object (win32com.client)
        """
        try:
            import win32com.client

            app = win32com.client.dynamic.Dispatch("PowerPoint.Application")

        except:
            raise Exception(
                "Could not launch PowerPoint, do you have Microsoft Office installed on Windows?"
            )

        if self.file_path:
            return app.Presentations.Open(self.file_path)
        else:
            return app.Presentations.Add()

    @activity
    def save_as(self, output_path):
        """Save PowerPoint
        
        Save PowerPoint Slidedeck

        :parameter file_path: Save the PowerPoint presentation.
        :type output_path: output_file
        :extension output_path: pptx

            :Example:
            
        >>> # Start PowerPoint
        >>> powerpoint = PowerPoint()
        >>> # Add a first slide
        >>> powerpoint.add_slide()
        >>> # Save the PowerPoint presentation
        >>> powerpoint.save_as('AmazingPresentation.pptx')

        Keywords
            powerpoint, ppt, save, save as, save powerpoint

        Icon
            las la-file-powerpoint

        """
        file_path = interpret_path(output_path)
        self.app.SaveAs(file_path)

        return file_path

    @activity
    def save(self):
        """Save PowerPoint
        
        Save PowerPoint Slidedeck

            :Example:
            
        >>> # Start PowerPoint
        >>> powerpoint = PowerPoint()
        >>> # Add a first slide
        >>> powerpoint.add_slide()
        >>> # Save the PowerPoint presentation
        >>> powerpoint.save_as('AmazingPresentation.pptx')

        Keywords
            powerpoint, ppt, save, save as, save powerpoint

        Icon
            las la-file-powerpoint

        """
        self.app.SaveAs(self.file_path)

        return self.file_path

    @activity
    def quit(self):
        """Close PowerPoint Application

        Close PowerPoint

        :parameter index: Index where the slide should be inserted. Default value is as final slide.
        :parmeter type: Type of the slide to be added. Supports following types: blank, chart, text, title and picture.

            :Example:
            
        >>> # Start PowerPoint
        >>> powerpoint = PowerPoint()
        >>> # Close PowerPoint
        >>> powerpoint.quit()


        Keywords
            powerpoint, ppt, quit, exit

        Icon
            las la-file-powerpoint

        """
        self.app.Application.Quit()

    @activity
    def add_slide(self, index=None, type="blank"):
        """Add PowerPoint Slides

        Adds slides to a presentation

        :parameter index: Index where the slide should be inserted. Default value is as final slide.
        :type index: int, optional
        :parmeter type: Type of the slide to be added. Supports following types: blank, chart, text, title and picture.
        :options type: ['blank', 'chart', 'text', 'title', 'picture']


            :Example:
            
        >>> # Start PowerPoint
        >>> powerpoint = PowerPoint()
        >>> # Add a first slide
        >>> powerpoint.add_slide()


        Keywords
            powerpoint, ppt, add, add slide  powerpoint, slides

        Icon
            las la-file-powerpoint

        """
        if type == "blank":
            type_id = 12
        if type == "chart":
            type_id = 8
        if type == "text":
            type_id = 2
        if type == "title":
            type_id = 1
        if type == "picture":
            type_id = 36

        if not index:
            index = self.app.Slides.Count + 1

        return self.app.Slides.Add(index, type_id)

    @activity
    def number_of_slides(self):
        """Slide count
        
        Returns the number of slides

            :Example:
            
        >>> # Start PowerPoint
        >>> powerpoint = PowerPoint()
        >>> # Add some slides
        >>> powerpoint.add_slide()
        >>> powerpoint.add_slide()
        >>> # Show number of slides
        >>> powerpoint.number_of_slides()

        Keywords
            powerpoint, ppt, slide count, number of slides

        Icon
            las la-file-powerpoint
        """
        return self.app.Slides.Count

    @activity
    def add_text(
        self,
        text,
        index=None,
        font_size=48,
        font_name=None,
        bold=False,
        margin_bottom=100,
        margin_left=100,
        margin_right=100,
        margin_top=100,
    ):
        """Text to slide
        
        Add text to a slide

        :parameter index: Slide index to add text. If none is specified, a new slide will be added as final slide
        :type index: int, optional
        :parameter text: Text to be added
        :type text: string, optional
        :parameter font_size: Fontsize, default value is 48
        :type font_size: int, optional
        :parameter font_name: Fontname, if not specified will take default PowerPoint font
        :type font_name: string, optional
        :parameter bold: Toggle bold with True or False, default value is False
        :type bold: bool, optional
        :parameter margin_bottom: Margin from the bottom in pixels, default value is 100 pixels
        :type margin_bottom: int, optional
        :parameter margin_left: Margin from the left in pixels, default value is 100 pixels
        :type margin_left: int, optional
        :parameter margin_right: Margin from the right in pixels, default value is 100 pixels
        :type margin_right: int, optional
        :parameter margin_top: Margin from the top in pixels, default value is 100 pixels
        :type margin_top: int, optional

            :Example:
            
        >>> # Start PowerPoint
        >>> powerpoint = PowerPoint()
        >>> # Add slide with text
        >>> powerpoint.add_text(text='Sample Text')


        Keywords
            powerpoint, ppt, text, add text, slides
        Icon
            las la-file-powerpoint

        """

        if not index:
            index = self.app.Slides.Count + 1
            self.app.Slides.Add(index, 12)
        text_box = (
            self.app.Slides(index)
            .Shapes.AddTextbox(1, 100, 100, 200, 50)
            .TextFrame.TextRange
        )
        text_box.Text = text
        text_box.Font.Size = font_size
        if font_name:
            text_box.Font.Name = font_name
        text_box.Font.Bold = bold

    @activity
    def delete_slide(self, index=None):
        """Delete slide

        Delete a slide

        :parameter index: Slide index to be deleted. If none is specified, last slide will be deleted
        :type index: int, optional

            :Example:

        >>> # Start PowerPoint
        >>> powerpoint = PowerPoint()
        >>> # Add some slides
        >>> powerpoint.add_slide()
        >>> powerpoint.add_slide()
        >>> # Delete last slide
        >>> powerpoint.delete_slide()

        Keywords
            powerpoint, ppt, delete, delete slide

        Icon
            las la-file-powerpoint

        """
        if not index:
            index = self.app.Slides.Count

        return self.app.Slides(index).Delete()

    @activity
    def replace_text(self, placeholder_text, replacement_text):
        """Replace all occurences of text in PowerPoint slides

        Can be used for example to replace arbitrary placeholder value in a PowerPoint. 
        For example when using a template slidedeck, using 'XXXX' as a placeholder.
        Take note that all strings are case sensitive.

        :parameter placeholder_text: Placeholder value (string) in the PowerPoint, this will be replaced, e.g. 'Company Name'
        :type placeholder_text: string
        :parameter replacement_text: Text (string) to replace the placeholder values with. It is recommended to make this unique in your PowerPoint to avoid wrongful replacement, e.g. 'XXXX_placeholder_XXX'
        :type replacement_text: string

            :Example:
            
        >>> # Start PowerPoint
        >>> powerpoint = PowerPoint()
        >>> # Add some slides with text
        >>> powerpoint.add_text(text='Hello, my name is placeholder')
        >>> # Change 'placeholder' to the word 'robot
        >>> powerpoint.replace_text(placeholder_text = 'placeholder', replacement_text ='robot')

        Keywords
            powerpoint, ppt, replace, placeholder

        Icon
            las la-file-powerpoint

        """
        for slide in self.app.Slides:
            for shape in slide.Shapes:
                shape.TextFrame.TextRange.Text = shape.TextFrame.TextRange.Text.replace(
                    placeholder_text, replacement_text
                )

    @activity
    def export_to_pdf(self, output_path=None):
        """PowerPoint to PDF
        
        Export PowerPoint presentation to PDF file

        :parameter path: Output path where PDF file will be exported to. Default path is home directory with filename 'pdf_export.pdf'.
        :type output_path: output_file
        :extension output_path: pdf

            :Example:
            
        >>> # Start PowerPoint
        >>> powerpoint = PowerPoint()
        >>> # Add some slides with text
        >>> powerpoint.add_text(text='Robots are cool')
        >>> # Export to pdf
        >>> powerpoint.export_to_pdf()

        Keywords
            powerpoint, ppt, export, pdf

        Icon
            las la-file-powerpoint

        """

        if self.app.Slides.Count == 0:
            raise Exception(
                "Please add a slide first bedore exporting the presentation."
            )

        if not path:
            import os

            path = os.path.join(os.path.expanduser("~"), "pdf_export.pdf")

        return self.app.ExportAsFixedFormat2(path, 2, PrintRange=None)

    @activity
    def export_slides_to_images(self, output_path=None, type="png"):
        """Slides to images
        
        Export PowerPoint slides to seperate image files

        :parameter output_path: Output path where image files will be exported to. Default path is home directory.
        :type output_put: output_dir
        :parameter type: Output type of the images, supports 'png' and 'jpg' with 'png' as default value
        :options type: ['jpg', 'png']

            :Example:
            
        >>> # Start PowerPoint
        >>> powerpoint = PowerPoint()
        >>> # Add some slides with text
        >>> powerpoint.add_text(text='Robots are cool')
        >>> powerpoint.add_text(text='Humans are cooler')
        >>> # Export slides to images
        >>> powerpoint.export_slides_to_images()

        Keywords
            powerpoint, ppt, export, png, image, slides to image

        Icon
            las la-file-powerpoint

        """

        if self.app.Slides.Count == 0:
            raise Exception(
                "Please add a slide first bedore exporting the presentation."
            )

        if not output_path:
            import os

            path = os.path.expanduser("~")

        return self.app.Export(path, "png")


"""
Office 365
Icon: las la-cloud
"""


@activity
def send_email_with_outlook365(client_id, client_secret, to_email, subject="", body=""):
    """Send email Office Outlook 365

    Send email Office Outlook 365

    :parameter client_id: Client id for office 365 account
    :type cliend_id: string
    :parameter client_secret: Client secret for office 365 account
    :type client_secret: string
    :parameter to_email: E-mail to send to
    :type to_email: string
    :parameter subject: Optional subject
    :type subject: string, optional
    :parameter body: Optional body of the email
    :type body: string, optional

        :Example:

    >>> # Send email to 'robot@automagica.com'
    >>> send_email_with_outlook365('SampleClientID', 'SampleClientSecret', 'robot@automagica.com')

    Keywords
        mail, office 365, outlook, email, e-mail

    Icon
        las la-envelope
    """
    from O365 import Account

    credentials = (client_id, client_secret)

    account = Account(credentials)
    m = account.new_message()
    m.to.add(to_email)
    m.subject = subject
    m.body = body
    m.send()


"""
Salesforce
Icon: lab la-salesforce
"""


@activity
def salesforce_api_call(action, key, parameters={}, method="get", data={}):
    """Salesforce API

    Activity to make calls to Salesforce REST API.

    :parameter action: Action (the URL)
    :type action: string
    :parameter key: Authorisation key 
    :type key: string
    :parameter parameters: URL params
    :type parameters: string
    :parameter method: Method (get, post or patch)
    :options method: ['get', 'post', 'patch']
    :parameter data: Data for POST/PATCH.
    :type data: string

    :return: API data

        :Example:

    >>> spf_api_call('action', 'key', 'parameters')
    Response

    Keywords
        salesforce

    Icon
        lab la-salesforce

    """
    from automagica.httpclient import http_client

    headers = {
        "Content-type": "application/json",
        "Accept-Encoding": "gzip",
        "Authorization": "Bearer " + key,
    }

    if method == "get":
        r = http_clientrequest(
            method,
            instance_url + action,
            headers=headers,
            params=parameters,
            timeout=30,
        )
    elif method in ["post", "patch"]:
        r = http_clientrequest(
            method,
            instance_url + action,
            headers=headers,
            json=data,
            params=parameters,
            timeout=10,
        )
    else:
        raise ValueError("Method should be get or post or patch.")
    print("Debug: API %s call: %s" % (method, r.url))
    if r.status_code < 300:
        if method == "patch":
            return None
        else:
            return r.json()
    else:
        raise Exception("API error when calling %s : %s" % (r.url, r.content))


"""
E-mail (SMTP)
Icon: las la-at
"""


@activity
def send_mail_smtp(
    smtp_host, smtp_user, smtp_password, to_address, subject="", message="", port=587
):
    """Mail with SMTP

    This function lets you send emails with an e-mail address. 

    :parameter smpt_host: The host of your e-mail account. 
    :type smpt_host: string
    :parameter smpt_user: The password of your e-mail account
    :type smtp_user: string
    :parameter smpt_password: The password of your e-mail account
    :type smpt_password: string
    :parameter to_address: The destination is the receiving mail address. 
    :type to_address: string
    :parameter subject: The subject 
    :type subject: string, optional
    :parameter message: The body of the mail
    :type message: text, optional
    :parameter port: The port variable is standard 587. In most cases this argument can be ignored, but in some cases it needs to be changed to 465.
    :type port: int, optional

        :Example:

    >>> send_mail_smpt('robot@automagica.com', 'SampleUser', 'SamplePassword', 'robotfriend@automagica.com')

    Keywords
        mail, e-mail, email smpt

    Icon
        las la-mail-bulk

    """
    BODY = "\r\n".join(
        [
            "To: %s" % destination,
            "From: %s" % user,
            "Subject: %s" % subject,
            "",
            message,
        ]
    )
    smtpObj = smtplib.SMTP(host, port)
    smtpObj.ehlo()
    smtpObj.starttls()
    smtpObj.login(user, password)
    smtpObj.sendmail(user, destination, BODY)
    smtpObj.quit()


"""
Windows OS
Icon: lab la-windows
"""


@activity
def find_window_title(searchterm, partial=True):
    """Find window with specific title

    Find a specific window based on the name, either a perfect match or a partial match.

    :parameter searchterm: Ttile to look for, e.g. 'Calculator' when looking for the Windows calculator
    :type searchterm: string
    :parameter partial: Option to look for titles partially, e.g. 'Edge' will result in finding 'Microsoft Edge' when partial is set to True. Default value is True
    :type pertial: bool, optional

        :Example:

    >>> # Make text file
    >>> testfile = make_text_file()
    >>> # Open the file
    >>> open_file(testfile)
    >>> #Find 'Notepad' in window titles
    >>> find_window_title('Notepad')
    'generated_text_file.txt - Notepad'

    Keywords
        windows, user, password, remote desktop, remote, citrix, vnc, remotedesktop

    Icon
        lab la-readme
    """

    import ctypes

    EnumWindows = ctypes.windll.user32.EnumWindows
    EnumWindowsProc = ctypes.WINFUNCTYPE(
        ctypes.c_bool, ctypes.POINTER(ctypes.c_int), ctypes.POINTER(ctypes.c_int),
    )
    GetWindowText = ctypes.windll.user32.GetWindowTextW
    GetWindowTextLength = ctypes.windll.user32.GetWindowTextLengthW
    IsWindowVisible = ctypes.windll.user32.IsWindowVisible

    titles = []

    def foreach_window(hwnd, lParam):
        if IsWindowVisible(hwnd):
            length = GetWindowTextLength(hwnd)
            buff = ctypes.create_unicode_buffer(length + 1)
            GetWindowText(hwnd, buff, length + 1)
            titles.append(buff.value)
        return True

    EnumWindows(EnumWindowsProc(foreach_window), 0)

    if partial:
        for title in titles:
            if searchterm in title:
                return title

    if not partial:
        for title in titles:
            if searchterm == title:
                return title

    else:
        return False


@activity
def start_remote_desktop(
    ip, username, password=None, desktop_width=1920, desktop_height=1080
):
    """Login to Windows Remote Desktop

    Create a RDP and login to Windows Remote Desktop

    :parameter ip: IP address of remote desktop
    :type ip: string
    :parameter username: Username
    :type username: string
    :parameter password: Password
    :type password: string, optional
    :parameter desktop_width: Resolution (width) of desktop, standard value is 1920 (full HD)
    :type desktop_width: int, optional
    :parameter desktop_height: Resolution (height) of desktop, standard value is 1080 (full HD)
    :type desktop_height: int, optional

        :Example:

    >>> start_remote_desktop('123.456.789.10','Administrator', 'SamplePassword')

    Keywords
        windows, user, password, remote desktop, remote, citrix, vnc, remotedesktop

    Icon
        las la-passport

    """
    only_supported_for("Windows")

    rdp_raw = """screen mode id:i:1
use multimon:i:0
session bpp:i:32
compression:i:1
keyboardhook:i:2
audiocapturemode:i:0
videoplaybackmode:i:1
connection type:i:7
networkautodetect:i:1
bandwidthautodetect:i:1
displayconnectionbar:i:1
enableworkspacereconnect:i:0
disable wallpaper:i:0
allow font smoothing:i:0
allow desktop composition:i:0
disable full window drag:i:1
disable menu anims:i:1
disable themes:i:0
disable cursor setting:i:0
bitmapcachepersistenable:i:1
audiomode:i:0
redirectprinters:i:1
redirectcomports:i:0
redirectsmartcards:i:1
redirectclipboard:i:1
redirectposdevices:i:0
autoreconnection enabled:i:1
authentication level:i:2
prompt for credentials:i:0
negotiate security layer:i:1
remoteapplicationmode:i:0
alternate shell:s:
shell working directory:s:
gatewayhostname:s:
gatewayusagemethod:i:4
gatewaycredentialssource:i:4
gatewayprofileusagemethod:i:0
promptcredentialonce:i:0
gatewaybrokeringtype:i:0
use redirection server name:i:0
rdgiskdcproxy:i:0
kdcproxyname:s:"""
    rdp_raw = rdp_raw + "\n" + "username:s:" + username
    rdp_raw = rdp_raw + "\n" + "full address:s:" + ip
    rdp_raw = rdp_raw + "\n" + "desktopwidth:i:" + str(desktop_width)
    rdp_raw = rdp_raw + "\n" + "desktopheight:i:" + str(desktop_height)

    import os

    output_path = os.path.join(os.path.expanduser("~"), "remote_desktop.rdp")

    with open(output_path, "w", encoding="utf-8") as file:
        file.write(rdp_raw)

    import subprocess  # nosec

    subprocess.Popen(["cmd.exe", "/c", output_path])

    return output_path


@activity
def close_remote_desktop():
    """Stop Windows Remote Desktop

    Stop Windows Remote Desktop

        :Example:

    >>> close_remote_desktop()

    Keywords
        windows, user, password, remote desktop, remote, citrix, vnc, remotedesktop, stop

    Icon
        las la-passport

    """
    only_supported_for("Windows")
    import os

    os.system("taskkill /f /im mstsc.exe >nul 2>&1")


@activity
def set_user_password(username, password):
    """Set Windows password

    Sets the password for a Windows user.

    :parameter username: Username
    :type username: string
    :parameter password: New password
    :type password: string

        :Example:

    >>> set_user_password('SampleUsername', 'SamplePassword')

    Keywords
        windows, user, password, account

    Icon
        las la-passport

    """
    only_supported_for("Windows")

    from win32com import adsi

    user = adsi.ADsGetObject("WinNT://localhost/%s,user" % username)
    user.SetPassword(password)


@activity
def validate_user_password(username, password):
    """Check Windows password

    Validates a Windows user password if it is correct

    :parameter username: Username
    :type username: string
    :parameter password: New password
    :type password: string

    :return: True if the password is correct

        :Example:

    >>> validate_user_password('SampleUsername', 'SamplePassword')
    False

    Keywords
        windows, user, password, account

    Icon
        las la-passport

    """
    only_supported_for("Windows")

    from win32security import LogonUser
    from win32con import LOGON32_LOGON_INTERACTIVE, LOGON32_PROVIDER_DEFAULT

    try:
        LogonUser(
            username,
            None,
            password,
            LOGON32_LOGON_INTERACTIVE,
            LOGON32_PROVIDER_DEFAULT,
        )
    except:
        return False
    return True


@activity
def lock_windows():
    """Lock Windows

    Locks Windows requiring login to continue.

        :Example:

    >>> lock_windows()

    Keywords
        windows, user, password, account, lock, freeze, hibernate, sleep, lockescreen

    Icon
        las la-user-lock

    """
    only_supported_for("Windows")

    import ctypes

    ctypes.windll.user32.LockWorkStation()


@activity
def is_logged_in():
    """Check if Windows logged in

    Checks if the current user is logged in and not on the lockscreen. Most automations do not work properly when the desktop is locked.

    :return: True if the user is logged in, False if not

        :Example:

    >>> is_logged_in()
    True

    Keywords
        windows, login, logged in, lockscreen, user, password, account, lock, freeze, hibernate, sleep

    Icon
        lar la-user
    """
    only_supported_for("Windows")

    import subprocess  # nosec

    output = subprocess.check_output("TASKLIST")

    if "LogonUI.exe" in str(output):
        return False
    else:
        return True


@activity
def is_desktop_locked():
    """Check if Windows is locked

    Checks if the current user is locked out and on the lockscreen. Most automations do not work properly when the desktop is locked.

    :return: True when the lockscreen is active, False if not.

        :Example:

    >>> desktop_locked()
    True

    Keywords
        windows, login, logged in, lockscreen, user, password, account, lock, locked, freeze, hibernate, sleep

    Icon
        las la-user

    """
    only_supported_for("Windows")

    return not is_logged_in()


@activity
def get_username():
    """Get Windows username

    Get current logged in user's username

        :Example:

    >>> get_username()
    'Automagica'

    Keywords
        windows, login, logged in, lockscreen, user, password, account, lock, locked, freeze, hibernate, sleep

    Icon
        las la-user

    """
    only_supported_for("Windows")

    import getpass

    return getpass.getuser()


@activity
def set_to_clipboard(text):
    """Set clipboard

    Set any text to the Windows clipboard. 

    :parameter text: Text to put in the clipboard
    :type text: string

        :Example:

    >>> # Create some sample text
    >>> sample_text = 'A robots favourite food must be computer chips'
    >>> # Set to clipboard
    >>> set_to_clipboard(sample_text)
    >>> # Print the clipboard to verify
    >>> print( get_from_clipboard() )

    Keywords
        copy, clipboard, clip board, ctrl c, ctrl v, paste

    Icon
        las la-clipboard-check
    """
    only_supported_for("Windows")

    import win32clipboard

    win32clipboard.OpenClipboard()
    win32clipboard.EmptyClipboard()
    win32clipboard.SetClipboardText(text, win32clipboard.CF_UNICODETEXT)
    win32clipboard.CloseClipboard()


@activity
def get_from_clipboard():
    """Get clipboard

    Get the text currently in the Windows clipboard

    :return: Text currently in the clipboard

        :Example:

    >>> # Create some sample text
    >>> sample_text = 'A robots favourite food must be computer chips'
    >>> # Set to clipboard
    >>> set_to_clipboard(sample_text)
    >>> # Get the clipboard to verify
    >>> get_from_clipboard()
    'A robots favourite food must be computer chips'

    Keywords
        copy, clipboard, clip board, ctrl c, ctrl v, paste

    Icon
        las la-clipboard-list

    """
    only_supported_for("Windows")

    import win32clipboard

    win32clipboard.OpenClipboard()
    try:
        data = str(win32clipboard.GetClipboardData(win32clipboard.CF_UNICODETEXT))
        return data

    except:
        return None

    finally:
        win32clipboard.CloseClipboard()


@activity
def clear_clipboard():
    """Empty clipboard

    Empty text from clipboard. Getting clipboard data after this should return in None

        :Example:

    >>> # Create some sample text
    >>> sample_text = 'A robots favourite food must be computer chips'
    >>> # Set to clipboard
    >>> set_to_clipboard(sample_text)
    >>> # Clear the clipboard
    >>> clear_clipboard()
    >>> # Get clipboard contents to verify
    >>> print( get_clipboard() )
    None

    Keywords
        copy, clipboard, clip board, ctrl c, ctrl v, paste

    Icon
        las la-clipboard
    """
    only_supported_for("Windows")

    from ctypes import windll

    if windll.user32.OpenClipboard(None):
        windll.user32.EmptyClipboard()
        windll.user32.CloseClipboard()
    return


@activity
def run_vbs_script(script_path, parameters=[]):
    """Run VBSscript

    Run a VBScript file

    :parameter script_path: Path to the .vbs-file
    :type script_path: input_file
    :parameter parameters: Additional arguments to pass to the VBScript

        :Example:

    >>> # Run a VBS script
    >>> run_vbs_script('Samplescript.vbs')

    Keywords
        vbs, VBScript

    Icon
        las la-cogs
    """
    only_supported_for("Windows")

    import subprocess  # nosec

    subprocess.call(["cscript.exe", script_path] + parameters)


@activity
def beep(frequency=1000, duration=500):
    """Beep

    Make a beeping sound. Make sure your volume is up and you have hardware connected.

    :parameter frequency: Integer to specify frequency (Hz), default value is 1000 Hz
    :type frequency: int, optional
    :parameter duration: Integer to specify duration of beep in miliseconds (ms), default value is 500 ms.
    :type duration: int, optional

    :return: Sound

        :Example:

    >>> beep()

    Keywords
        beep, sound, noise, speaker, alert

    Icon
        las la-volume-up

    """
    only_supported_for("Windows")

    import winsound

    winsound.Beep(frequency, duration)


@activity
def get_all_network_interface_names():
    """Get all network interface names

    Returns a list of all network interfaces of the current machine

        :Example:

    >>> get_all_network_interface_names()
    ['Microsoft Kernel Debug Network Adapter', 'Realtek Gaming GbE Family Controller', 'WAN Miniport (SSTP)']

    Keywords
        networking, connection, list

    Icon
        las la-ethernet

    """
    only_supported_for("Windows")

    import subprocess  # nosec

    rows = subprocess.check_output("wmic nic get name")

    results = [row.strip() for row in rows.decode("utf-8").split("\n")[1:]]

    return results


@activity
def enable_network_interface(name):
    """Enable network interface

    Enables a network interface by its name.

    :parameter name: Name of the network
    :type name: string

        :Example:

    >>> enable_network_interface('Realtek Gaming GbE Family Controller')

    Keywords
        networking, connection, enable

    Icon
        las la-ethernet
    """
    only_supported_for("Windows")
    import subprocess  # nosec

    subprocess.check_output(
        'wmic path win32_networkadapter where name="{}" call enable'.format(name)
    )


@activity
def disable_network_interface(name):
    """Disable network interface

    Disables a network interface by its name.

    :parameter name: Name of the network interface
    :type name: string

        :Example:

    >>> disable_network_interface('Realtek Gaming GbE Family Controller')
    
    Keywords
        networking, connection, disable

    Icon
        las la-ethernet
    """
    only_supported_for("Windows")

    import subprocess  # nosec

    subprocess.check_output(
        'wmic path win32_networkadapter where name="{}" call disable'.format(name)
    )


@activity
def get_default_printer_name():
    """Get default printer

    Returns the name of the printer selected as default

        :Example:

    >>> get_default_printer_name()
    'Epson MF742C/744C'

    Keywords
        printing, get default printer name, default printer

    Icon
        las la-print
    """
    only_supported_for("Windows")

    import win32print

    return win32print.GetDefaultPrinter()


@activity
def set_default_printer(name):
    """Set default printer

    Set the default printer.

    :parameter name: Printer name
    :type name: string

        :Example:

    >>> set_default_printer('Epson MF742C/744C')

    Keywords
        printing, set default printer name, default printer

    Icon
        las la-print
    """
    only_supported_for("Windows")

    import win32print

    return win32print.SetDefaultPrinter(name)


@activity
def remove_printer(name):
    """Remove printer

    Removes a printer by its name

    :parameter name: Printer name to remove
    :type name: string

        :Example:

    >>> remove_printer('Epson MF742C/744C')

    Keywords
        printing, remove printer, printer

    Icon
        las la-print
    """
    only_supported_for("Windows")

    import win32print

    return win32print.DeletePrinter(name)


@activity
def get_service_status(name):
    """Get service status

    Returns the status of a service on the machine

    :parameter name: Name of service
    :type name: string

        :Example:

    >>> get_service_status('Windows Backup')
    'stopped'

    Keywords
        services, get service status, status

    Icon
        las la-cog
    """
    only_supported_for("Windows")

    import psutil

    for s in psutil.win_service_iter():
        if s.name() == name or s.display_name() == name:
            return s.status()


@activity
def start_service(name):
    """Start a service

    Starts a Windows service

    :parameter name: Name of service
    :type name: string

        :Example:

    >>> start_service('Windows Backup')

    Keywords
        services, start a service, start

    Icon
        las la-cog
    """
    only_supported_for("Windows")

    import win32serviceutil

    win32serviceutil.StartService(name)


@activity
def stop_service(name):
    """Stop a service

    Stops a Windows service

    :parameter name: Name of service
    :type name: string

        :Example:

    >>> stop_service('Windows Backup')

    Keywords
        services, stop a service, stop

    Icon
        las la-cog
    """
    only_supported_for("Windows")

    import win32serviceutil

    win32serviceutil.StopService(name)


@activity
def set_window_to_foreground(title):
    """Set window to foreground

    Sets a window to foreground by its title.

    :parameter name: Name of service
    :type name: string

        :Example:

    >>> set_window_to_foreground('Notepad - Untitled')

    Keywords
        window, foreground

    Icon
        las la-window-restore
    """
    only_supported_for("Windows")

    import win32gui

    handle = win32gui.FindWindow(None, title)

    if not handle:
        raise Exception('Could not find a window with title "{}"'.format(title))

    win32gui.SetForegroundWindow(handle)


@activity
def get_foreground_window_title():
    """Get foreground window title
    
    Retrieve the title of the current foreground window

        :Example:

    >>> get_foreground_window_title()
    'IPython'

    Keywords
        window, foreground, title

    Icon
        las la-window-restore
    """
    only_supported_for("Windows")

    import win32gui

    handle = win32gui.GetForegroundWindow()

    return win32gui.GetWindowText(handle)


@activity
def close_window(title):
    """Close window
    
    Closes a window by its title

    :parameter title: Title of window
    :type title: string

        :Example:

    >>> close_window('Untitled - Notepad')

    Keywords
        window, close, title

    Icon
        las la-window-restore
    """
    only_supported_for("Windows")

    import win32gui

    handle = win32gui.FindWindow(None, title)

    if not handle:
        raise Exception('Could not find a window with title "{}"'.format(title))

    win32gui.DestroyWindow(handle)


@activity
def maximize_window(title):
    """Maximize window
    
    Maximizes a window by its title

    :parameter title: Title of window
    :type title: string

        :Example:

    >>> maximize_window('Untitled - Notepad')

    Keywords
        window, maximize, title

    Icon
        las la-window-restore
    """
    only_supported_for("Windows")

    import win32con
    import win32gui

    handle = win32gui.FindWindow(None, title)

    if not handle:
        raise Exception('Could not find a window with title "{}"'.format(title))

    win32gui.ShowWindow(handle, win32con.SW_SHOWMAXIMIZED)
    win32gui.SetForegroundWindow(handle)


@activity
def restore_window(title):
    """Restore window
    
    Restore a window by its title

    :parameter title: Title of window
    :type title: string

        :Example:

    >>> restore_window('Untitled - Notepad')

    Keywords
        window, restore, title

    Icon
        las la-window-restore
    """
    only_supported_for("Windows")

    import win32con
    import win32gui

    handle = win32gui.FindWindow(None, title)

    if not handle:
        raise Exception('Could not find a window with title "{}"'.format(title))

    win32gui.ShowWindow(handle, win32con.SW_RESTORE)
    win32gui.SetForegroundWindow(handle)


@activity
def minimize_window(title):
    """Minimize window
    
    Minimizes a window by its title

    :parameter title: Title of window
    :type title: string

        :Example:

    >>> minimize_window(title)

    Keywords
        window, minimize, title

    Icon
        las la-window-restore
    """
    only_supported_for("Windows")

    import win32gui

    handle = win32gui.FindWindow(None, title)

    if not handle:
        raise Exception('Could not find a window with title "{}"'.format(title))

    win32gui.CloseWindow(handle)


@activity
def resize_window(title, x, y, width, height):
    """Resize window
    
    Resize a window by its title

    :parameter title: Title of window
    :type title: string
    :parameter x: Starting x position
    :type x: int
    :parameter y: Starting y position
    :type y: int
    :parameter width: Width
    :type width: int
    :parameter height: Height
    :type height: int


        :Example:

    >>> resize_window('Untitled - Notepad', 100, 200, 300, 400)

    Keywords
        window, resize, title

    Icon
        las la-window-restore
    """
    only_supported_for("Windows")

    import win32gui

    handle = win32gui.FindWindow(None, title)

    if not handle:
        raise Exception('Could not find a window with title "{}"'.format(title))

    win32gui.MoveWindow(handle, x, y, width, height, True)
    win32gui.SetForegroundWindow(handle)


@activity
def hide_window(title):
    """Hide window
    
    Hides a window from the user desktop by using it's title

    :parameter title: Title of window
    :type title: string

        :Example:

    >>> hide_window('Untitled - Notepad')

    Keywords
        window, hide, title

    Icon
        las la-window-restore
    """
    only_supported_for("Windows")

    import win32con
    import win32gui

    handle = win32gui.FindWindow(None, title)

    if not handle:
        raise Exception('Could not find a window with title "{}"'.format(title))

    win32gui.ShowWindow(handle, win32con.SW_HIDE)


"""
Terminal
Icon: las la-terminal
"""


@activity
def run_ssh_command(user, host, command):
    """Run SSH command

    Runs a command over SSH (Secure Shell)

    :parameter user: User
    :type user: string
    :parameter host: Host
    :type host: string
    :parameter command: Command
    :type command: string

        :Example:

    >>> run_ssh_command('root', 'machine', 'ls -a')
    '. .. .bashrc'

    Keywords
        ssh, command

    Icon
        las la-terminal
    """
    import subprocess  # nosec

    return subprocess.Popen(
        "ssh {user}@{host} {cmd}".format(user=user, host=host, cmd=command),
        shell=True,  # nosec
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
    ).communicate()


"""
SNMP
Icon: las la-ethernet
"""


@activity
def snmp_get(target, oids, credentials, port=161, engine=None, context=None):
    """SNMP Get
    
    Retrieves data from an SNMP agent using SNMP (Simple Network Management Protocol)

    :parameter target: Target
    :type target: string
    :parameter oids: oids
    :type oids: string
    :parameter credentials: credentials
    :type credentials: string
    :parameter port: Port (default 161)
    :type port: int, optional
    :parameter engine: Engine (default none)
    :type engine: string, optional
    :parameter context: Contect (default none)
    :type context: string, optional

        :Example:

    >>> snmp_get()

    Keywords
        snmp, simple network management protocol, protocols, get

    Icon
        las la-ethernet

    """
    # Adaptation of Alessandro Maggio's implementation in QuickSNMP (MIT License)
    # (Copyright (c) 2018 Alessandro Maggio) (https://github.com/alessandromaggio/quicksnmp)

    from pysnmp import hlapi

    if not engine:
        engine = hlapi.SnmpEngine()

    if not context:
        context = hlapi.ContextData()

    def construct_object_types(list_of_oids):
        object_types = []
        for oid in list_of_oids:
            object_types.append(hlapi.ObjectType(hlapi.ObjectIdentity(oid)))
        return object_types

    def construct_value_pairs(list_of_pairs):
        pairs = []
        for key, value in list_of_pairs.items():
            pairs.append(hlapi.ObjectType(hlapi.ObjectIdentity(key), value))
        return pairs

    handler = hlapi.getCmd(
        engine,
        credentials,
        hlapi.UdpTransportTarget((target, port)),
        context,
        *construct_object_types(oids),
    )

    return fetch(handler, 1)[0]


"""
Active Directory
Icon: las la-user
"""


class ActiveDirectory:
    @activity
    def __init__(self, ldap_server=None, username=None, password=None):
        """AD interface
        
        Interface to Windows Active Directory through ADSI. Connects to the AD domain to which the machine is joined by default.

        :parameter ldap_server: LDAP server
        :type ldap_server: string, optional
        :parameter username: Username
        :type username: string, optional
        :parameter password: Password
        :type password: string, optional

            :Example:

        >>> ad = ActiveDirectory()

        Keywords
            AD, active directory, activedirectory

        Icon
            las la-audio-description

        """
        import pyad

        self.pyad = pyad

        if ldap_server:
            self.pyad.set_defaults(ldap_server=ldap_server)

        if username:
            self.pyad.set_defaults(username=username)

        if password:
            self.pyad.set_defaults(password=password)

    @activity
    def get_object_by_distinguished_name(self, distinguished_name):
        """Get AD object by name
        
        Interface to Windows Active Directory through ADSI

        :parameter distinguished_name: Name
        :type distinguished_name: string

            :Example:

        >>> ad = ActiveDirectory()
        >>> ad.get_object_by_distinguished_name('SampleDN')

        Keywords
            AD, active directory, activedirectory

        Icon
            las la-audio-description

        """
        return self.pyad.from_dn(distinguished_name)


"""
Utilities
Icon: las la-toolbox
"""


@activity
def home_path(filename=None):
    """Get user home path

    Returns the current user's home path

    :parameter subdir: Optional filename to add to the path. Can also be a subdirectory
    :type filename: string, optional

    :return: Path to the current user's home folder

        :Example:

    >>> # Home_path without arguments will return the home path
    >>> print( home_path() )
    >>> # When looking for a file in the home path, we can specify it
    >>> # First make a sample text file
    >>> make_text_file()
    >>> # Refer to it
    >>> home_path('generated_text_file.txt')
    'C:\\Users\\<username>\\generated_text_file.txt'

    Keywords
        home, home path, homepath, home directory, homedir

    Icon
        las la-home

    """
    import os

    if filename:
        return os.path.join(os.path.expanduser("~"), filename)
    return os.path.expanduser("~")


@activity
def desktop_path(filename=None):
    """Get desktop path

    Returns the current user's desktop path

    :parameter filename: Optional filename to add to the path. Can also be a subdirectory
    :type filename: string, optional

    :return: Path to the current user's desktop folder

        :Example:

    >>> # Desktop_path without arguments will return the home path
    >>> print( desktop_path() )
    >>> # When looking for a file on the desktop, we can specify it
    >>> # First make a sample text file
    >>> make_text_file()
    >>> # Refer to it
    >>> desktop_path('generated_text_file.txt')
    'C:\\Users\\<username>\\Desktop\\generated_text_file.txt'

    Keywords
        desktop, desktop path, desktoppath, desktop directory, desktopdir

    Icon
        lar la-desktop
    """
    import os

    if filename:
        return os.path.join(os.path.join(os.path.expanduser("~"), "Desktop"), filename)
    return os.path.join(os.path.expanduser("~"), "Desktop")


@activity
def downloads_path():
    """Get downloads path

    Returns the current user's default download path

    :return: Path to the current user's downloads folder

        :Example:

    >>> # Find downloads path
    >>> downloads_path()

    Keywords
        download, download path, downloadpath, download directory, download dir, downloaddir

    Icon
        lar la-download
    """
    import os

    if os.name == "nt":
        import winreg

        sub_key = r"SOFTWARE\Microsoft\Windows\CurrentVersion\Explorer\Shell Folders"
        downloads_guid = "{374DE290-123F-4565-9164-39C4925E467B}"
        with winreg.OpenKey(winreg.HKEY_CURRENT_USER, sub_key) as key:
            location = winreg.QueryValueEx(key, downloads_guid)[0]
        return location
    else:
        return os.path.join(os.path.expanduser("~"), "downloads")


@activity
def open_file(input_path):
    """Open file

    Opens file with default programs

    :parameter input_path: Path to file. 
    :type input_path: input_file

    :return: Path to file

        :Example:

    >>> # Make text file
    >>> testfile = make_text_file()
    >>> # Open the file
    >>> open_file(testfile)

    Keywords
        file, open, open file, show, reveal, explorer, run, start

    Icon
        lar la-file

    """
    import sys

    path = interpret_path(input_path)

    if sys.platform == "win32":
        import os

        os.startfile(path)
    else:
        from subprocess import call

        opener = "open" if sys.platform == "darwin" else "xdg-open"
        call([opener, path])

    return path


@activity
def set_wallpaper(image_path):
    """Set wallpaper

    Set Windows desktop wallpaper with the the specified image

    :parameter image_path: Path to the image. This image will be set as desktop wallpaper
    :type image_path: input_file

        :Example:

    >>> # Caution: this example will change your wallpaper
    >>> # Take a screenshot of current screen
    >>> screenshot = take_screenshot()
    >>> # Flip it hozirontally for fun
    >>> mirror_image_horizontally(screenshot)
    >>> # Set flipped image as wallpaper
    >>> set_wallpaper(screenshot)

    Keywords
        desktop, desktop path, desktoppath, desktop directory, desktopdir, wallpaper, wall paper, wall

    Icon
        las la-desktop

    """
    only_supported_for("Windows")

    import ctypes

    image_path = interpret_path(image_path)
    ctypes.windll.user32.SystemParametersInfoW(20, 0, image_path, 0)


@activity
def download_file_from_url(url, output_path=None):
    """Download file from a URL

    Download file from a URL

    :parameter url: Source URL to download file from
    :type url: string
    :parameter output_path: Target path, default to homedir with name '_download_' + random addition
    :type output_path: output_dir, optional

    :return: Target path as string

        :Example:

    >>> # Download robot picture from the wikipedia robot page
    >>> picture_url = 'https://upload.wikimedia.org/wikipedia/commons/thumb/6/6c/Atlas_from_boston_dynamics.jpg/220px-Atlas_from_boston_dynamics.jpg'
    >>> download_file_from_url(url = picture_url, output_path ='robot.jpg')
    'C:\\Users\\<username>\\robot.jpg'

    Keywords
        download, download url, save, request

    Icon
        las la-cloud-download-alt
    """
    from automagica.httpclient import http_client
    import re
    import os
    from urllib.parse import urlparse

    a = urlparse(url)
    filename = os.path.basename(a.path)

    output_path = interpret_path(output_path, default_filename=filename)

    r = http_client.get(url, stream=True)

    if r.status_code == 200:
        with open(output_path, "wb") as f:
            f.write(r.content)

        return output_path

    else:
        raise Exception("Could not download file from {}".format(url))


"""
System
Icon: las la-laptop
"""


@activity
def rename_file(input_path, output_name=None):
    """Rename a file

    This activity will rename a file. If the the desired name already exists in the folder file will not be renamed. Make sure to add the exstention to specify filetype.

    :parameter input_path: Full path to file that will be renamed
    :type input_path: input_file
    :parameter output_name: New name of the file e.g. 'newfile.txt'. By default file will be renamed to original folder name with '_renamed' added to the folder name.
    :type output_name: string, optional

    :return: Path to renamed file as a string. None if folder could not be renamed.

        :Example:

    >>> # Make new text file in home directory
    >>> text_file = make_text_file()
    >>> # Rename the file
    >>> rename_file(text_file, output_name='brand_new_filename.txt')
    C:\\Users\\<username>\\brand_new_filename.txt'

    Keywords
        file, rename, rename file, organise file, files, file manipulation, explorer, nautilus

    Icon
        las la-file-contract
    """
    import os

    input_path = interpret_path(input_path, required=True)
    if not output_name:
        output_path = interpret_path(input_path, addition="_renamed")
    else:
        output_path = interpret_path(input_path, replace_filename=output_name)

    os.rename(input_path, output_path)

    return output_path


@activity
def move_file(input_path, output_path=None):
    """Move a file

    If the new location already contains a file with the same name.

    :parameter input_path: Full path to the file that will be moved
    :type input_path: input_file
    :parameter output_path: Path to the folder where file will be moved to, defaults to input_path with '_moved' added
    :type output_path: output_dir

    :return: Path to renamed file as a string. None if folder could not be moved.

        :Example:

    >>> # Make new text file in home directory
    >>> text_file = make_text_file()
    >>> # Make a folder to move the file to
    >>> new_folder = create_folder()
    >>> # Move text file to the folder
    >>> move_file(text_file, new_folder)

    Keywords
        file, move, move file, organise file, files, file manipulation, explorer, nautilus

    Icon
        las la-file-export

    """

    import shutil

    input_path = interpret_path(input_path, required=True)
    if not output_path:
        output_path = interpret_path(input_path, addition="_moved")
    else:
        output_path = interpret_path(output_path)

    shutil.move(input_path, output_path)

    return output_path


@activity
def remove_file(path):
    """Remove a file

    Remove a file 

    :parameter path: Full path to the file that will be deleted.
    :type path: input_file

    :return: Path to removed file as a string.

        :Example:

    >>> # Make new text file in home directory
    >>> text_file = make_text_file()
    >>> # Remove the file
    >>> remove_file(text_file)

    Keywords
        file, delete, erase, delete file, organise file, files, file manipulation, explorer, nautilus

    Icon
        las la-trash
    """

    import os

    path = interpret_path(path)
    if os.path.isfile(path):
        os.remove(path)
    return path


@activity
def file_exists(path):
    """Check if file exists

    This function checks whether the file with the given path exists.

    :parameter path: Full path to the file to check.
    :type path: input_file

    return: True or False (boolean)

        :Example:

    >>> # Make new text file in home directory
    >>> text_file = make_text_file()
    >>> # Check if file exists
    >>> file_exists(text_file)
    True

    Keywords
        file, exists, files, file manipulation, explorer, nautilus

    Icon
        las la-tasks
    """
    import os

    path = interpret_path(path)
    return os.path.isfile(path)


@activity
def wait_file_exists(path, timeout=60):
    """Wait until a file exists.

    Not that this activity is blocking and will keep the system waiting.

    :parameter path: Full path to file.
    :type path: input_file
    :parameter timeout: Maximum time in seconds to wait before continuing. Default value is 60 seconds.
    :type timeout: int, optional

        :Example:

    >>> # Make new text file in home directory
    >>> text_file = make_text_file()
    >>> # Wait untile file exists # Should pass immediatly
    >>> wait_file_exists(text_file)

    Keywords
        file, wait, wait till exists, files, file manipulation, explorer, nautilus

    Icon
        las la-list-alt

    """
    from time import sleep
    import os

    path = interpret_path(path)
    for _ in range(timeout):
        if os.path.exists(path):
            break
        sleep(1)


@activity
def write_list_to_file(list_to_write, file_path):
    """List to .txt
    
    Writes a list to a  text (.txt) file. 
    Every element of the entered list is written on a new line of the text file.

    :parameter list_to_write: List to write to .txt file
    :parameter file_path: Path to the text-file. 
    :type file_path: output_file
    :extension file_path: txt

        :Example:
    
    >>> # Make a list to write
    >>> robot_names = ['WALL-E', 'Terminator', 'R2D2']
    >>> # Create a new text file
    >>> text_file = make_text_file()
    >>> write_list_to_file(robot_names, text_file)
    >>> # Open the file for illustration
    >>> open_file(text_file)

    Keywords
        list, text, txt, list to file, write list, write

    Icon
        las la-list

    """
    file_path = interpret_path(file_path)
    with open(file_path, "w") as filehandle:
        filehandle.writelines("%s\n" % place for place in list_to_write)
    return


@activity
def read_list_from_txt(input_path):
    """Read list from .txt file

    This activity reads the content of a .txt file to a list and returns that list. 
    Every new line from the .txt file becomes a new element of the list. The activity will 
    not work if the entered path is not attached to a .txt file.

    :parameter input_path: Path to the .txt file
    :type input_path: input_file
    :extension input_path: txt

    :return: List with contents of specified .txt file

        :Example:
    
    >>> # Make a list to write
    >>> robot_names = ['WALL-E', 'Terminator', 'R2D2']
    >>> # Create a new text file
    >>> text_file = make_text_file()
    >>> write_list_to_file(robot_names, text_file)
    >>> # Read list from file
    >>> read_list_from_txt(text_file)
    ['WALL-E', 'Terminator', 'R2D2']

    Keywords
        list, text, txt, list to file, write list, read, read txt, read text

    Icon
        las la-th-list

    """
    file_path = interpret_path(input_path)
    written_list = []
    with open(file_path, "r") as filehandle:
        filecontents = filehandle.readlines()
        for line in filecontents:
            current_place = line[:-1]
            written_list.append(current_place)
    return written_list


@activity
def read_from_txt(file_path):
    """Read .txt file

    This activity reads a .txt file and returns the content

    :parameter input_path: Path to the .txt file
    :type input_path: input_file
    :extension input_path: txt

    :return: Contents of specified .txt file

        :Example:
    
    >>> # Create a new text file
    >>> text_file = make_text_file()
    >>> # Read list from file
    >>> read_from_txt(text_file)
    'Sample text'

    Keywords
        list, text, txt, list to file, read, read txt, read text

    Icon
        las la-th-list

    """
    file_path = interpret_path(file_path)
    written_list = []
    with open(file_path, "r") as filehandle:
        filecontents = filehandle.readlines()

    return filecontents


@activity
def append_line(text, file_path):
    """Append to .txt
    
    Append a text line to a file and creates the file if it does not exist yet.

    :parameter text: The text line to write to the end of the file
    :type text: string
    :parameter file_path: Path to the file to write to
    :type file_path: input_file
    :extension file_path: txt

        :Example:
    
    >>> # Create a new text file
    >>> text_file = make_text_file()
    >>> # Append a few lines to the file
    >>> append_line('Line 1', text_file)
    >>> append_line('Line 2', text_file)
    >>> append_line('Line 3', text_file)
    >>> # Open the file for illustration
    >>> open_file(text_file)

    Keywords
        list, text, txt, list to file, write list, read, write txt, append text, append line, append, add to file, add

    Icon
        las la-tasks
    """

    import os

    file_path = interpret_path(file_path)
    if not os.path.isfile(file_path):
        with open(file_path, "a"):
            os.utime(file_path, None)

    with open(file_path, "a") as f:
        f.write("\n" + text)


@activity
def make_text_file(text="Sample text", output_path=None):
    """Make text file

    Initialize text file

    :parameter text: The text line to write to the end of the file. Default text is 'Sample text'
    :type text: string, optional
    :parameter output_path: Ouput path. Will write to home directory on default
    :type output_path: output_file, optional
    :extension output_path: txt

    :return: Path as string

        :Example:

    >>> # Create a new text file
    >>> text_file = make_text_file()
    C:\\Users\\<username>\\generated_text_file.txt'

    Keywords
        make text file, text_file, testfile, exampel file, make file, make, new file, new text_file, txt, new txt

    Icon
        las la-file-alt

    """

    # Set to user home if no path specified
    import os

    if not output_path:
        output_path = interpret_path(
            output_path, default_filename="generated_text_file.txt"
        )
    else:
        import pathlib

        if pathlib.Path(output_path).is_dir():
            output_path = interpret_path(
                output_path, default_filename="generated_text_file.txt"
            )
        else:
            output_path = interpret_path(output_path)
    with open(output_path, "w", encoding="utf-8") as file:
        file.write(text)

    return output_path


@activity
def read_text_file_to_list(file_path):
    """Read .txt file with newlines to list

    Read a text file to a Python list-object

    :parameter file_path: Path to the text file which should be read to a list
    :type file_path: input_file
    :extension file_path: txt

    :return: List with the lines in the text file

        :Example:

    >>> # Create a new text file
    >>> text_file = make_text_file(text="First line!\nSecond line!")
    >>> # Read the text file to a list
    >>> lines = read_text_file_to_list(text_file)
    >>> lines
    ['First line!', 'Second line!']

    Keywords
        read text file, list, reading text file

    Icon
        las la-copy

    """
    file_path = interpret_path(file_path)
    with open(file_path, "r") as f:
        result = [line.strip() for line in f.readlines()]

    return result


@activity
def copy_file(input_path, output_path=None):
    """Copy a file

    Copies a file from one place to another.
    If the new location already contains a file with the same name, a random 4 character uid is added to the name.

    :parameter input_path: Full path to the source location of the file
    :type input_path: input_file
    :parameter output_path: Optional full path to the destination location of the folder. If not specified file will be copied to the same location with a random 8 character uid is added to the name.
    :type output_path: output_dir, optional

    :return: New path as string

        :Example:

    >>> # Create a new text file
    >>> text_file = make_text_file()
    >>> # Copy the text file
    >>> copy_file(text_file)
    C:\\Users\\<username>\\generated_text_file.txt'

    Keywords
        make text file, text_file, testfile, example file, make file, make, new file, new text_file, txt, new txt

    Icon
        las la-copy
    
    """

    import shutil

    input_path = interpret_path(input_path, required=True)
    if not output_path:
        output_path = interpret_path(input_path, addition="_copied")
    else:
        output_path = interpret_path(output_path)

    shutil.copy(input_path, output_path)

    return output_path


@activity
def get_file_extension(file_path):
    """Get file extension
    
    Get extension of a file

    :parameter file_path: Path to file to get extension from
    :type file_path: input_file

    :return: String with extension, e.g. '.txt'

        :Example:

    >>> # Create a new text file
    >>> text_file = make_text_file()
    >>> # Get file extension of this text file
    >>> get_file_extension(text_file)
    '.txt'

    Keywords
        file, extension, file extension, details

    Icon
        las la-info

    """

    import os

    file_path = interpret_path(file_path)

    filename, file_extension = os.path.splitext(file_path)

    return file_extension


@activity
def send_to_printer(file_path):
    """Print
    
    Send file to default printer to priner. This activity sends a file to the printer. Make sure to have a default printer set up.

    :parameter file_path: Path to the file to print, should be a printable file
    :type file_path: input_file

        :Example:

    >>> # Caution as this example could result in a print from default printer
    >>> # Create a new text file
    >>> text_file = make_text_file(text = 'What does a robot do at lunch? Take a megabyte!')
    >>> # Print the text file
    >>> send_to_printer(text_file)

    Keywords
        print, printer, printing, ink, export

    Icon
        las la-print
    """
    import os

    file_path = interpret_path(file_path)
    os.startfile(file_path, "print")


"""
PDF
Icon: las la-file-pdf
"""


@activity
def read_text_from_pdf(file_path):
    """Text from PDF
    
    Extracts the text from a PDF. This activity reads text from a pdf file. Can only read PDF files that contain a text layer.

    :parameter file_path: Path to the PDF (either relative or absolute)
    :type file_path: input_file
    :extension file_path: pdf

    :return: The text from the PDF

        :Example:
        
    >>> # Caution, for this example to work a .pdf example file will be downloaded from automagica.com FTP
    >>> example_pdf = download_file_from_url('http://automagica.com/examples/example_document.pdf')
    >>> # Open example pdf for illustration
    >>> open_file(example_pdf)
    >>> # Read the text
    >>> read_text_from_pdf(example_pdf)

    Keywords
        PDF, read, text, extract text, PDF file

    Icon
        las la-glasses
    """
    from PyPDF2 import PdfFileReader

    file_path = interpret_path(file_path)
    text = ""

    with open(file_path, "rb") as f:
        reader = PdfFileReader(f)
        for i in range(reader.numPages):
            page = reader.getPage(i)
            text += page.extractText()

    return text


@activity
def join_pdf_files(
    first_file_path, second_file_path, third_file_path=None, output_path=None
):
    """Merge PDF
    
    Merges multiple PDFs into a single file

    :parameter first_file_path: Path to first PDF file
    :type first_file_path: input_file
    :extension first_file_path: pdf
    :parameter second_file_path: Path to second PDF file
    :type second_file_path: input_file, optional
    :extension second_file_path: pdf
    :parameter third_file_path: Path to third PDF file, optional
    :type third_file_path: input_file, optional
    :extension third_file_path: pdf
    :parameter output_path: Full path where joined pdf files can be written. If no path is given will write to home dir as 'merged_pdf.pdf'
    :type output_path: output_file, optional
    :extension output_path: pdf

    :return: Output path as string
    
        :Example:
        
    >>> # Caution, for this example to work a .pdf example file will be downloaded from automagica.com FTP
    >>> example_pdf = download_file_from_url('http://automagica.com/examples/example_document.pdf')
    >>> # Join the PDF file with itself for illustration, could also be different files
    >>> merged_pdf = join_pdf_files(example_pdf, example_pdf)
    >>> # Open resulting PDF file for illustration
    >>> open_file(merged_pdf)

    Keywords
        PDF, read, text, extract text, PDF file, join PDF, join, merge, merge PDF

    Icon
        las la-object-ungroup

    """
    from PyPDF2 import PdfFileMerger, PdfFileReader

    if not output_path:
        output_path = interpret_path(first_file_path, addition="_merged_pdf.pdf")
    else:
        output_path = interpret_path(output_path)

    first_file_path = interpret_path(first_file_path)
    second_file_path = interpret_path(second_file_path)

    file_paths = []
    file_paths.append(first_file_path)
    file_paths.append(second_file_path)
    if third_file_path:
        file_paths.append(third_file_path)

    merger = PdfFileMerger()
    for file_path in file_paths:
        with open(file_path, "rb") as f:
            merger.append(PdfFileReader(f))

    merger.write(output_path)

    return output_path


@activity
def extract_page_range_from_pdf(file_path, start_page, end_page, output_path=None):
    """Extract page from PDF
    
    Extracts a particular range of a PDF to a separate file.

    :parameter file_path: Path to the PDF (either relative or absolute)
    :type file_path: input_file
    :extension file_path: pdf
    :parameter start_page: Page number to start from, with 0 being the first page
    :type start_page: int
    :parameter end_page: Page number to end with, with 0 being the first page
    :type end_page: int
    :param output_path: Output path, if no path is provided same path as input will be used with 'extracted' added to the name
    :type output_path: output_file, optional
    :extension output_path: pdf


        :Example:

    >>> # Caution, for this example to work a .pdf example file will be downloaded from automagica.com FTP
    >>> example_pdf = download_file_from_url('http://automagica.com/examples/example_document.pdf')
    >>> # Join the PDF file three times to create multi page
    >>> multi_page_pdf_example = join_pdf_files(example_pdf, example_pdf, example_pdf)
    >>> # Extract some pages from it
    >>> new_file = extract_page_range_from_pdf(multi_page_pdf_example, 1, 2 )
    >>> # Open resulting PDF file for illustration
    >>> open_file(new_file)

    Keywords
        PDF, read, extract text, PDF file, extract PDF, join, cut, cut PDF, extract pages, extract from pdf, select page, page
    Icon
        las la-cut
    
    """
    from PyPDF2 import PdfFileWriter, PdfFileReader

    if not output_path:
        output_path = interpret_path(file_path, replace_filename="extracted_paged.pdf")
    else:
        output_path = interpret_path(output_path)

    file_path = interpret_path(file_path)
    with open(file_path, "rb") as f:

        reader = PdfFileReader(f)
        writer = PdfFileWriter()

        for i in range(start_page, end_page):
            writer.addPage(reader.getPage(i))

        with open(output_path, "wb") as f:
            writer.write(f)

    return output_path


@activity
def extract_images_from_pdf(file_path):
    """Extract images from PDF
    
    Save a specific page from a PDF as an image

    :parameter file_path: Full path to store extracted images
    :type file_path: output_dir

        :Example:
        
    >>> # Caution, for this example to work a .pdf example file will be downloaded from automagica.com FTP
    >>> example_pdf = download_file_from_url('http://automagica.com/examples/example_document.pdf')
    >>> # Extract the images
    >>> extract_images_from_pdf(example_pdf)

    Keywords
        PDF, extract images, images, extract text, PDF file, image

    Icon
        las la-icons
        
    """

    # Snippet from silvain https://stackoverflow.com/a/34116472/10949633
    # Snippet available on his GitHub under the BSD license https://github.com/sylvainpelissier/PyPDF2

    from PyPDF2 import PdfFileReader
    from PIL import Image

    file_path = interpret_path(file_path)

    extracted_images = []

    try:
        with open(file_path, "rb") as f:

            r = PdfFileReader(f)

            for i in range(r.getNumPages()):
                page = r.getPage(i)
                items = page["/Resources"]["/XObject"].getObject()

                for item in items:
                    if items[item]["/Subtype"] == "/Image":
                        size = (
                            items[item]["/Width"],
                            items[item]["/Height"],
                        )
                        data = items[item].getData()

                        if items[item]["/ColorSpace"] == "/DeviceRGB":
                            type = "RGB"
                        else:
                            type = "P"

                        if items[item]["/Filter"] == "/FlateDecode":
                            img = Image.frombytes(type, size, data)
                            img.save(item[1:] + ".png")
                            extracted_images.append(item[1:] + ".png")

                        elif items[item]["/Filter"] == "/JPXDecode":
                            img = open(item[1:] + ".jp2", "wb")
                            extracted_images.append(item[1:] + ".jp2")
                            img.write(data)
                            img.close()
    except:
        return None

    return extracted_images


@activity
def apply_watermark_to_pdf(file_path, watermark_path, output_path=""):
    """Watermark a PDF

    Watermark a PDF 

    :parameter file_path: Filepath to the document that will be watermarked. Should be pdf file.
    :type file_path: input_file
    :extension file_path: pdf
    :parameter watermark_path: Filepath to the watermark. Should be pdf file.
    :type watermark_path: input_file
    :extension watermark_path: pdf
    :parameter output_path: Path to save watermarked PDF. If no path is provided same path as input will be used with 'watermarked' added to the name
    :type output_path: output_file, optional
    :extension output_path: pdf

    :return: Output path as a string

        :Example:
        
    >>> # Caution, for this example to work a .pdf example file will be downloaded from automagica.com FTP
    >>> example_pdf = download_file_from_url('http://automagica.com/examples/example_document.pdf')
    >>> # Download the watermark
    >>> example_watermark = download_file_from_url('http://automagica.com/examples/approved_stamp.pdf')
    >>> # Apply the watermark
    >>> watermarked_file = apply_watermark_to_pdf(example_pdf, example_watermark)
    >>> # Open the file for illustration
    >>> open_file(watermarked_file)

    Keywords
        PDF, extract images, images, extract text, PDF file, image

    Icon
        las la-stamp
    """
    from PyPDF2 import PdfFileWriter, PdfFileReader
    import os

    if not output_path:
        base, file_extension = os.path.splitext(file_path)
        output_path = base + "_watermarked_" + file_extension

    watermark = PdfFileReader(open(watermark_path, "rb"))

    input_file = PdfFileReader(open(file_path, "rb"))

    page_count = input_file.getNumPages()

    output_file = PdfFileWriter()

    for page_number in range(page_count):
        input_page = input_file.getPage(page_number)
        input_page.mergePage(watermark.getPage(0))
        output_file.addPage(input_page)

    with open(output_path, "wb") as outputStream:
        output_file.write(outputStream)

    return output_path


"""
System Monitoring
Icon: las la-wave-square
"""


@activity
def get_cpu_load(measure_time=1):
    """CPU load

    Get average CPU load for all cores.

    :parameter measure_time: Time (seconds) to measure load. Standard measure_time is 1 second.
    :type measure_time: int, optional

    :return: Displayed load is an average over measured_time.

        :Example:

    >>> get_cpu_load()
    10.1

    Keywords
        cpu, load, cpuload

    Icon
        las la-microchip

    """
    import psutil

    cpu_measurements = []
    for _ in range(measure_time):
        cpu_measurements.append(psutil.cpu_percent(interval=1))
    return sum(cpu_measurements) / len(cpu_measurements)


@activity
def get_number_of_cpu(logical=True):
    """Count CPU

    Get the number of CPU's in the current system. 

    :parameter logical: Determines if only logical units are added to the count, default value is True.
    :type logical: bool, optional

    :return: Number of CPU Integer

        :Example:

    >>> get_number_of_cpu()
    2

    Keywords
        cpu, count, number of cpu

    Icon
        las la-calculator

    """
    import psutil

    return psutil.cpu_count(logical=logical)


@activity
def get_cpu_frequency():
    """CPU frequency

    Get frequency at which CPU currently operates.

    :return: minimum and maximum frequency

        :Example:

    >>> get_cpu_frequency()
    scpufreq(current=3600.0, min=0.0, max=3600.0)

    Keywords
        cpu, load, cpu frequency

    Icon
        las la-wave-square

    """
    import psutil

    return psutil.cpu_freq()


@activity
def get_cpu_stats():
    """CPU Stats

    Get CPU statistics

    :return: Number of CTX switches, intterupts, soft-interrupts and systemcalls.

        :Example:

    >>> get_cpu_stats()
    scpustats(ctx_switches=735743826, interrupts=1540483897, soft_interrupts=0, syscalls=2060595131)

    Keywords
        cpu, load, cpu frequency, stats, cpu statistics

    Icon
        las la-server
    """
    import psutil

    return psutil.cpu_stats()


@activity
def get_memory_stats(mem_type="swap"):
    """Memory statistics

    Get  memory statistics

    :parameter mem_type: Choose mem_type = 'virtual' for virtual memory, and mem_type = 'swap' for swap memory (standard).
    :options mem_type: ['virtual', 'swap']

    :return: Total, used, free and percentage in use.

        :Example:

    >>> get_memory_stats()
    sswap(total=24640016384, used=18120818688, free=6519197696, percent=73.5, sin=0, sout=0)

    Keywords
        memory, statistics, usage, ram

    Icon
        las la-memory

    """
    import psutil

    if mem_type == "virtual":
        return psutil.virtual_memory()
    else:
        return psutil.swap_memory()


@activity
def get_disk_stats():
    """Disk stats

    Get disk statistics of main disk

    :return: Total, used, free and percentage in use.

        :Example:

    >>> get_disk_stats()
    sdiskusage(total=999559262208, used=748696350720, free=250862911488, percent=74.9)

    Keywords
        disk usage, disk stats, disk, harddisk, space

    Icon
        las la-save
    """
    import psutil

    return psutil.disk_usage("/")


@activity
def get_disk_partitions():
    """Partition info

    Get disk partition info

    :return: tuple with info for every partition.

        :Example:

    >>> get_disk_paritions()
    [sdiskpart(device='C:\\', mountpoint='C:\\', fstype='NTFS', opts='rw,fixed')]

    Keywords
        disk usage, disk stats, disk, harddisk, space

    Icon
        las la-save

    """
    import psutil

    return psutil.disk_partitions()


@activity
def get_boot_time():
    """Boot time

    Get most recent boot time

    :return: time PC was booted in seconds after the epoch.

        :Example:

    >>> get_boot_time()
    123456789.0

    Keywords
        boot, boot time, boottime, startup, timer

    Icon
        lar la-clock
    """
    import psutil

    return psutil.boot_time()


@activity
def get_time_since_last_boot():
    """Uptime

    Get uptime since last boot

    :return: time since last boot in seconds.

        :Example:

    >>> get_time_since_last_boot()
    1337.0

    Keywords
        boot, boot time, boottime, startup, timer

    Icon
        lar la-clock
    """
    import time
    import psutil

    return time.time() - psutil.boot_time()


"""
Image Processing
Icon: las la-photo-video
"""


@activity
def show_image(file_path):
    """Show image

    Displays an image specified by the path variable on the default imaging program.

    :parameter file_path: Path to image
    :type file_path: input_file

        :Example:

    >>> # Take screenshot of current screen to use as test image
    >>> testimage = take_screenshot()
    >>> # Show the image
    >>> show_image(testimage)

    Keywords
        image, show image, reveal, open image, open

    Icon
        las la-images

    """
    from PIL import Image

    file_path = interpret_path(file_path)
    im = Image.open(file_path)

    return im.show()


@activity
def rotate_image(file_path, angle=90):
    """Rotate image

    Rotate an image

    :parameter file_path: Path to image
    :type file_path: input_file
    :parameter angle: Degrees to rotate image. Note that angles other than 90, 180, 270, 360 can resize the picture. 
    :type angle: int, optional

        :Example:

    >>> # Take screenshot of current screen to use as test image
    >>> testimage = take_screenshot()
    >>> # Rotate the image
    >>> rotate_image(testimage)
    >>> # Show the image
    >>> show_image(testimage)

    Keywords
        image, rotate image, 90 degrees, image manipulation, photoshop, paint

    Icon
        las la-undo

    """
    from PIL import Image

    file_path = interpret_path(file_path)
    im = Image.open(file_path)

    return im.rotate(angle, expand=True).save(file_path)


@activity
def resize_image(file_path, size):
    """Resize image

    Resizes the image specified by the path variable. 

    :parameter file_path: Path to the image
    :type file_path: input_file
    :parameter size: Tuple with the width and height in pixels. E.g.  (300, 400) gives the image a width of 300 pixels and a height of 400 pixels.
    :type size: tuple

        :Example:

    >>> # Take screenshot of current screen to use as test image
    >>> testimage = take_screenshot()
    >>> # Resize the image
    >>> resize_image(testimage, size=(100,100))
    >>> # Show the image
    >>> show_image(testimage)

    Keywords
        image, resize image, resize, size, image manipulation, photoshop, paint

    Icon
        las la-expand-arrows-alt
    """
    from PIL import Image

    file_path = interpret_path(file_path)
    im = Image.open(file_path)

    return im.resize(size).save(file_path)


@activity
def get_image_width(file_path):
    """Get image width

    Get with of image

    :parameter file_path: Path to image
    :type file_path: input_file

        :Example:

    >>> # Take screenshot of current screen to use as test image
    >>> testimage = take_screenshot()
    >>> # get image height
    >>> get_image_width(testimage)
    1000

    Keywords
        image, height, width, image height, image width

    Icon
        las la-expand-arrows-alt
    """
    from PIL import Image

    file_path = interpret_path(file_path)
    im = Image.open(file_path)

    width, _ = im.size

    return width


@activity
def get_image_height(file_path):
    """Get image height

    Get height of image

    :parameter file_path: Path to image
    :type file_path: input_file

    :return: Height of image

        :Example:

    >>> # Take screenshot of current screen to use as test image
    >>> testimage = take_screenshot()
    >>> # get image height
    >>> get_image_height(testimage)
    1000

    Keywords
        image, height, width, image height, image width

    Icon
        las la-arrows-alt-v

    """
    from PIL import Image

    file_path = interpret_path(file_path)
    im = Image.open(file_path)

    _, height = im.size

    return height


@activity
def crop_image(file_path, box=None):
    """Crop image

        Crops the image specified by path to a region determined by the box variable.

    :parameter file_path: Path to image
    :type file_path: input_file
    :parameter box:  A tuple that defines the left, upper, right and lower pixel coördinate e.g.: (left, upper, right, lower)
    :type box: tuple

        :Example:

    >>> # Take screenshot of current screen to use as test image
    >>> testimage = take_screenshot()
    >>> # Crop the image
    >>> crop_image(testimage, box = (10,10,100,100))
    >>> # Show the image
    >>> show_image(testimage)

    Keywords
        image, crop, crop image

    Icon
        las la-crop
    """
    from PIL import Image

    file_path = interpret_path(file_path)
    im = Image.open(file_path)

    return im.crop(box).save(file_path)


@activity
def mirror_image_horizontally(file_path):
    """Mirror image horizontally

    Mirrors an image with a given path horizontally from left to right.

    :parameter file_path: Path to image
    :type file_path: input_file

        :Example:

    >>> # Take screenshot of current screen to use as test image
    >>> testimage = take_screenshot()
    >>> # Mirror image horizontally
    >>> mirror_image_horizontally(testimage)
    >>> # Show the image
    >>> show_image(testimage)

    Keywords
        image, flip, flip image, mirror, mirror image, horizon, horizontally

    Icon
        las la-caret-up
    """
    from PIL import Image

    file_path = interpret_path(file_path)
    im = Image.open(file_path)
    return im.transpose(Image.FLIP_LEFT_RIGHT).save(file_path)


@activity
def mirror_image_vertically(file_path):
    """Mirror image vertically

    Mirrors an image with a given path vertically from top to bottom.

    :parameter file_path: Path to image
    :type file_path: input_file

        :Example:

    >>> # Take screenshot of current screen to use as test image
    >>> testimage = take_screenshot()
    >>> # Mirror image vertically
    >>> mirror_image_vertically(testimage)
    >>> # Show the image
    >>> show_image(testimage)

    Keywords
        image, flip, flip image, mirror, mirror image, vertical, vertically

    Icon
        las la-caret-right
    """
    from PIL import Image

    file_path = interpret_path(file_path)
    im = Image.open(file_path)

    return im.transpose(Image.FLIP_TOP_BOTTOM).save(file_path)


"""
Process
Icon: las la-play
"""


@activity
def run_manual(task):
    """Windows run

    Use Windows Run to boot a process
    Note this uses keyboard inputs which means this process can be disrupted by interfering inputs

    :parameter task: Name of the task to run e.g. 'mspaint.exe'
    :type task: string

        :Example:

    >>> # Open paint with Windows run
    >>> run_manual('mspaint.exe')
    >>> # Open home directory with Windows run
    >>> run_manual(home_path())

    Keywords
        run, open, task, win r, windows run, shell, cmd

    Icon
        las la-cog
    """

    import time

    press_key_combination("win", "r")
    time.sleep(0.5)

    import platform

    # Set keyboard layout for Windows platform
    if platform.system() == "Windows":
        from win32api import LoadKeyboardLayout

        LoadKeyboardLayout("00000409", 1)

    type_text(task)
    press_key("enter")


@activity
def run(process):
    """Run process

    Use subprocess to open a windows process

    :parameter process: Process to open e.g: 'calc.exe', 'notepad.exe', 'control.exe', 'mspaint.exe'.
    :type process: string

        :Example:

    >>> # Open paint with Windows run
    >>> run('mspaint.exe')

    Keywords
        run, open, task, win r, windows run, shell, cmd

    Icon
        las la-play
    """
    import subprocess  # nosec

    subprocess.Popen(process)


@activity
def is_process_running(name):
    """Check if process is running

    Check if process is running. Validates if given process name (name) is currently running on the system.

    :parameter name: Name of process
    :type name: string

    :return: Boolean

        :Example:

    >>> # Open paint with Windows run
    >>> run('mspaint.exe')
    >>> # Check if paint is running
    >>> is_process_running('mspaint.exe')
    True

    Keywords
        run, open, task, win r, windows run, shell, cmd

    Icon
        las la-cogs
    """
    import psutil

    if name:
        for p in psutil.process_iter():
            if name in p.name():
                return True

    return False


@activity
def get_running_processes():
    """Get running processes

    Get names of unique processes currently running on the system.

    :return: List of unique running processes

         :Example:

    >>> # Show all running processes
    >>> get_running_processes()
    ['cmd.exe', 'chrome.exe', ... ]

    Keywords
        process, processes, list processes, running, running processes

    Icon
        las la-list
    """
    import psutil

    process_list = []

    for p in psutil.process_iter():
        process_list.append(p.name())

    return list(set(process_list))


@activity
def kill_process(name=None):
    """Kill process

    Kills a process forcefully

    :parameter name: Name of the process
    :type name: string

        :Example:

    >>> # Open paint with Windows run
    >>> run('mspaint.exe')
    >>> # Force paint to close
    >>> kill_process('mspaint.exe')


    Keywords
        run, open, task, win r, windows run, shell, cmd, kill, stop, kill process, stop process, quit, exit

    Icon
        las la-window-close
    """
    import os

    return os.system(f"taskkill /f /im {name} >nul 2>&1")  # nosec


"""
Optical Character Recognition (OCR)
Icon: las la-glasses
"""


@activity
def extract_text_ocr(file_path=None):
    """Get text with OCR

    This activity extracts all text from the current screen or an image if a path is specified.

    :parameter file_path: Path to image from where text will be extracted. If no path is specified a screenshot of current screen will be used.
    :type file_path: input_file

    :return: String with all text from current screen

        :Example:

    >>> # Make a text file with some text to recognize
    >>> testfile = make_text_file(text='OCR Example')
    >>> # Open the text file
    >>> open_file(testfile)
    >>> # Find the text with OCR
    >>> extracted_text = find_text_on_screen_ocr(text='OCR Example')
    >>> # Check if the extracted_text contains the original word
    >>> 'OCR Example' in extracted_text
    True

    Keywords
        OCR, vision, AI, screen, citrix, read, optical character recognition

    Icon
        lab la-readme
    """

    from automagica.httpclient import http_client
    import base64
    import os
    import json

    if not file_path:
        file_path = interpret_path(file_path, default_filename="ocr_temp.jpg")
        import PIL.ImageGrab

        img = PIL.ImageGrab.grab()
        img.save(file_path, "JPEG")
    else:
        file_path = interpret_path(file_path)

    # Open file and encode as Base 64
    with open(file_path, "rb") as f:
        image_base64 = base64.b64encode(f.read()).decode("utf-8")

    # Get Bot API_key
    config_path = os.path.join(os.path.expanduser("~"), "automagica.json")

    # Read JSON
    with open(config_path) as json_file:
        local_data = json.load(json_file)
        api_key = str(local_data.get("bot_secret"))  # Your API key

    # Prepare data for request
    data = {"image_base64": image_base64, "api_key": api_key}

    # Post request to API
    url = (
        os.environ.get("AUTOMAGICA_PORTAL_URL", "https://portal.automagica.com")
        + "/api/ocr/find-text-locations"
    )

    r = http_client.post(url, json=data)

    # Print results
    return r.json()["text"]


@activity
def find_text_on_screen_ocr(text, criteria=None):
    """Find text on screen with OCR

    This activity finds position (coordinates) of specified text on the current screen using OCR.

    :parameter text: Text to find. Only exact matches are returned.
    :type text: string
    :parameter criteria: Criteria to select on if multiple matches are found. If no criteria is specified all matches will be returned. Options are 'first', which returns the first match closest to the upper left corner, 'last' returns the last match closest to the lower right corner, random selects a random match.
    :options criteria: ['first', 'last', 'random']

    :return: Dictionary or list of dictionaries with matches with following elements: 'h' height in pixels, 'text' the matched text,'w' the width in pixels, 'x' absolute x-coördinate , 'y' absolute y-coördinate. Returns nothing if no matches are found

        :Example:

    >>> # Make a text file with some text to recognize
    >>> testfile = make_text_file(text='OCR Example')
    >>> # Open the text file
    >>> open_file(testfile)
    >>> # Find the text with OCR
    >>> find_text_on_screen_ocr(text='OCR Example')

    Keywords
        OCR, vision, AI, screen, citrix, read, optical character recognition

    Icon
        las la-glasses

    """

    from automagica.httpclient import http_client
    import base64
    import os
    import json

    import PIL.ImageGrab

    img = PIL.ImageGrab.grab()
    path = os.path.join(os.path.expanduser("~"), "ocr_capture.jpg")
    img.save(path, "JPEG")

    # Open file and encode as Base 64
    with open(path, "rb") as f:
        image_base64 = base64.b64encode(f.read()).decode("utf-8")

    # Get Bot API_key
    config_path = os.path.join(os.path.expanduser("~"), "automagica.json")

    # Read JSON
    with open(config_path) as json_file:
        local_data = json.load(json_file)
        api_key = str(local_data.get("bot_secret"))  # Your API key

    # Prepare data for request
    data = {"image_base64": image_base64, "api_key": api_key}

    # Post request to API
    url = (
        os.environ.get("AUTOMAGICA_PORTAL_URL", "https://portal.automagica.com")
        + "/api/ocr/find-text-locations"
    )

    r = http_client.post(url, json=data)

    # Print results
    data = r.json()["locations"]

    # Find all matches
    matches = []
    for item in data:
        if item["text"].lower() == text.lower():
            matches.append(item)

    if not matches:
        return None

    if criteria:
        if len(matches) > 0:
            if criteria == "first":
                best_match = matches[0]
            if criteria == "last":
                best_match = matches[-1]
            if criteria == "random":
                import random

                best_match = random.choice(matches)

            return best_match

    else:
        return matches


@activity
def click_on_text_ocr(text, delay=1):
    """Click on text with OCR

    This activity clicks on position (coordinates) of specified text on the current screen using OCR.

    :parameter text: Text to find. Only exact matches are returned.
    :type text: string
    :parameter delay: Delay before clicking in seconds
    :type delay: int, optional

        :Example:

    >>> # Make a text file with some text to recognize
    >>> testfile = make_text_file(text='OCR Example')
    >>> # Open the text file
    >>> open_file(testfile)
    >>> # Find the text with OCR and click on it
    >>> click_on_text_ocr(text='OCR Example')

    Keywords
        OCR, vision, AI, screen, citrix, read, optical character recognition, click

    Icon
        las la-mouse-pointer
    """
    if delay:
        from time import sleep

        sleep(delay)

    position = find_text_on_screen_ocr(text, criteria="first")
    if position:

        x = int(position["x"] + position["w"] / 2)
        y = int(position["y"] + position["h"] / 2)

        from mouse import move, click

        move(x, y)
        click()

        return


@activity
def double_click_on_text_ocr(text, delay=1):
    """Double click on text with OCR

    This activity double clicks on position (coordinates) of specified text on the current screen using OCR.

    :parameter text: Text to find. Only exact matches are returned.
    :type text: string
    :parameter delay: Delay before clicking in seconds
    :type delay: int, optional

        :Example:

    >>> # Make a text_file with some text to recognize
    >>> testfile = make_text_file(text='OCR Example')
    >>> # Open the text file
    >>> open_file(testfile)
    >>> # Find the text with OCR and double click on it
    >>> double_click_on_text_ocr(text='OCR Example')

    Keywords
        OCR, vision, AI, screen, citrix, read, optical character recognition, click, double click

    Icon
        las la-mouse-pointer

    """
    if delay:
        from time import sleep

        sleep(delay)

    position = find_text_on_screen_ocr(text, criteria="first")
    if position:

        x = int(position["x"] + position["w"] / 2)
        y = int(position["y"] + position["h"] / 2)

        from mouse import move, double_click

        move(x, y)
        double_click()
        return


@activity
def right_click_on_text_ocr(text, delay=1):
    """Right click on text with OCR

    This activity Right clicks on position (coordinates) of specified text on the current screen using OCR.

    :parameter text: Text to find. Only exact matches are returned.
    :type text: string
    :parameter delay: Delay before clicking in seconds
    :type delay: int, optional

        :Example:

    >>> # Make a text file with some text to recognize
    >>> testfile = make_text_file(text='OCR Example')
    >>> # Open the text file
    >>> open_file(testfile)
    >>> # Find the text with OCR and right click on it
    >>> right_click_on_text_ocr(text='OCR Example')

    Keywords
        OCR, vision, AI, screen, citrix, read, optical character recognition, click, right click

    Icon
        las la-mouse-pointer
    """

    if delay:
        from time import sleep

        sleep(delay)

    position = find_text_on_screen_ocr(text, criteria="first")
    if position:

        x = int(position["x"] + position["w"] / 2)
        y = int(position["y"] + position["h"] / 2)

        from mouse import move, right_click

        move(x, y)
        right_click()

        return


"""
UiPath
Icon: las la-robot
"""


@activity
def execute_uipath_process(project_file_path, arguments=None, uirobot_exe_path=None):
    """Execute a UiPath process

    This activity allows you to execute a process designed with the UiPath Studio. All console output from the Write Line activity (https://docs.uipath.com/activities/docs/write-line) will be printed as output.

    :parameter project_file_path: path to the project file (as created within the UiPath Studio)
    :type project_file_path: input_file
    :parameter arguments: dictionary with input arguments/parameters for the process to use in UiPath (optional)
    :type arguments: string, optional
    :parameter uirobot_exe_path: path to UiPath's UiRobot.exe (optional)
    :type uirobot_exe_path: input_file, optional
    :extension uirobot_exe_path: exe
    

        :Example:

    >>> # Run a UiPath process
    >>> arguments = {'firstname': 'John', 'lastname': 'Doe'}
    >>> execute_uipath_process(r"C:\\Processes UiPath\\my_process.xaml", arguments=arguments)
    Completed UiPath process "C:\\Processes UiPath\\my_process.xaml"

    Keywords
        RPA, UiPath, Studio, robot, orchestrator, xaml, ui path

    Icon
        las la-robot
    """
    only_supported_for("Windows")

    import subprocess  # nosec
    import json

    project_file_path = interpret_path(project_file_path)

    if not uirobot_exe_path:
        uirobot_exe_path = r"C:\Program Files (x86)\UiPath\Studio\UiRobot.exe"
    else:
        uirobot_exe_path = interpret_path(uirobot_exe_path)

    cmd = ' -f "{}"'.format(project_file_path)

    if arguments:
        cmd += ' --input "{}"'.format(json.dumps(arguments))

    uirobot_exe_path = '"' + uirobot_exe_path + '"'

    process = subprocess.Popen(uirobot_exe_path + cmd)

    out, err = process.communicate()

    if out:
        print("Output:")
        print(out)

    if err:
        print("Errors:")
        print(err)

    print('Completed UiPath process "{}"'.format(project_file_path))


"""
AutoIt 
Icon: las la-robot
"""


@activity
def run_autoit_script(script_path, arguments=None, autoit_exe_path=None):
    """Execute a AutoIt script

    This activity allows you to run an AutoIt script. If you use the ConsoleWrite function (https://www.autoitscript.com/autoit3/docs/functions/ConsoleWrite.htm), the output will be presented to you.

    :parameter script_path: path to the '.au3' script file
    :type script_path: input_file
    :extension script_path: au3
    :parameter arguments: string with input arguments/parameters for the script (optional)
    :type arguments: string, optional
    :parameter autoit_exe_path: path to AutoIt.exe (optional)
    :type autoit_exe_path: input_file, optional
    :extension autoit_exe_path: exe
    

        :Example:

    >>> # Run an AutoIt script
    >>> arguments = 'John'
    >>> run_autoit_script(r"C:\\AutoIt\\Scripts\\MyScript.au3", arguments=arguments) # Point this to your AutoIt Script
    Completed AutoIt script "C:\\AutoIt\\Scripts\\MyScript.au3"

    Keywords
        RPA, AutoIt, au3, au

    Icon
        las la-robot
    """
    only_supported_for("Windows")

    import subprocess  # nosec  # nosec
    import json

    script_path = interpret_path(script_path)

    if not autoit_exe_path:
        autoit_exe_path = r"C:\Program Files (x86)\AutoIt3\AutoIt3_x64.exe"
    else:
        autoit_exe_path = interpret_path(autoit_exe_path)

    cmd = ' "{}"'.format(script_path)

    if arguments:
        cmd = +' "{}"'.format(json.dumps(arguments))

    autoit_exe_path = '"' + autoit_exe_path + '"'

    process = subprocess.Popen(autoit_exe_path + cmd)

    out, err = process.communicate()

    if out:
        print("Output:")
        print(out)

    if err:
        print("Errors:")
        print(err)

    print('Completed AutoIt script "{}"'.format(script_path))


"""
Alternative frameworks
Icon: las la-robot
"""


@activity
def execute_robotframework_test(test_case_path, variables=None):
    """Execute a Robot Framework test case

    This activity allows you to run a Robot Framework test case. Console output of the test case will be printed.

    :parameter test_case_path: path to the '.robot' test case file
    :type test_case_path: input_file
    :extension test_case_path: robot
    :parameter variables: dictionary with variable declarations
    :type variables: string, optional

        :Example:

    >>> # Run an Robot Framework test case
    >>> variables = {'FIRSTNAME': 'John', 'LASTNAME': 'Doe'}
    >>> execute_robotframework_test(r"C:\\Test Cases\\my_test_case.robot", variables=variables) # Point this to your Robot Framework test case
    Completed Robot Framework test case "C:\\Test Cases\\my_test_case.robot"

    Keywords
        RPA, robot framework, robotframework, robot

    Icon
        las la-robot
    """

    import subprocess  # nosec  # nosec  # nosec
    import json

    test_case_path = interpret_path(test_case_path)

    cmd = ' "{}"'.format(test_case_path)

    if variables:
        variables_parameter = " --variables ".join(
            ["{}:{}".format(key, value) for key, value in variables.items()]
        )
        cmd = +variables_parameter

    process = subprocess.Popen("robot" + cmd)

    out, err = process.communicate()

    if out:
        print("Output:")
        print(out)

    if err:
        print("Errors:")
        print(err)

    print('Completed Robot Framework test case "{}"'.format(test_case_path))


@activity
def run_blueprism_process(
    process_name,
    username="",
    password="",
    sso=False,
    inputs=None,
    automatec_exe_path=None,
):
    """Run a Blue Prism process

    This activity allows you to run a Blue Prism process.

    :parameter process_name: name of the process in Blue Prism
    :type process_name: string
    :parameter username: Blue Prism username
    :type username: string, optional
    :parameter password: Blue Prism password
    :type password: string, optional
    :parameter sso: Run as single-sign on user with Blue Prism
    :type sso: bool, optional
    :parameter inputs: dictionary with inputs declarations (optional)
    :type inputs: string, optional
    :parameter automatec_exe_path: path to Blue Prism's AutomateC.exe (optional)
    :type automatec_exe_path: input_file
    :extension automatec_exe_path: exe
    

        :Example:

    >>> # Run a Blue Prism process
    >>> inputs = {'firstname': 'John', 'lastname': 'Doe'}
    >>> run_blueprism_process("My Example Process", username="user", password="password", inputs=inputs)
    Completed Blue Prism process "My Example Process"

    Keywords
        RPA, blueprism, blue prism, robot

    Icon
        las la-robot
    """
    only_supported_for("Windows")

    import subprocess  # nosec  # nosec
    import json

    cmd = ' /run "{}"'.format(process_name)

    if not sso:
        cmd += " /user {} {}".format(username, password)
    else:
        cmd += " /sso"

    if inputs:
        inputs_parameters = "".join(
            [
                "<input name='{}' type='text' value='{}' /></inputs>".format(key, value)
                for key, value in inputs.items()
            ]
        )
        cmd = +" " + inputs_parameters

    if automatec_exe_path:
        automatec_exe_path = interpret_path(automatec_exe_path)

    automatec_exe_path = '"' + automatec_exe_path + '"'

    process = subprocess.Popen(automatec_exe_path + cmd)

    out, err = process.communicate()

    if out:
        print("Output:")
        print(out)

    if err:
        print("Errors:")
        print(err)

    print('Completed Blue Prism process "{}"'.format(test_case_path))


@activity
def run_automationanywhere_task(task_file_path, aaplayer_exe_path=None):
    """Run an Automation Anywhere task

    This activity allows you to run an Automation Anywhere task.

    :parameter task_file_path: path to the task file of Automation Anywhere
    :type task_file_path: input_file
    :parameter aaplayer_exe_path: path to the AAPlayer.exe (optional)
    :type aaplayer_exe_path: input_file
    :extension aaplayer_exe_path: exe

        :Example:

    >>> # Run an Automation Anywhere task
    >>> run_automationanywhere_task(r"C:\AutomationAnywhereTasks\MyTask.atmx")
    Completed Automation Anywhere task "C:\AutomationAnywhereTasks\MyTask.atmx"

    Keywords
        RPA, automation anywhere, aa, robot

    Icon
        las la-robot
    """
    only_supported_for("Windows")

    import subprocess  # nosec   # nosec
    import json

    task_file_path = interpret_path(task_file_path)
    if aaplayer_exe_path:
        aaplayer_exe_path = interpret_path(aaplayer_exe_path)

    cmd = ' "/f{}/e"'.format(task_file_path)

    aaplayer_exe_path = '"' + aaplayer_exe_path + '"'

    process = subprocess.Popen(aaplayer_exe_path + cmd)

    out, err = process.communicate()

    if out:
        print("Output:")
        print(out)

    if err:
        print("Errors:")
        print(err)

    print('Completed Automation Anywhere task "{}"'.format(test_case_path))


"""
General
Icon: las la-briefcase
"""


@activity
def raise_exception(message="Exception", exception=Exception):
    """Raise exception

    Raises an exception

    :parameter message: Message
    :type: string, optional
    :parameter exception: Exception to raise
    :type exception: exception, optional

    Keywords
        sap, sap gui, sap client
    
    Icon
        las la-exclamation
    """
    raise Exception(message)


"""
SAP GUI
Icon: las la-briefcase
"""


class SAPGUI:
    def __init__(self, sap_logon_exe_path=None, delay=1):
        """Start SAP GUI

        For this activity to work, SAP GUI needs to be installed on the system.

        :parameter sap_logon_exe_path: Specifiy the installation location of the saplogon.exe if not at the default location.
        :type sap_logon_exe_path: input_file, optional
        :extension sap_logon_exe_path: exe
        :parameter delay: Number of seconds to wait between tries for attaching to the SAP process
        :type delay: int, optional

            :Example:

        >>> # Log in to SAP GUI
        >>> sap = SAPGUI()
        >>> sap.login('System', '001', 'username', 'password')

        Keywords
            sap, sap gui, sap client
        
        Icon
            las la-briefcase
        """
        from subprocess import Popen
        import win32com.client
        from time import sleep

        # Run SAP process
        if not sap_logon_exe_path:
            self.sap_logon_exe_path = (
                r"C:\Program Files (x86)\SAP\FrontEnd\SAPgui\saplogon.exe"
            )
        else:
            self.sap_logon_exe_path = interpret_path(sap_logon_exe_path)

        self.process = Popen(self.sap_logon_exe_path)

        # Try to connect to SAP GUI
        for _ in range(10):
            try:
                self.sapgui = win32com.client.GetObject("SAPGUI").GetScriptingEngine
                break
            except:
                sleep(delay)
        else:
            raise Exception(
                "Could not connect to the SAP GUI. Did you enable scripting in the SAP GUI?"
            )

    @activity
    def quit(self):
        """Quit SAP GUI

        Quits the SAP GUI completely and forcibly.

            :Example:
            
        >>> # Log in to SAP GUI
        >>> sap = SAPGUI()
        >>> sap.login('System', '001', 'username', 'password')
        >>> # Quit SAP
        >>> sap.quit()

        Keywords
            sap, sap gui, sap client, quit
        
        Icon
            las la-briefcase
        """
        self.process.kill()

    @property
    def connections(self):
        """Returns connections for SAP GUI
        """
        connections = []

        for connection_id in range(0, self.sapgui.Children.Count):
            connections.append(self.sapgui.Children(connection_id))

        return connections

    @activity
    def login(self, environment, client, username, password, force=True):
        """Log in to SAP GUI

        Logs in to an SAP system on SAP GUI.

        :parameter environment: Environment
        :type environment: string
        :parameter client: Client
        :type client: string
        :parameter username: Username
        :type username: string
        :parameter password: Password
        :type password: string
        :parameter force: Force optional
        :type force: bool, optional

            :Example:
            
        >>> # Log in to SAP GUI
        >>> sap = SAPGUI()
        >>> sap.login('System', '001', 'username', 'password')

        Keywords
            sap, sap gui, sap client, login
        
        Icon
            las la-briefcase
        """
        # Open the connection window
        self.sapgui.OpenConnection(environment, True)

        # Identify SAP session
        self.session = self.sapgui.FindById("ses[0]")

        # Log in to SAP
        self.session.findById("wnd[0]/usr/txtRSYST-MANDT").text = client
        self.session.findById("wnd[0]/usr/txtRSYST-BNAME").text = username
        self.session.findById("wnd[0]/usr/pwdRSYST-BCODE").text = password
        self.session.findById("wnd[0]").sendVKey(0)

        # Continue even if other logged in sessions detected
        if force:
            try:
                self.session.findById("wnd[1]/usr/radMULTI_LOGON_OPT2").select()
                self.session.findById("wnd[1]/usr/radMULTI_LOGON_OPT2").setFocus()
                self.session.findById("wnd[1]/tbar[0]/btn[0]").press()
            except:
                pass

    @activity
    def click_sap(self, identifier):
        """Click on a SAP GUI element

        Clicks on an identifier in the SAP GUI.

        :parameter identifier: Technical identifier of the element
        :type identifier: string

            :Example:
            
        >>> # Log in to SAP GUI
        >>> sap = SAPGUI()
        >>> sap.login('System', '001', 'username', 'password')
        >>> logout_button = '/app/con[0]/ses[0]/wnd[0]/tbar[0]/btn[15]'
        >>> sap.highlight(logout_button)
        >>> sap.click_sap(logout_button)

        Keywords
            sap, sap gui, sap client, click
        
        Icon
            las la-briefcase
        """
        self.sapgui.findById(identifier).Press()

    @activity
    def get_text(self, identifier):
        """Get text from a SAP GUI element

        Retrieves the text from a SAP GUI element.

        :parameter identifier: Technical identifier of the element
        :type identifier: string

            :Example:
            
        >>> # Log in to SAP GUI
        >>> sap = SAPGUI()
        >>> sap.login('System', '001', 'username', 'password')
        >>> status_bar = '/app/con[0]/ses[0]/wnd[0]/sbar/pane[0]'
        >>> sap.get_text(status_bar)

        Keywords
            sap, sap gui, sap client, get text
        
        Icon
            las la-briefcase
        """
        return self.sapgui.findById(identifier).text

    @activity
    def set_text(self, identifier, text):
        """Set text of a SAP GUI element

        Sets the text of a SAP GUI element.

        :parameter text: Text to set
        :type text: string
        :parameter identifier: Technical identifier of the element
        :type identifier: string

            :Example:
            
        >>> # Log in to SAP GUI
        >>> sap = SAPGUI()
        >>> sap.login('System', '001', 'username', 'password')
        >>> sap.set_text('/app/con[0]/ses[0]/wnd[0]/tbar[0]/okcd', 'Hello!')

        Keywords
            sap, sap gui, sap client, set text
        
        Icon
            las la-briefcase
        """
        self.sapgui.FindById(identifier).text = text

    @activity
    def highlight(self, identifier, duration=1):
        """Highlights a SAP GUI element

        Temporarily highlights a SAP GUI element

        :parameter identifier: Technical identifier of the element
        :type identifier: string
        :parameter duration: Duration of the highlight
        :type duration: int, optional

            :Example:
            
        >>> # Log in to SAP GUI
        >>> sap = SAPGUI()
        >>> sap.login('System', '001', 'username', 'password')
        >>> sap.highlight('/app/con[0]/ses[0]/wnd[0]/tbar[0]/okcd', 'Hello!')

        Keywords
            sap, sap gui, sap client, highlight
        
        Icon
            las la-briefcase
        """
        from time import sleep

        self.sapgui.FindById(identifier).Visualize(1)

        sleep(duration)
        self.sapgui.FindById(identifier).Visualize(0)


"""
Portal
Icon: las la-robot
"""


@activity
def create_new_job_in_portal(
    process_name, process_version_id=None, priority=0, parameters=None
):
    """Create a new job in the Automagica Portal

    This activity creates a new job in the Automagica Portal for a given process. The bot performing this activity needs to be in the same team as the process it creates a job for.

    :parameter process_name: name of the process
    :type process_name: string
    :parameter process_version_id: id of a specific version of the process, if not provided it will use the latest version (optional)
    :type process_version_id: string, optional
    :parameter priority: priority level of the process. higher priority levels are performed first. (optional)
    :type priority: int, optional
    :parameter parameters: parameters for the process (optional)
    :type parameters: text, optional

        :Example:

    >>> # Create a job in the Automagica Portal
    >>> create_new_job_in_portal('My process')
    Job 1234567890 created

    Keywords
        queueing, process, job, create job, new job

    Icon
        las la-robot
    """
    from automagica.httpclient import http_client
    import os
    import json

    # Get Bot API_key
    config_path = os.path.join(os.path.expanduser("~"), "automagica.json")

    # Read JSON
    with open(config_path) as json_file:
        local_data = json.load(json_file)
        bot_secret = str(local_data["bot_secret"])

    headers = {"bot_secret": bot_secret, "process": process_name}

    if process_version_id:
        headers["version_id"] = process_version_id

    data = {}

    if priority:
        data["priority"] = priority

    if parameters:
        data["parameters"] = parameters

    r = http_client.post(
        os.environ.get("AUTOMAGICA_PORTAL_URL", "https://portal.automagica.com")
        + "/api/job/new",
        json=data,
        headers=headers,
    )

    try:

        result = r.json()

    except:

        raise Exception("Could not create job in Portal for unknown reason.")

    if result.get("error"):

        raise Exception(result["error"])

    else:
        print(result["message"])


@activity
def get_credential_from_portal(credential_name):
    """Get a credential from the Automagica Portal

    This activity retrieves a credential from the Automagica Portal.

    :parameter credential_name: name of the credential
    :type credential_name: string

    :return: Credential

        :Example:

    >>> # Get a credential from the Portal
    >>> print(get_credential_from_portal('My credential'))
    'secretpassword'

    Keywords
        password, credential, portal, login, username

    Icon
        las la-key
    """
    from automagica.httpclient import http_client
    import os
    import json

    # Get Bot API_key
    config_path = os.path.join(os.path.expanduser("~"), "automagica.json")

    # Read JSON
    with open(config_path) as json_file:
        local_data = json.load(json_file)
        bot_secret = str(local_data["bot_secret"])

    headers = {"bot_secret": bot_secret}

    data = {"name": credential_name}

    r = http_client.post(
        os.environ.get("AUTOMAGICA_PORTAL_URL", "https://portal.automagica.com")
        + "/api/credential/get",
        json=data,
        headers=headers,
    )

    try:
        result = r.json()
    except:

        raise Exception("Could not get credential from Portal for unknown reason.")

    if result.get("error"):
        raise Exception(result["error"])

    else:
        return result["contents"]


"""
Vision
Icon: las la-eye
"""


def get_screen_dimensions():
    """
    Returns primary screen width and height in pixels
    """
    import mss

    with mss.mss() as sct:

        # Find primary monitor
        for monitor in sct.monitors:
            if monitor["left"] == 0 and monitor["top"] == 0:
                break

    return monitor["width"], monitor["height"]


def capture_screen():
    """
    Captures the screen to a Pillow Image object
    """
    from PIL import Image
    import mss

    with mss.mss() as sct:

        # Find primary monitor
        for monitor in sct.monitors:
            if monitor["left"] == 0 and monitor["top"] == 0:
                break

        sct_img = sct.grab(monitor)

    img = Image.frombytes("RGB", sct_img.size, sct_img.bgra, "raw", "BGRX")

    return img


def insert_cell_below(content, type_="code"):
    """
    Inserts a cell in the current Jupyter Notebook below the currently
    activated cell
    """
    from IPython.display import Javascript, display

    javascript = """
    var cell;
    cell = Jupyter.notebook.insert_cell_below('{}');
    cell.set_text("{}");
    """.format(
        type_, content
    )

    if type_ == "markdown":
        javascript += """
        cell.execute();
        """

    display(Javascript(javascript))


def detect_vision(automagica_id, detect_target=True):
    from automagica.httpclient import http_client
    from io import BytesIO
    import os
    import base64
    import json

    screenshot = capture_screen()

    # Convert to base64
    buffered = BytesIO()
    screenshot.save(buffered, format="PNG")
    image_base64 = base64.b64encode(buffered.getvalue()).decode("utf-8")

    # Get Bot API_key
    config_path = os.path.join(os.path.expanduser("~"), "automagica.json")

    # Read JSON
    with open(config_path) as json_file:
        local_data = json.load(json_file)
        bot_secret = str(local_data.get("bot_secret"))

    data = {
        "bot_secret": bot_secret,  # Automagica Bot Secret
        "automagica_id": automagica_id,
        "image_base64": image_base64,  # Screenshot of the example screen
        "detect_target": detect_target,
    }

    portal_url = local_data.get("portal_url")

    if not portal_url:
        portal_url = "https://portal.automagica.com"

    url = portal_url + "/api/wand/detect"

    r = http_client.post(url, json=data)

    try:
        data = r.json()

    except Exception:
        raise Exception(
            "An unknown error occurred accessing the Automagica Portal API. Please try again later."
        )

    if data.get("error"):
        raise Exception(data["error"])

    return data["location"]


def get_center_of_rectangle(rectangle):
    """
    Returns center of rectangle in carthesian coordinate system
    """
    return (
        int((rectangle[0] + rectangle[2]) / 2),
        int((rectangle[1] + rectangle[3]) / 2),
    )


@activity
def is_visible(automagica_id, delay=1, timeout=30):
    """Check if element is visible on screen

    This activity can be used to check if a certain element is visible on the screen. 
    Note that this uses Automagica Portal and uses some advanced an fuzzy matching algorithms for finding identical elements.

    :parameter automagica_id: Element ID provided by the recorder
    :type automagica_id: automagica_id
    :parameter delay: Delay before checking visibility in seconds
    :type delay: int, optional
    :parameter timeout: Time before timeout (maximum time to wait) in seconds
    :type timeout: int, optional

    :return: True if visble, False if not

        :Example:

    >>> # Use the recorder to find an element ID
    >>> # Run the Windows calculator and try to perform the activity
    >>> run('calc.exe')
    >>> # Use the element ID found by the recorder, in this case ID 'qf41'. You can also view this on automagica.id/qf41 
    >>> #  If you have a vastly different version or layout the element might not be found, use the recorder 
    >>> is_visible('qf41')

    Keywords
        click, visible, is visible, appear,  computer vision, vision, AI

    Icon
        las la-eye
    """

    try:
        _ = detect_vision(automagica_id)
        return True
    except Exception:
        return False


@activity
def wait_appear(automagica_id, delay=1, timeout=30):
    """Wait for an element to appear

    Wait for an element that is defined the recorder

    :parameter automagica_id: The element ID provided by the recorder
    :type automagica_id: automagica_id
    :parameter delay: Delay before waiting to appear in seconds
    :type delay: int, optional
    :parameter timeout: Maximum time to wait for an element in seconds
    :type timeout: int, optional

    :return: Blocks while element not visible

        :Example:

    >>> # Use the recorder to find the element ID to wait for
    >>> # Run the Windows calculator and try to perform the activity
    >>> run('calc.exe')
    >>> # Use the element ID found by the recorder, in this case ID 'qf41'. You can also view this on automagica.id/qf41 
    >>> #  If you have a vastly different version or layout the element might not be found, use the recorder 
    >>> wait_appear('qf41')

    Keywords
        click, computer vision, vision, AI

    Icon
        las la-eye
    """
    from time import sleep

    sleep(delay)  # Default delay

    increment = 5

    for _ in range(int(timeout / increment)):
        try:
            _ = detect_vision(automagica_id)
            break
        except Exception:
            logging.exception()

        sleep(increment)

    else:
        raise Exception("Element did not appear within {} seconds".format(timeout))


@activity
def wait_vanish(automagica_id, delay=1, timeout=30):
    """Wait Vanish

    This activity allows the bot to wait for an element to vanish.

    :parameter automagica_id: The element ID provided by the recorder
    :type automagica_id: automagica_id
    :parameter delay: Delay before waiting for vanish in seconds
    :type delay: int, optional
    :parameter timeout: Maximum time to wait for an element in seconds
    :type timeout: int, optional

        :Example:

    >>> # Use the recorder to find the element ID for the vanishing element
    >>> # Run the Windows calculator and try to perform the activity
    >>> run('calc.exe')
    >>> # Use the element ID found by the recorder, in this case ID 'qf41'. You can also view this on automagica.id/qf41 
    >>> #  If you have a vastly different version or layout the element might not be found, use the recorder 
    >>> wait_vanish('qf41')

    Keywords
        wait, disappear, computer vision, vision, AI

    Icon
        las la-eye
    """
    from time import sleep

    sleep(delay)  # Default delay

    increment = 5

    for _ in range(int(timeout / increment)):
        try:
            _ = detect_vision(automagica_id)
        except Exception:
            break

        sleep(increment)

    else:
        raise Exception("Element did not disappear within {} seconds".format(timeout))


@activity
def read_text(automagica_id, delay=1):
    """Read Text with Automagica Wand

    This activity allows the bot to detect and read the text of an element by using the Automagica Portal API with a provided sample ID.

    :parameter automagica_id: the sample ID provided by Automagica Wand
    :type automagica_id: automagica_id
    :parameter delay: Delay before reading text for vanish in seconds
    :type delay: int, optional

    :return: Text

        :Example:

    >>> # Record an element to read with the recorder
    >>> # Run the Windows calculator and try to perform the activity
    >>> run('calc.exe')
    >>> # Use the element ID found by the recorder, in this case ID 'qf41'. You can also view this on automagica.id/qf41 
    >>> #  If you have a vastly different version or layout the element might not be found, use the recorder 
    >>> read_text('qf41')

    Keywords
        click, computer vision, vision, AI

    Icon
        las la-eye
    """
    from io import BytesIO
    from automagica.httpclient import http_client
    import base64
    import os
    import json

    from time import sleep

    sleep(delay)  # Default delay

    location = detect_vision(automagica_id, detect_target=False)

    screenshot = capture_screen()

    image = screenshot.crop(location)

    buffered = BytesIO()
    image.save(buffered, format="PNG")
    image_base64 = base64.b64encode(buffered.getvalue()).decode("utf-8")

    # Get Bot API_key
    config_path = os.path.join(os.path.expanduser("~"), "automagica.json")

    # Read JSON
    with open(config_path) as json_file:
        local_data = json.load(json_file)
        api_key = str(local_data.get("bot_secret"))  # Your API key

    # Prepare data for request
    data = {"image_base64": image_base64, "api_key": api_key}

    # Post request to API
    url = (
        os.environ.get("AUTOMAGICA_PORTAL_URL", "https://portal.automagica.com")
        + "/api/ocr/find-text-locations"
    )

    r = http_client.post(url, json=data)

    # Print results
    return r.json()["text"]
